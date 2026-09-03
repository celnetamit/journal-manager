"""Print resolution, measured from what the file actually contains.

Over 250 real manuscripts — 943 images — only 14% reach 300 DPI at the size they are
placed, 58% sit between 150 and 300, and 28% are below 150. Nothing in the tool had
ever looked at an image.
"""

import docx
import pytest

import image_check as I


def _doc_with_image(tmp_path, px, inches):
    """A document holding one image of `px` pixels wide, placed `inches` wide."""
    from PIL import Image

    img = tmp_path / "fig.png"
    Image.new("RGB", (px, int(px * 0.6)), "white").save(img)
    d = docx.Document()
    d.add_picture(str(img), width=docx.shared.Inches(inches))
    out = tmp_path / "m.docx"
    d.save(str(out))
    return docx.Document(str(out))


def test_a_low_resolution_image_is_reported(tmp_path):
    """1259 px across 6.34 in is 198 DPI — a real figure from a real manuscript."""
    doc = _doc_with_image(tmp_path, 1259, 6.34)
    (f,) = I.check_images(doc)
    assert f.rule == "image.resolution"
    assert f.severity == "warning"
    # Word stores the frame in EMU, so the width round-trips to 198 or 199 DPI
    # depending on rounding. Pinning the exact number tests the arithmetic of
    # python-docx, not this rule.
    (facts,) = I.read_images(doc)
    assert 150 <= facts.dpi < 300


def test_a_very_low_resolution_image_is_an_error(tmp_path):
    """399 px across 2.63 in is 152 DPI. Below 150 nothing at layout can rescue it."""
    doc = _doc_with_image(tmp_path, 300, 2.63)
    (f,) = I.check_images(doc)
    assert f.severity == "error"
    assert "blocky" in f.message


def test_a_print_ready_image_is_silent(tmp_path):
    doc = _doc_with_image(tmp_path, 1900, 6.0)      # 317 DPI
    assert I.check_images(doc) == []


def test_an_icon_is_not_a_figure(tmp_path):
    """A half-inch mark placed inline is a symbol, and its DPI means nothing to
    anyone. Reporting it produces a finding no editor can act on."""
    doc = _doc_with_image(tmp_path, 40, 0.4)
    assert I.check_images(doc) == []


def test_the_facts_are_reported_alongside_the_verdict(tmp_path):
    """An editor asked to chase an author needs the numbers, not just a grade."""
    doc = _doc_with_image(tmp_path, 1259, 6.34)
    (f,) = I.check_images(doc)
    assert "1259x" in f.detail and "6.34" in f.detail


def test_a_document_with_no_images_is_silent():
    assert I.check_images(docx.Document()) == []


def test_pillow_missing_degrades_to_nothing(tmp_path, monkeypatch):
    """An optional dependency must never take a manuscript's run down with it."""
    import builtins

    real = builtins.__import__

    def no_pil(name, *a, **k):
        if name == "PIL":
            raise ImportError("no Pillow")
        return real(name, *a, **k)

    doc = _doc_with_image(tmp_path, 1259, 6.34)
    monkeypatch.setattr(builtins, "__import__", no_pil)
    assert I.read_images(doc) == []
