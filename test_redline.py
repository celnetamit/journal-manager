"""What the redline writer puts in the file.

Written **before** `generate_redline_docx` was touched, so it describes the behaviour
that already worked rather than the behaviour I hoped for. Extending that function is the
riskiest change in this repository: it pairs `doc.paragraphs` with `edited_paragraphs` by
index, and a mistake there does not raise — it writes each track change onto the wrong
paragraph, in a file that opens cleanly and looks like a redline.

The table tests are the new part. Table cells are not in `doc.paragraphs`, which is why
11.3% of a real manuscript has never been copyedited; they are edited through a separate
address (table, row, cell, paragraph) so the body list keeps its exact meaning.
"""

import docx
import pytest
from docx.oxml.ns import qn

import editor as E


def _changes(path):
    """Every insertion and deletion in the file, in document order."""
    doc = docx.Document(path)
    ins, dele = [], []
    for el in doc.element.body.iter():
        if el.tag == qn("w:ins"):
            ins.append("".join(t.text or "" for t in el.iter(qn("w:t"))))
        elif el.tag == qn("w:del"):
            dele.append("".join(t.text or "" for t in el.iter(qn("w:delText"))))
    return ins, dele


def _body_text(path):
    return [p.text for p in docx.Document(path).paragraphs]


def _paragraph_insertions(path, index):
    """Text inserted into one paragraph.

    Not `paragraph.text` — python-docx only walks `w:r` elements that are direct
    children of `w:p`, and an insertion is wrapped in `w:ins`. Reading `.text` shows
    the deletion took effect and the insertion silently missing, which looks exactly
    like a broken redline and is not one.
    """
    p = docx.Document(path).paragraphs[index]
    return "".join(t.text or ""
                   for ins in p._p.iter(qn("w:ins"))
                   for t in ins.iter(qn("w:t")))


def _comment_text(path):
    """Comments live in `word/comments.xml`; the body only carries the anchor."""
    import zipfile
    with zipfile.ZipFile(path) as z:
        if "word/comments.xml" not in z.namelist():
            return ""
        return z.read("word/comments.xml").decode("utf8", "replace")


@pytest.fixture
def source(tmp_path):
    d = docx.Document()
    d.add_paragraph("The result was significant.")
    d.add_paragraph("Scale inhibitors are the primary defence.")
    d.add_paragraph("Table 1. Reagents.")
    t = d.add_table(rows=2, cols=2)
    t.cell(0, 0).text = "Reagent"
    t.cell(0, 1).text = "Purity"
    t.cell(1, 0).text = "Sodium chlorid"          # a typo, inside a table
    t.cell(1, 1).text = "99%"
    path = tmp_path / "src.docx"
    d.save(str(path))
    return str(path)


# ------------------------------------------------------------------- body, unchanged

def test_an_unchanged_paragraph_produces_no_track_change(source, tmp_path):
    out = tmp_path / "a.docx"
    E.generate_redline_docx(source, _body_text(source), str(out))
    ins, dele = _changes(str(out))
    assert ins == [] and dele == []


def test_an_edited_paragraph_is_marked_up(source, tmp_path):
    edits = list(_body_text(source))
    edits[0] = "The result was highly significant."
    out = tmp_path / "b.docx"
    E.generate_redline_docx(source, edits, str(out))

    ins, dele = _changes(str(out))
    assert "highly" in "".join(ins)
    # The rest of the sentence is not deleted and re-inserted: only what changed.
    assert "significant" not in "".join(dele)


def test_edits_land_on_the_paragraph_they_belong_to(source, tmp_path):
    """The failure this whole file exists for. If the pairing ever slips, this is
    what catches it — the second paragraph's edit appearing in the first."""
    edits = list(_body_text(source))
    edits[1] = "Scale inhibitors are the principal defence."
    out = tmp_path / "c.docx"
    E.generate_redline_docx(source, edits, str(out))

    assert "principal" in _paragraph_insertions(str(out), 1)
    assert "principal" not in _paragraph_insertions(str(out), 0)


def test_a_query_becomes_a_comment_on_its_paragraph(source, tmp_path):
    out = tmp_path / "d.docx"
    E.generate_redline_docx(
        source, _body_text(source), str(out),
        queries=[{"index": 1, "query": "Confirm the year.", "suggestion": "2002"}])
    comments = _comment_text(str(out))
    assert "Confirm the year." in comments
    assert "2002" in comments


# --------------------------------------------------------------------------- tables

def test_table_text_can_be_edited(source, tmp_path):
    """The 11.3%. Cell paragraphs are addressed separately from the body list."""
    out = tmp_path / "e.docx"
    E.generate_redline_docx(
        source, _body_text(source), str(out),
        table_edits={(0, 1, 0, 0): "Sodium chloride"})

    ins, dele = _changes(str(out))
    assert "chloride" in "".join(ins)
    assert "chlorid" in "".join(dele)


def test_editing_a_table_leaves_the_body_alone(source, tmp_path):
    out = tmp_path / "f.docx"
    E.generate_redline_docx(
        source, _body_text(source), str(out),
        table_edits={(0, 1, 0, 0): "Sodium chloride"})

    doc = docx.Document(str(out))
    assert doc.paragraphs[0].text == "The result was significant."
    assert doc.paragraphs[1].text == "Scale inhibitors are the primary defence."


def test_an_unchanged_cell_produces_no_track_change(source, tmp_path):
    out = tmp_path / "g.docx"
    E.generate_redline_docx(source, _body_text(source), str(out),
                            table_edits={(0, 1, 0, 0): "Sodium chlorid"})
    ins, dele = _changes(str(out))
    assert ins == [] and dele == []


def test_an_address_that_does_not_exist_is_ignored(source, tmp_path):
    """A stale address must not raise mid-write and lose the whole redline, and must
    not silently write into a neighbouring cell either."""
    out = tmp_path / "h.docx"
    E.generate_redline_docx(source, _body_text(source), str(out),
                            table_edits={(9, 9, 9, 9): "nonsense"})
    ins, dele = _changes(str(out))
    assert ins == [] and dele == []
    assert docx.Document(str(out)).tables[0].cell(1, 0).text == "Sodium chlorid"


def test_table_addresses_round_trip_through_the_reader():
    """`collect_table_texts` must produce addresses `generate_redline_docx` accepts.
    Two halves invented separately is how the edits end up nowhere."""
    import docxmodel as M
    import tempfile
    import os

    d = docx.Document()
    d.add_paragraph("Body.")
    t = d.add_table(rows=1, cols=2)
    # Prose, not labels: `collect_table_texts` deliberately skips short and numeric
    # cells, so a round-trip test built on "alpha"/"beta" would test nothing.
    t.cell(0, 0).text = "The first column describes the sample preparation"
    t.cell(0, 1).text = "The second column reports the measured yield"
    path = os.path.join(tempfile.mkdtemp(), "r.docx")
    d.save(path)

    items = E.collect_table_texts(M.read_structure(path))
    assert [text for _, text in items] == [
        "The first column describes the sample preparation",
        "The second column reports the measured yield",
    ]

    out = os.path.join(os.path.dirname(path), "out.docx")
    E.generate_redline_docx(path, [p.text for p in docx.Document(path).paragraphs],
                            out, table_edits={items[0][0]:
                                "The first column describes how the sample was prepared"})
    # The diff is word-level, so only what changed is inserted: "the sample" was
    # already there and is not re-inserted. Asserting on the whole new sentence
    # would be asserting that the redline is wasteful.
    inserted = "".join(_changes(out)[0])
    assert "how" in inserted and "was prepared" in inserted


# ------------------------------------------------------- which cells are worth editing

@pytest.mark.parametrize("text,editable", [
    ("Table 1 Number of Ranks Given by Sample Respondents", True),
    ("The samples were held at constant temperature", True),
    ("0.15", False),
    ("30", False),
    ("4.5 N/mm2", False),
    ("−0.469 to 5.133", False),
    ("16.67%", False),
    ("N", False),
    ("Mean", False),
    ("Standard deviation", False),      # two words is a label, not a sentence
    ("", False),
    ("   ", False),
])
def test_only_prose_cells_go_to_the_model(text, editable):
    """Measured before this existed: of 1,556 cell paragraphs in one real manuscript,
    71% were pure numbers. Sending them cost 311 LLM calls — and an LLM asked to
    copyedit `0.15` may return `0.150`, which is a silently altered result."""
    assert E.is_editable_cell(text) is editable


def test_a_merged_cell_is_collected_once(tmp_path):
    """python-docx returns the same cell for every grid position it spans, so a
    caption merged across a row arrives once per column. Left alone it is edited,
    billed and written back that many times."""
    import docxmodel as M

    d = docx.Document()
    t = d.add_table(rows=2, cols=4)
    merged = t.cell(0, 0).merge(t.cell(0, 3))
    merged.text = "Table 1 Number of Ranks Given by Sample Respondents"
    for c in range(4):
        t.cell(1, c).text = str(c)

    path = tmp_path / "m.docx"
    d.save(str(path))

    items = E.collect_table_texts(M.read_structure(str(path)))
    captions = [text for _, text in items if text.startswith("Table 1")]
    assert len(captions) == 1, items


# ------------------------------------------------------------------ serper failures

class _Resp:
    def __init__(self, code, payload=None, text=""):
        self.status_code, self._payload, self.text = code, payload, text

    def json(self):
        if self._payload is None:
            raise ValueError("no json")
        return self._payload


def test_an_exhausted_serper_account_is_not_reported_as_a_bad_key(monkeypatch):
    """Serper answers 400 "Not enough credits" when the account runs dry. Reported as
    a bare "HTTP 400" it reads like a malformed request or a dead token, and sends
    somebody to reissue a key that was never the problem — which is exactly what
    happened."""
    monkeypatch.setattr(E.requests, "post",
                        lambda *a, **k: _Resp(400, {"message": "Not enough credits"}))
    ok, msg = E.verify_serper_key("a" * 40)
    assert ok is False
    assert "out of credits" in msg
    assert "valid" in msg


def test_a_genuinely_rejected_key_still_says_so(monkeypatch):
    monkeypatch.setattr(E.requests, "post", lambda *a, **k: _Resp(401))
    ok, msg = E.verify_serper_key("a" * 40)
    assert ok is False and "rejected" in msg


def test_an_unexplained_failure_carries_the_body(monkeypatch):
    monkeypatch.setattr(E.requests, "post",
                        lambda *a, **k: _Resp(503, {"message": "upstream unavailable"}))
    ok, msg = E.verify_serper_key("a" * 40)
    assert ok is False and "upstream unavailable" in msg


# ------------------------------------- the author's own character formatting

def _redline_runs(build_source, edited):
    """(text, superscript, italic, bold) for every run of the first paragraph."""
    import os
    import tempfile

    import docx as _docx
    from docx.oxml.ns import qn as _qn

    import editor as _E

    src = tempfile.mktemp(suffix=".docx")
    build_source().save(src)
    out = tempfile.mktemp(suffix=".docx")
    try:
        _E.generate_redline_docx(src, edited, out)
        para = _docx.Document(out).paragraphs[0]
        rows = []
        for r in para._p.iter(_qn("w:r")):
            text = ("".join(n.text or "" for n in r.iter(_qn("w:t")))
                    or "".join(n.text or "" for n in r.iter(_qn("w:delText"))))
            rpr = r.find(_qn("w:rPr"))
            has = lambda tag: rpr is not None and rpr.find(_qn(tag)) is not None
            rows.append((text, has("w:vertAlign"), has("w:i"), has("w:b")))
        return rows
    finally:
        for f in (src, out):
            if os.path.exists(f):
                os.unlink(f)


def _thermo_source():
    import docx as _docx
    d = _docx.Document()
    p = d.add_paragraph()
    p.add_run("Enthalpy (ΔH")
    p.add_run("#").font.superscript = True     # the activation-parameter marker
    p.add_run(") = 41.49 KJ/mol for ")
    sp = p.add_run("Solanum viarum")
    sp.italic = True
    p.add_run(" samples.")
    return d


def test_an_edited_paragraph_keeps_its_superscripts_and_italics():
    """`_mark_up_paragraph` clears the paragraph and rebuilds it, and the rebuilt runs
    carried no formatting at all. Measured on three real manuscripts: 156, 187 and 96
    edited paragraphs with **zero** italic, bold, superscript or subscript left in any
    of them, while their untouched paragraphs still had 13, 96 and 17.

    A superscript `#` marking an activation parameter and an italic species name are
    exactly what a science manuscript loses, in a redline that otherwise looks
    perfect.
    """
    rows = _redline_runs(
        _thermo_source,
        ["Enthalpy (ΔH#) = 41.49 kJ/mol for Solanum viarum samples."])

    assert any(t == "#" and sup for t, sup, _, _ in rows), rows
    assert any("Solanum viarum" in t and ital for t, _, ital, _ in rows), rows


def test_the_edit_itself_still_happens():
    """Formatting must not be preserved by simply not editing."""
    rows = _redline_runs(
        _thermo_source,
        ["Enthalpy (ΔH#) = 41.49 kJ/mol for Solanum viarum samples."])
    texts = [t for t, _, _, _ in rows]
    assert "KJ" in texts and "kJ" in texts


def test_formatting_does_not_bleed_past_its_run():
    """A token that begins in an ordinary run and ends in a superscript one has to
    become two runs, or the superscript swallows the rest of the word."""
    rows = _redline_runs(
        _thermo_source,
        ["Enthalpy (ΔH#) = 41.49 kJ/mol for Solanum viarum samples."])
    for text, sup, _, _ in rows:
        if sup:
            assert text == "#", f"superscript leaked onto {text!r}"


def test_an_insertion_at_the_very_end_of_a_paragraph():
    """`i1 == len(orig_tokens)` — the commonest edit there is, a full stop or a
    citation appended to the end of a paragraph.

    The formatting-preservation change guarded this with `0 < i1 <= len(starts)` and
    then indexed `starts[i1]` with it, so every such edit raised `list index out of
    range`. Four production jobs failed on it: 34, 35, 36 and 37, all reported as
    "list index out of range" with nothing to act on.
    """
    import docx as _docx

    def source():
        d = _docx.Document()
        p = d.add_paragraph()
        p.add_run("The reaction was complete")
        return d

    rows = _redline_runs(source, ["The reaction was complete [17]."])
    texts = "".join(t for t, _, _, _ in rows)
    assert "[17]." in texts


def test_an_insertion_at_the_very_start_of_a_paragraph():
    """The other end of the same off-by-one."""
    import docx as _docx

    def source():
        d = _docx.Document()
        d.add_paragraph().add_run("reaction was complete.")
        return d

    rows = _redline_runs(source, ["The reaction was complete."])
    assert "The" in "".join(t for t, _, _, _ in rows)


def test_a_paragraph_emptied_by_the_edit_does_not_raise():
    """A degenerate case worth pinning: every token deleted, nothing inserted."""
    import docx as _docx

    def source():
        d = _docx.Document()
        d.add_paragraph().add_run("some text that goes away")
        return d

    _redline_runs(source, [""])


# ------------------------------- the bibliography must survive its own re-sort

def _bib(n, drop=None, dup=None):
    """A tiny numbered bibliography, optionally with an entry lost or repeated."""
    entries = [f"{i}. Author{i} A. A study of thing {i}. J Test. 200{i % 10}; 1: 1."
               for i in range(1, n + 1)]
    if drop is not None:
        entries.pop(drop - 1)
    if dup is not None:
        entries.append(entries[dup - 1])
    return ["Body text before the list."] + entries


def _align(before, after, monkeypatch):
    """Run align_global_citations with the model returning `after`."""
    import json as _json

    import editor as _E

    updates = {str(i): t for i, t in enumerate(after) if t != before[i]}
    monkeypatch.setattr(_E, "_generate_text",
                        lambda *a, **k: _json.dumps(updates))
    warnings = []
    got = _E.align_global_citations(before, {}, "Vancouver", warnings=warnings)
    return got, warnings


def test_a_resort_that_loses_a_reference_is_rejected(monkeypatch):
    """On a real manuscript the model was handed a 75-entry bibliography to re-sort
    and returned it with reference 18 — "Sixth Annual Report on Carcinogens" — gone
    and reference 75 written twice. Nothing checked, so the redline looked perfect: a
    lost reference is invisible unless you count."""
    before = _bib(10)
    after = _bib(10, drop=4, dup=9)             # one lost, one repeated
    got, warnings = _align(before, after, monkeypatch)

    assert got == before, "the original order must be kept intact"
    assert warnings and "left in its original order" in warnings[0]


def test_a_resort_that_preserves_every_reference_is_applied(monkeypatch):
    """Narrowed, not disabled: a correct re-sort still goes through."""
    before = _bib(6)
    after = list(before)
    after[1], after[6] = after[6], after[1]     # a genuine reorder
    got, warnings = _align(before, after, monkeypatch)

    assert got == after
    assert warnings == []


def test_a_manuscript_with_no_bibliography_is_unaffected(monkeypatch):
    before = ["Just body text.", "And more of it."]
    after = ["Just body text [1].", "And more of it."]
    got, warnings = _align(before, after, monkeypatch)
    assert got == after and warnings == []
