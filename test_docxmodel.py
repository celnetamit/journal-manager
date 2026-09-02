"""The structured reader, and the house-spec checks built on it.

The test that matters most is `test_read_docx_contract_is_unchanged`. `generate_redline_docx`
walks `zip(doc.paragraphs, edited_paragraphs)`, so if the plain list ever gains, loses or
reorders an entry, every track change after that point lands on the wrong paragraph — in a
file that still opens, still looks like a redline, and is wrong. Nothing else in the suite
would catch it.

Fixtures are built here rather than committed: a .docx in the repository is a binary
nobody can review, and these need specific formatting rather than a real manuscript.
"""

import docx
import pytest
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

import docxmodel as M
import house_layout as H


@pytest.fixture
def manuscript(tmp_path):
    """A small paper with one correct heading, one wrong one, and a table."""
    d = docx.Document()

    h1 = d.add_paragraph("INTRODUCTION", style="Heading 1")
    h1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for r in h1.runs:
        r.bold = True
        r.font.name = "Times New Roman"
        r.font.size = Pt(11)

    d.add_paragraph("Scale inhibition matters. See Table 1.")

    bad = d.add_paragraph("Materials and Methods", style="Heading 1")   # wrong case
    for r in bad.runs:
        r.font.name = "Calibri"                                          # wrong font
        r.font.size = Pt(14)                                             # wrong size

    d.add_paragraph("Table 1. Reagents used.")
    t = d.add_table(rows=2, cols=2)
    t.cell(0, 0).text = "Reagent"
    t.cell(0, 1).text = "Purity"
    t.cell(1, 0).text = "NaCl"
    t.cell(1, 1).text = "99%"

    path = tmp_path / "m.docx"
    d.save(str(path))
    return str(path)


def test_read_docx_contract_is_unchanged(manuscript):
    """`Structure.texts` must equal `[p.text for p in doc.paragraphs]`, exactly.

    Same length, same order, same strings. The redline writer pairs the two by index
    and has no way to notice a mismatch.
    """
    plain = [p.text for p in docx.Document(manuscript).paragraphs]
    s = M.read_structure(manuscript)

    assert s.texts == plain
    assert len(s.paragraphs) == len(plain)
    for i, text in enumerate(plain):
        assert s.paragraphs[i].index == i
        assert s.paragraphs[i].text == text


def test_table_text_is_read(manuscript):
    """The old reader could not see any of this; 11.3% of a real manuscript sat here."""
    s = M.read_structure(manuscript)
    assert len(s.tables) == 1
    assert s.tables[0].rows == 2 and s.tables[0].cols == 2
    assert "NaCl" in s.tables[0].text
    # …and it is still absent from the plain list, which is what keeps the redline safe.
    assert not any("NaCl" in t for t in s.texts)


def test_a_table_knows_which_paragraph_it_follows(manuscript):
    s = M.read_structure(manuscript)
    above = s.paragraphs[s.tables[0].after_paragraph]
    assert above.text.startswith("Table 1.")


def test_formatting_survives(manuscript):
    s = M.read_structure(manuscript)
    intro = next(p for p in s.paragraphs if p.text == "INTRODUCTION")
    assert intro.is_bold is True
    assert intro.dominant_font == "Times New Roman"
    assert intro.dominant_size_pt == 11.0
    assert intro.outline_level == 0


def test_font_is_resolved_through_the_style(tmp_path):
    """A correctly-formatted heading has no font on the run — it is on the style.

    Reading only the run returns None for every paragraph, which reads as uniformity
    and is a dead measurement: nothing about "Times New Roman, 11 pt" is checkable.
    """
    d = docx.Document()
    style = d.styles["Heading 1"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)
    p = d.add_paragraph("METHODS", style="Heading 1")
    assert all(r.font.name is None for r in p.runs)      # nothing on the run itself

    path = tmp_path / "s.docx"
    d.save(str(path))

    para = M.read_structure(str(path)).paragraphs[0]
    assert para.dominant_font == "Times New Roman"
    assert para.dominant_size_pt == 11.0


@pytest.mark.parametrize("text,expected", [
    ("INTRODUCTION", "upper"),
    ("Materials and Methods", "title"),
    ("Results of the compatibility test", "sentence"),
    ("Findings and Analysis", "title"),
    ("mixed Case heading Here", "mixed"),
    ("2025", "unknown"),          # no letters: not silently "upper"
    ("", "unknown"),
])
def test_case_detection(text, expected):
    assert H.case_of(text) == expected


def test_house_checks_find_the_planted_faults(manuscript):
    findings = H.check_all(M.read_structure(manuscript))
    rules = {f.rule for f in findings}
    assert "heading.case" in rules       # "Materials and Methods" is not upper case
    assert "heading.font" in rules       # Calibri
    assert "heading.size" in rules       # 14 pt


def test_a_correct_heading_raises_nothing(manuscript):
    findings = H.check_headings(M.read_structure(manuscript))
    intro = [f for f in findings if f.paragraph == 0]
    assert intro == [], intro


def test_body_text_wearing_a_heading_style_is_caught(tmp_path):
    d = docx.Document()
    d.add_paragraph("A" * 40 + " and then a great deal more text besides, "
                    "which is a paragraph and not a heading at all.", style="Heading 1")
    path = tmp_path / "b.docx"
    d.save(str(path))

    findings = H.check_headings(M.read_structure(str(path)))
    assert [f.rule for f in findings] == ["heading.body-text-as-heading"]


def test_skipped_heading_level_is_reported(tmp_path):
    d = docx.Document()
    d.add_paragraph("INTRODUCTION", style="Heading 1")
    d.add_paragraph("Sub Sub Heading", style="Heading 3")
    path = tmp_path / "h.docx"
    d.save(str(path))

    findings = H.check_headings(M.read_structure(str(path)))
    assert any(f.rule == "heading.skipped-level" for f in findings)


def test_an_uncaptioned_table_is_reported(tmp_path):
    d = docx.Document()
    d.add_paragraph("Some body text with no caption.")
    t = d.add_table(rows=1, cols=1)
    t.cell(0, 0).text = "x"
    path = tmp_path / "t.docx"
    d.save(str(path))

    findings = H.check_tables(M.read_structure(str(path)))
    assert any(f.rule == "table.caption" for f in findings)


def test_oversized_artwork_is_reported(monkeypatch):
    s = M.Structure(paragraphs=[], tables=[], sections=[],
                    images=[M.Image(index=0, width_in=10.0, height_in=7.0)])
    findings = H.check_artwork(s)
    assert len(findings) == 1 and findings[0].rule == "artwork.size"


def test_artwork_at_the_limit_is_fine():
    s = M.Structure(paragraphs=[], tables=[], sections=[],
                    images=[M.Image(index=0, width_in=9.0, height_in=6.0)])
    assert H.check_artwork(s) == []


def test_a4_is_not_reported_as_the_wrong_page_size():
    """Word stores A4 as 8.268 x 11.693. A checker that calls that a deviation from
    8.27 x 11.69 reports every correct manuscript as broken, and stops being read."""
    s = M.Structure(paragraphs=[], tables=[], images=[], sections=[
        M.Section(page_width_in=8.268, page_height_in=11.693,
                  left_margin_in=1.0, right_margin_in=1.0,
                  top_margin_in=0.6, bottom_margin_in=0.5)])
    assert H.check_page(s) == []


def test_wrong_margins_are_reported():
    s = M.Structure(paragraphs=[], tables=[], images=[], sections=[
        M.Section(page_width_in=8.268, page_height_in=11.693,
                  left_margin_in=0.59, right_margin_in=0.59,
                  top_margin_in=0.319, bottom_margin_in=0.194)])
    findings = H.check_page(s)
    assert len(findings) == 4
    assert all(f.rule == "page.geometry" for f in findings)


# --------------------------------------------------------------- table formatting

@pytest.fixture
def table_doc(tmp_path):
    """A table set the way the house asks, so a deviation is the exception."""
    d = docx.Document()

    cap = d.add_paragraph()
    label = cap.add_run("Table 1. ")
    label.bold = True
    label.font.name = "Times New Roman"
    label.font.size = Pt(9)
    rest = cap.add_run("Frequency and percentage distribution.")
    rest.font.name = "Times New Roman"
    rest.font.size = Pt(11)

    t = d.add_table(rows=2, cols=2)
    for (r, c), text, bold in (((0, 0), "S.N.", True), ((0, 1), "Description", True),
                               ((1, 0), "1.", False), ((1, 1), "Age", False)):
        para = t.cell(r, c).paragraphs[0]
        run = para.add_run(text)
        run.bold = bold
        run.font.name = "Times New Roman"
        run.font.size = Pt(9)

    path = tmp_path / "t.docx"
    d.save(str(path))
    return str(path)


def test_table_cells_keep_their_formatting(table_doc):
    t = M.read_structure(table_doc).tables[0]
    head = t.grid[0][0].paragraphs[0]
    assert head.text == "S.N."
    assert head.is_bold is True
    assert head.dominant_size_pt == 9.0
    assert head.in_table is True
    # Cell paragraphs are not in `doc.paragraphs`, so they carry no index there.
    assert head.index == -1


def test_a_correctly_set_table_is_quiet_about_its_text(table_doc):
    findings = H.check_table_format(M.read_structure(table_doc))
    rules = {f.rule for f in findings}
    assert "table.font" not in rules
    assert "table.size" not in rules
    assert "table.column-head" not in rules


def test_wrong_table_text_size_is_reported(tmp_path):
    d = docx.Document()
    d.add_paragraph("Table 1. Something.")
    t = d.add_table(rows=1, cols=1)
    run = t.cell(0, 0).paragraphs[0].add_run("value")
    run.font.name = "Times New Roman"
    run.font.size = Pt(8.5)          # the Biophilic manuscript's real defect
    path = tmp_path / "s.docx"
    d.save(str(path))

    findings = H.check_table_format(M.read_structure(str(path)))
    assert any(f.rule == "table.size" and "8.5 pt" in f.message for f in findings)


def test_a_caption_that_is_bold_throughout_is_reported(tmp_path):
    """The house sets 'Table N.' bold and the caption text normal. A paragraph that
    is uniformly one or the other cannot be right, and no per-paragraph bold check
    can see it — it has to look run by run."""
    d = docx.Document()
    p = d.add_paragraph()
    r = p.add_run("Table 1. Everything here is bold.")
    r.bold = True
    r.font.size = Pt(11)
    t = d.add_table(rows=1, cols=1)
    t.cell(0, 0).text = "x"
    path = tmp_path / "c.docx"
    d.save(str(path))

    findings = H.check_table_format(M.read_structure(str(path)))
    assert any(f.rule == "table.caption-weight" for f in findings)


def test_border_and_margins_are_read_in_the_units_the_spec_uses(tmp_path):
    """Word stores borders in eighths of a point and margins in twentieths.
    Comparing a spec of "1/2 pt" against a stored 4 is confidently wrong."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    d = docx.Document()
    t = d.add_table(rows=1, cols=1)
    t.cell(0, 0).text = "x"
    tblPr = t._tbl.find(qn("w:tblPr"))

    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right"):
        e = OxmlElement("w:" + edge)
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), "4")            # 4 eighths = 0.5 pt
        borders.append(e)
    tblPr.append(borders)

    mar = OxmlElement("w:tblCellMar")
    for edge, dxa in (("top", "29"), ("bottom", "29"), ("left", "58"), ("right", "58")):
        e = OxmlElement("w:" + edge)
        e.set(qn("w:w"), dxa)             # 58/1440 = 0.04"
        e.set(qn("w:type"), "dxa")
        mar.append(e)
    tblPr.append(mar)

    path = tmp_path / "b.docx"
    d.save(str(path))

    fmt = M.read_structure(str(path)).tables[0].fmt
    assert fmt.border_style == "single"
    assert fmt.border_size_pt == 0.5
    assert fmt.margin_left_in == 0.04
    assert fmt.margin_top_in == 0.02

    # …and a table set exactly to spec says nothing about its borders or margins.
    findings = H.check_table_format(M.read_structure(str(path)))
    assert not [f for f in findings if f.rule in ("table.border", "table.cell-margin")]


def test_the_same_deviation_across_many_tables_is_collapsed():
    """Twelve tables with one wrong margin produced 34 findings on a real paper —
    enough to bury the four that were about something else."""
    findings = [
        H.Finding("table.cell-margin", "info", None,
                  f"table {i} left cell margin is 0.03\", house is 0.04\"")
        for i in range(1, 13)
    ] + [H.Finding("heading.case", "error", 3, "H1 should be upper case")]

    collapsed = H.collapse_repeats(findings)
    assert len(collapsed) == 2
    margin = next(f for f in collapsed if f.rule == "table.cell-margin")
    assert "12 tables" in margin.message
    assert "table 1" in margin.detail          # somewhere to go and look


def test_the_same_deviation_across_many_headings_is_collapsed():
    """Measured on 1,597 real manuscripts: `heading.size` averages 11 findings a
    paper and reaches 97 in one, every line word for word the one above it. The
    heading rules were the largest single source of noise in the report."""
    findings = [
        H.Finding("heading.font", "error", i,
                  "H3 is Cambria, house font is Times New Roman")
        for i in range(64)
    ]
    collapsed = H.collapse_repeats(findings)

    assert len(collapsed) == 1
    # The count has to be stated. The heading message carries no number of its own, so
    # without it the collapsed finding reads exactly like a single one and the other
    # sixty-three are simply gone.
    assert "64 headings" in collapsed[0].message
    assert "Cambria" in collapsed[0].message
    assert collapsed[0].paragraph == 0         # anchored at the first one


def test_two_heading_levels_are_not_folded_together():
    """`_COLLAPSIBLE`'s key normalises "table N" and "section N". It must not do the
    same to the heading level: folding H1 into H3 would name a level the editor then
    finds nothing wrong with."""
    findings = ([H.Finding("heading.size", "error", i, "H1 is 12.0 pt, house size is 11.0 pt")
                 for i in range(5)]
                + [H.Finding("heading.size", "error", 10 + i, "H3 is 12.0 pt, house size is 11.0 pt")
                   for i in range(3)])

    collapsed = H.collapse_repeats(findings)
    assert len(collapsed) == 2
    assert {"5 headings: H1 is 12.0 pt, house size is 11.0 pt",
            "3 headings: H3 is 12.0 pt, house size is 11.0 pt"} == {f.message for f in collapsed}


def test_over_long_headings_fold_despite_differing_lengths():
    """`heading.body-text-as-heading` is the one heading message carrying a varying
    number — "(143 characters)". Unnormalised, every one is its own group and the rule
    that reaches 118 findings in a manuscript would not fold at all."""
    findings = [
        H.Finding("heading.body-text-as-heading", "error", i,
                  f"styled 'Heading 2' but reads as body text ({100 + i * 7} characters)")
        for i in range(9)
    ]
    collapsed = H.collapse_repeats(findings)
    assert len(collapsed) == 1
    assert collapsed[0].message.startswith("9 headings:")
    assert "(100 characters)" in collapsed[0].message      # the first one, exactly


def test_a_heading_level_with_one_deviation_is_left_alone():
    """One wrong heading should read as one wrong heading, not "1 headings:"."""
    findings = [H.Finding("heading.case", "error", 3,
                          "H1 should be upper case, reads as title")]
    assert H.collapse_repeats(findings) == findings


def test_skipped_level_is_never_collapsed():
    """It names one specific place in the hierarchy. A count would say nothing."""
    findings = [H.Finding("heading.skipped-level", "warning", i,
                          "H2 is followed by H4 — H3 is missing") for i in (4, 19)]
    assert len(H.collapse_repeats(findings)) == 2


# ------------------------------------------------------- body text & front matter

def _front(tmp_path, name="f.docx"):
    d = docx.Document()
    t = d.add_paragraph()
    r = t.add_run("Safety Analysis of a Three Span Reinforced Concrete Beam")
    r.font.name = "Calibri Light"
    r.font.size = Pt(20)
    r.bold = True
    t.alignment = WD_ALIGN_PARAGRAPH.LEFT

    a = d.add_paragraph()
    ar = a.add_run("Sule S., T.C. Nwofor, Matthew F.")
    ar.font.name = "Garamond"
    ar.font.size = Pt(12)
    a.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return d


def test_a_correct_title_and_author_line_say_nothing(tmp_path):
    d = _front(tmp_path)
    d.add_paragraph("Abstract")
    d.add_paragraph("Keywords: one, two")
    path = tmp_path / "ok.docx"
    d.save(str(path))

    findings = H.check_front_matter(M.read_structure(str(path)))
    assert not [f for f in findings if f.rule.startswith(("title.", "authors."))], findings


def test_a_title_in_the_wrong_font_is_reported(tmp_path):
    d = docx.Document()
    p = d.add_paragraph()
    r = p.add_run("Safety Analysis of a Three Span Reinforced Concrete Beam")
    r.font.name = "Times New Roman"          # house is Calibri Light
    r.font.size = Pt(20)
    d.add_paragraph("Sule S.")
    path = tmp_path / "t.docx"
    d.save(str(path))

    findings = H.check_front_matter(M.read_structure(str(path)))
    assert any(f.rule == "title.font" for f in findings)


def test_a_run_on_abstract_is_not_reported_as_missing(tmp_path):
    """`Abstract— This study examines…` is an abstract that is formatted wrong, not a
    missing one. Saying "no abstract was found" sends an editor looking for something
    that is on the page."""
    d = _front(tmp_path)
    d.add_paragraph("Abstract— This study examines gender inequality and its influence "
                    "on economic empowerment across three districts of the region.")
    d.add_paragraph("Keywords: gender, economics")
    path = tmp_path / "r.docx"
    d.save(str(path))

    rules = {f.rule for f in H.check_front_matter(M.read_structure(str(path)))}
    assert "front.abstract-runon" in rules
    assert "front.abstract-missing" not in rules


def test_a_genuinely_missing_abstract_is_still_reported(tmp_path):
    d = _front(tmp_path)
    d.add_paragraph("Keywords: gender, economics")
    path = tmp_path / "n.docx"
    d.save(str(path))

    rules = {f.rule for f in H.check_front_matter(M.read_structure(str(path)))}
    assert "front.abstract-missing" in rules


def test_the_keywords_finding_names_the_separator_that_is_there(tmp_path):
    d = _front(tmp_path)
    d.add_paragraph("Abstract")
    d.add_paragraph("Keywords – Biophilic Design, Vastu Shastra, Interior Design")
    path = tmp_path / "k.docx"
    d.save(str(path))

    findings = H.check_front_matter(M.read_structure(str(path)))
    kw = next(f for f in findings if f.rule == "front.keywords-colon")
    # Saying only "should be a colon" is a complaint; naming what is there is an
    # instruction.
    assert "–" in kw.message


def test_body_text_is_reported_as_a_share_not_per_paragraph(tmp_path):
    """A manuscript on the wrong template has every paragraph wrong, and three
    hundred identical findings is one piece of information told badly."""
    d = docx.Document()
    for _ in range(10):
        p = d.add_paragraph()
        r = p.add_run("Reinforced concrete beam is one of the most widely used "
                      "structural members and beams are used in buildings and "
                      "bridges throughout the region under study here. " * 2)
        r.font.name = "Arial"                # house is Times New Roman
        r.font.size = Pt(11)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    path = tmp_path / "b.docx"
    d.save(str(path))

    findings = H.check_body_text(M.read_structure(str(path)))
    font = [f for f in findings if f.rule == "body.font"]
    assert len(font) == 1
    assert "10 of 10" in font[0].message


def test_one_stray_paragraph_does_not_trip_the_body_check(tmp_path):
    d = docx.Document()
    for i in range(10):
        p = d.add_paragraph()
        r = p.add_run("Reinforced concrete beam is one of the most widely used "
                      "structural members and beams are used in buildings and "
                      "bridges throughout the region under study here. " * 2)
        r.font.name = "Arial" if i == 0 else "Times New Roman"
        r.font.size = Pt(11)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    path = tmp_path / "one.docx"
    d.save(str(path))

    findings = H.check_body_text(M.read_structure(str(path)))
    assert not [f for f in findings if f.rule == "body.font"]


def test_unnumbered_references_are_reported(tmp_path):
    d = docx.Document()
    d.add_paragraph("REFERENCES", style="Heading 1")
    for name in ("Afolayan JO, Abubakar I. Reliability-based design program for slabs.",
                 "Amartey YD. Reliability evaluation of in-situ strength of members.",
                 "Biondini F. Probabilistic limit states analysis of framed structures."):
        d.add_paragraph(name)
    path = tmp_path / "ref.docx"
    d.save(str(path))

    findings = H.check_references(M.read_structure(str(path)))
    assert any(f.rule == "references.numbering" for f in findings)


def test_numbered_references_are_accepted(tmp_path):
    d = docx.Document()
    d.add_paragraph("REFERENCES", style="Heading 1")
    for i, name in enumerate((
            "Afolayan JO, Abubakar I. Reliability-based design program for slabs.",
            "Amartey YD. Reliability evaluation of in-situ strength of members.",
            "Biondini F. Probabilistic limit states analysis of framed structures."), 1):
        d.add_paragraph(f"{i}. {name}")
    path = tmp_path / "ref2.docx"
    d.save(str(path))

    findings = H.check_references(M.read_structure(str(path)))
    assert not [f for f in findings if f.rule == "references.numbering"]


# ------------------------------------------------- glyphs and labels the corpus uses

def _bullet_para(index, glyph, num_id=1):
    """A list paragraph as `read_structure` would return it."""
    return M.Para(index=index, text="a list item", style="List Paragraph",
                  listing={"num_id": num_id, "level": 0, "kind": "bullet",
                           "house_code": M.BULLET_GLYPHS.get(glyph), "glyph": glyph})


def test_the_solid_round_bullet_is_a_house_mark():
    """The checker printed "bullet '●' is not one of the house marks (B1 ● …)" — it
    rejected the very glyph its own message shows as B1. `●` is U+25CF; only U+F0B7 and
    U+2022 were mapped. 196 findings over 400 manuscripts, and nothing was wrong."""
    assert M.BULLET_GLYPHS["●"] == "B1"
    st = M.Structure(paragraphs=[_bullet_para(0, "●")], tables=[], sections=[], images=[])
    assert not [f for f in H.check_listings(st) if f.rule == "listing.bullet"]


def test_a_bullet_that_is_not_a_house_mark_is_still_reported():
    """Wingdings' tick. Narrowing must not empty the rule."""
    st = M.Structure(paragraphs=[_bullet_para(0, "")], tables=[], sections=[], images=[])
    assert [f for f in H.check_listings(st) if f.rule == "listing.bullet"]


def _one_table(caption):
    cap = M.Para(index=0, text=caption, style="Normal")
    tbl = M.Table(index=0, rows=1, cols=1, cells=[["x"]], after_paragraph=0)
    return M.Structure(paragraphs=[cap], tables=[tbl], sections=[], images=[])


def test_a_roman_numbered_caption_is_a_caption():
    """`Table I – Summary of Cybersecurity Threats` was reported as "table 1 has no
    'Table N' caption immediately above it", with the caption sitting right there. The
    numbering style is a different question from whether a caption exists."""
    assert not [f for f in H.check_tables(_one_table("Table I – Summary of Threats"))
                if f.rule == "table.caption"]


def test_a_table_with_no_caption_is_still_reported():
    for text in ("N=150", "This case study is located in Cairo."):
        assert [f for f in H.check_tables(_one_table(text)) if f.rule == "table.caption"]


def test_a_sentence_starting_with_table_is_not_a_caption():
    """`[ivxlc]+` must not swallow ordinary words — "Table Illustrating…" begins with
    an I and is prose."""
    assert [f for f in H.check_tables(_one_table("Table Illustrating the overall trend"))
            if f.rule == "table.caption"]
