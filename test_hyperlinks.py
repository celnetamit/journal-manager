"""Hyperlinks in the redline.

The tests that matter here are not the ones about which strings are recognised — they
are the ones asserting that linking changed *nothing else*. The redline aligns tracked
changes by paragraph index, so a paragraph gained or lost silently moves every later
change onto the wrong text, and the file still opens and still looks right. Two of
these tests exist purely to make that impossible to ship.

`test_never_nests_hyperlink_inside_ins` is the other one worth keeping. The first
working version produced `<w:ins><w:hyperlink>`, which is invalid OOXML — `w:ins` has
no `hyperlink` in its content group — and Word's response to invalid markup is to
repair the file, dropping the links without an error. Every DOI in a redline is inside
`w:ins`, because the copyedit reformats the references, so that bug affected 26 links
out of 30 and nothing in the output looked wrong.
"""

import docx
import pytest
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

import hyperlinks as HL


def _doc(*paragraph_texts):
    d = docx.Document()
    for t in paragraph_texts:
        d.add_paragraph(t)
    return d


def _saved(doc, tmp_path, name="out.docx"):
    path = tmp_path / name
    doc.save(str(path))
    return docx.Document(str(path))


def _links(doc):
    return [h for p in doc.paragraphs for h in p._p.iter(qn("w:hyperlink"))]


def _link_text(h):
    return "".join(t.text or "" for t in h.iter(qn("w:t")))


def _insert_tracked(paragraph, text, author="AI Editor"):
    """A `w:ins` run, the shape `generate_redline_docx` produces."""
    ins = OxmlElement("w:ins")
    ins.set(qn("w:id"), "1")
    ins.set(qn("w:author"), author)
    ins.set(qn("w:date"), "2026-09-04T00:00:00Z")
    run = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    run.append(t)
    ins.append(run)
    paragraph._p.append(ins)
    return ins


def _delete_tracked(paragraph, text):
    dele = OxmlElement("w:del")
    dele.set(qn("w:id"), "2")
    dele.set(qn("w:author"), "AI Editor")
    dele.set(qn("w:date"), "2026-09-04T00:00:00Z")
    run = OxmlElement("w:r")
    t = OxmlElement("w:delText")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    run.append(t)
    dele.append(run)
    paragraph._p.append(dele)
    return dele


# --- what counts as an address -------------------------------------------------

@pytest.mark.parametrize("text, expected_href", [
    ("See https://doi.org/10.1234/abc for detail.", "https://doi.org/10.1234/abc"),
    ("doi:10.1016/j.jcis.2019.05.001", "https://doi.org/10.1016/j.jcis.2019.05.001"),
    ("10.3390/ma17235997", "https://doi.org/10.3390/ma17235997"),
    ("Write to jaya@gbu.ac.in today", "mailto:jaya@gbu.ac.in"),
    ("Visit www.nanoschool.in now", "https://www.nanoschool.in"),
])
def test_recognises_addresses(text, expected_href):
    found = HL.find_targets(text)
    assert [h for _, _, h in found] == [expected_href]


def test_trailing_sentence_punctuation_is_not_part_of_the_address():
    """A DOI ending a reference is followed by a full stop, and the stop is not its."""
    (start, end, href), = HL.find_targets("... Chem. 12:44. doi:10.1234/xy.2020.")
    assert href == "https://doi.org/10.1234/xy.2020"
    assert not href.endswith(".")


def test_a_doi_url_is_matched_once_not_twice():
    """`https://doi.org/10.x/y` contains a bare DOI. Wrapping both would nest links."""
    assert len(HL.find_targets("https://doi.org/10.1234/abc")) == 1


def test_prose_with_no_address_yields_nothing():
    assert HL.find_targets("The samples were calcined at 550°C for 2 h.") == []


# --- the invariants ------------------------------------------------------------

def test_paragraph_count_is_unchanged(tmp_path):
    """The whole redline is aligned by paragraph index. This may never move."""
    doc = _doc("Intro.", "See doi:10.1234/a.", "Mail x@y.org.", "End.")
    before = len(doc.paragraphs)
    HL.linkify_document(doc)
    assert len(doc.paragraphs) == before
    assert len(_saved(doc, tmp_path).paragraphs) == before


def test_text_is_never_rewritten(tmp_path):
    doc = _doc("Contact jaya@gbu.ac.in for the dataset and code.")
    before = doc.paragraphs[0].text
    HL.linkify_document(doc)
    assert _saved(doc, tmp_path).paragraphs[0].text == before


def test_tracked_changes_survive_intact(tmp_path):
    """Inserted and deleted text must read back byte-identical after linking."""
    doc = _doc("Reference list:")
    _insert_tracked(doc.paragraphs[0], " See https://doi.org/10.1234/new.")
    _delete_tracked(doc.paragraphs[0], " old text ")
    ins_before = "".join(t.text or "" for t in doc.paragraphs[0]._p.iter(qn("w:t")))
    del_before = "".join(t.text or "" for t in doc.paragraphs[0]._p.iter(qn("w:delText")))

    HL.linkify_document(doc)
    out = _saved(doc, tmp_path).paragraphs[0]
    assert "".join(t.text or "" for t in out._p.iter(qn("w:t"))) == ins_before
    assert "".join(t.text or "" for t in out._p.iter(qn("w:delText"))) == del_before


def test_links_inserted_text(tmp_path):
    """Nearly every DOI in a real redline is inserted text, because the copyedit
    reformats the references. A version that skipped `w:ins` would look like it
    worked and link almost nothing."""
    doc = _doc("Refs:")
    _insert_tracked(doc.paragraphs[0], " https://doi.org/10.1234/inserted")
    HL.linkify_document(doc)
    links = _links(_saved(doc, tmp_path))
    assert len(links) == 1
    assert "10.1234/inserted" in _link_text(links[0])


def test_never_nests_hyperlink_inside_ins(tmp_path):
    """`<w:ins><w:hyperlink>` is invalid OOXML; Word repairs it by dropping links."""
    doc = _doc("Refs:")
    _insert_tracked(doc.paragraphs[0], " https://doi.org/10.1234/x")
    HL.linkify_document(doc)
    out = _saved(doc, tmp_path)
    assert not [h for p in out.paragraphs
                for ins in p._p.iter(qn("w:ins")) for h in ins.iter(qn("w:hyperlink"))]
    # and the correct nesting is present
    assert [i for h in _links(out) for i in h.iter(qn("w:ins"))]


def test_the_insertion_keeps_its_author(tmp_path):
    """Splitting a `w:ins` around the link must carry the attributes over, or Word
    shows the change as authored by nobody."""
    doc = _doc("Refs:")
    _insert_tracked(doc.paragraphs[0], " https://doi.org/10.1234/x rest", author="AI Editor")
    HL.linkify_document(doc)
    out = _saved(doc, tmp_path)
    every = [i for p in out.paragraphs for i in p._p.iter(qn("w:ins"))]
    assert every and all(i.get(qn("w:author")) == "AI Editor" for i in every)


def test_deleted_text_never_becomes_clickable(tmp_path):
    """A DOI the redline proposes to remove must not invite the editor to follow it."""
    doc = _doc("Refs:")
    _delete_tracked(doc.paragraphs[0], " https://doi.org/10.1234/removed")
    HL.linkify_document(doc)
    out = _saved(doc, tmp_path)
    assert not [h for p in out.paragraphs
                for d in p._p.iter(qn("w:del")) for h in d.iter(qn("w:hyperlink"))]
    assert _links(out) == []


def test_only_the_address_is_linked_not_the_sentence(tmp_path):
    doc = _doc("Please contact jaya@gbu.ac.in for the raw data.")
    HL.linkify_document(doc)
    link, = _links(_saved(doc, tmp_path))
    assert _link_text(link) == "jaya@gbu.ac.in"


def test_repeated_url_reuses_one_relationship(tmp_path):
    doc = _doc("First https://doi.org/10.1/aaaa here.",
               "Again https://doi.org/10.1/aaaa there.")
    HL.linkify_document(doc)
    out = _saved(doc, tmp_path)
    assert len(_links(out)) == 2
    rids = {h.get(qn("r:id")) for h in _links(out)}
    assert len(rids) == 1


def test_run_formatting_is_preserved(tmp_path):
    doc = docx.Document()
    p = doc.add_paragraph()
    run = p.add_run("Bold https://doi.org/10.1/bbbb tail")
    run.bold = True
    HL.linkify_document(doc)
    out = _saved(doc, tmp_path).paragraphs[0]
    assert all(r.bold for r in out.runs), "splitting a run must copy its properties"


def test_running_twice_does_not_double_wrap(tmp_path):
    doc = _doc("See https://doi.org/10.1/cccc now.")
    HL.linkify_document(doc)
    HL.linkify_document(doc)
    out = _saved(doc, tmp_path)
    assert len(_links(out)) == 1
    assert not [h for outer in _links(out) for h in outer.iter(qn("w:hyperlink"))
                if h is not outer]


def test_document_with_no_addresses_is_byte_stable(tmp_path):
    doc = _doc("No addresses.", "None here either.")
    assert HL.linkify_document(doc) == 0


def test_table_cells_are_linked_too(tmp_path):
    doc = docx.Document()
    table = doc.add_table(rows=1, cols=1)
    table.rows[0].cells[0].paragraphs[0].text = "Mail a@b.org"
    assert HL.linkify_document(doc) == 1


def test_address_split_across_runs_is_linked_whole(tmp_path):
    """Word splits runs wherever it likes. In a real manuscript
    `doi:10.1016/j.ajhg.2016.09.015` arrived as `doi:10.1016/j.ajhg` +
    `.2016.09.015.`, and per-run detection linked the first half to a DOI that does
    not exist — a link that looks right in Word and resolves to nothing."""
    doc = docx.Document()
    p = doc.add_paragraph()
    p.add_run("Am J Hum Genet. 2016. ")
    p.add_run("doi:10.1016/j.ajhg")
    p.add_run(".2016.09.015.")
    HL.linkify_document(doc)
    out = _saved(doc, tmp_path)
    link, = _links(out)
    assert _link_text(link) == "doi:10.1016/j.ajhg.2016.09.015"
    assert out.paragraphs[0].text == "Am J Hum Genet. 2016. doi:10.1016/j.ajhg.2016.09.015."


def test_address_split_across_runs_inside_an_insertion(tmp_path):
    doc = _doc("Refs:")
    ins = _insert_tracked(doc.paragraphs[0], " https://doi.org/10.1016/")
    run = OxmlElement("w:r")
    t = OxmlElement("w:t"); t.text = "j.jcis.2019.05.001"; t.set(qn("xml:space"), "preserve")
    run.append(t); ins.append(run)
    HL.linkify_document(doc)
    out = _saved(doc, tmp_path)
    link, = _links(out)
    assert _link_text(link) == "https://doi.org/10.1016/j.jcis.2019.05.001"
    assert not [h for p in out.paragraphs
                for i in p._p.iter(qn("w:ins")) for h in i.iter(qn("w:hyperlink"))]


def test_an_address_broken_by_a_deletion_is_not_stitched_together(tmp_path):
    """Text the redline is deleting must not be spliced into a link."""
    doc = _doc("Refs:")
    p = doc.paragraphs[0]
    _insert_tracked(p, " doi:10.1016/j.ajhg")
    _delete_tracked(p, "REMOVED")
    _insert_tracked(p, ".2016.09.015")
    HL.linkify_document(doc)
    out = _saved(doc, tmp_path)
    for h in _links(out):
        assert "REMOVED" not in _link_text(h)
