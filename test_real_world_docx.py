"""What real manuscripts do that clean test fixtures never do.

Every test here comes from a measured failure on a 400-manuscript sweep of the WordPress
media corpus (`wpimport/media/*.docx`) — law, materials science, agronomy, nursing,
homoeopathy. Before these fixes, 39 of those 400 (9.8%) could not be read at all; after,
3 (0.8%), and all three are files whose embedded images have a corrupt CRC.

The failure mode was never a crash the user saw. `run_pipeline` already catches a
structure-read failure and carries on with the copyedit, so the job finished `done` and
the House Style panel — the flagship feature — was simply absent, under a warning that
read "House-style layout check was skipped: " with nothing after the colon. That is the
shape worth testing: not "does it raise", but "does the check still happen".
"""

import json
import zipfile

import docx
from docx.oxml.ns import qn
import pytest

import docxmodel as M

#: The minimum `run_pipeline` insists on. Building it here rather than in each test
#: keeps the tests about the failure they describe.
_OPTS = {"user_id": 1, "edit_style": "CMOS", "ref_style": "Vancouver",
         "lang_type": "US English"}


@pytest.fixture
def plain(tmp_path):
    """A document with no lists — so Word writes no numbering part at all."""
    d = docx.Document()
    d.add_paragraph("Introduction")
    d.add_paragraph("The samples were held at constant mass throughout the trial.")
    p = tmp_path / "plain.docx"
    d.save(str(p))
    return str(p)


# --------------------------------------------------------------- no numbering part

def test_a_document_with_no_list_is_still_read(plain, monkeypatch):
    """32 of 400. python-docx raises a bare `NotImplementedError` when a document has
    never contained a list, and the original guard named KeyError/AttributeError/
    ValueError — every exception except the one that actually fires."""
    import docx.parts.document as dp

    def boom(self):
        raise NotImplementedError

    monkeypatch.setattr(dp.DocumentPart, "numbering_part",
                        property(boom), raising=False)

    st = M.read_structure(plain)
    assert len(st.paragraphs) == 2
    assert st.paragraphs[0].text == "Introduction"


def test_no_numbering_means_no_list_marker_not_a_failure(plain, monkeypatch):
    import docx.parts.document as dp
    monkeypatch.setattr(dp.DocumentPart, "numbering_part",
                        property(lambda self: (_ for _ in ()).throw(NotImplementedError)),
                        raising=False)
    st = M.read_structure(plain)
    assert all(not p.listing for p in st.paragraphs)


# ------------------------------------------------------------ fractional font sizes

def test_a_fractional_half_point_size_is_read_not_raised(plain):
    """LibreOffice and Google Docs export `w:sz` values like `26.666666666666668`.
    python-docx's typed accessor does `int(value)` on that and raises, which took the
    whole manuscript down over a single heading."""
    from docx.oxml.ns import qn

    d = docx.Document(plain)
    run = d.paragraphs[0].runs[0] if d.paragraphs[0].runs else d.paragraphs[0].add_run("x")
    rpr = run._r.get_or_add_rPr()
    sz = rpr.makeelement(qn("w:sz"), {qn("w:val"): "26.666666666666668"})
    rpr.append(sz)
    out = plain.replace(".docx", "-frac.docx")
    d.save(out)

    st = M.read_structure(out)
    # 26.666… half-points is 13.3 pt. The point is that it is read at all.
    assert st.paragraphs[0].runs[0].font_size_pt == pytest.approx(13.3, abs=0.1)


def test_an_unreadable_size_falls_back_to_what_the_run_inherits(plain):
    """A `w:sz` nobody can parse means the run states no size of its own, so the
    style and then `w:docDefaults` apply — exactly as if the attribute were absent.
    Returning None here instead would report "no size" for a run that plainly has
    one, and every body-text size check would then fire on it."""
    from docx.oxml.ns import qn

    d = docx.Document(plain)
    run = d.paragraphs[0].runs[0] if d.paragraphs[0].runs else d.paragraphs[0].add_run("x")
    rpr = run._r.get_or_add_rPr()
    rpr.append(rpr.makeelement(qn("w:sz"), {qn("w:val"): "large"}))
    out = plain.replace(".docx", "-bad.docx")
    d.save(out)

    st = M.read_structure(out)          # the point: this does not raise
    assert st.paragraphs[0].runs[0].font_size_pt == 11.0


# ----------------------------------------------------------- irregular merged tables

def test_a_row_whose_grid_does_not_resolve_still_yields_its_cells():
    """`row.cells` walks to the cell above for every vertically merged continuation
    and raises `ValueError: no 'tc' element at grid_offset=N` when the grid does not
    line up — which Word tolerates and writes anyway. 4 of 400."""
    class _Tr:
        tc_lst = ["tc0", "tc1", "tc2"]

    class _Row:
        _tr = _Tr()
        table = object()

        @property
        def cells(self):
            raise ValueError("no `tc` element at grid_offset=4")

    cells = M._row_cells(_Row())
    assert len(cells) == 3


def test_a_healthy_row_is_untouched(tmp_path):
    """The fallback must not become the normal path — it loses merge resolution."""
    d = docx.Document()
    t = d.add_table(rows=1, cols=3)
    for i in range(3):
        t.cell(0, i).text = f"c{i}"
    p = tmp_path / "t.docx"
    d.save(str(p))

    row = docx.Document(str(p)).tables[0].rows[0]
    assert [c.text for c in M._row_cells(row)] == ["c0", "c1", "c2"]


# ------------------------------------------------------------------- damaged files

def test_a_corrupt_docx_is_explained_in_words_an_author_can_act_on(tmp_path):
    """3 of 400 have a bad CRC on an embedded image. The author used to be shown
    "Bad CRC-32 for file 'word/media/image1.png'" — true, and no help at all."""
    import pipeline

    bad = tmp_path / "corrupt.docx"
    bad.write_bytes(b"this is not a zip archive")

    with pytest.raises(ValueError) as exc:
        pipeline.run_pipeline(_OPTS, str(bad), lambda *a, **k: None)

    msg = str(exc.value)
    assert "damaged" in msg
    assert "Save As" in msg
    assert not isinstance(exc.value, zipfile.BadZipFile)


# ------------------------------------------------- the warning that said nothing

def test_a_skipped_house_check_always_gives_a_reason():
    """`str(NotImplementedError())` is "", so the warning read "…skipped: " and the
    editor had no way to tell a bug from a document that genuinely had no layout."""
    import pipeline


    assert pipeline.skip_reason(NotImplementedError()) == "NotImplementedError"
    assert pipeline.skip_reason(zipfile.BadZipFile()) == "BadZipFile"
    # A real message is still preferred over the class name.
    assert pipeline.skip_reason(ValueError("no `tc` element")) == "no `tc` element"


# --------------------------------------------------- one numbering, not two

def test_every_surface_numbers_paragraphs_the_same_way():
    """`app.py` and `pipeline.py` both render `paragraph + 1`; `Finding.__str__` used
    to render `paragraph`. The same paragraph then appeared as ¶142 in one line of a
    report and ¶143 in the next, which reads as a tool that cannot count — and an
    editor sent to the wrong paragraph finds nothing wrong there."""
    import house_layout as H

    f = H.Finding("heading.size", "error", 141, "H3 is 13.0 pt, house size is 11.0 pt")
    assert "¶142" in str(f)

    # A document-wide finding has no paragraph and must not become "¶1".
    g = H.Finding("page.margin", "warning", None, "top margin is 1.2 in")
    assert "document" in str(g)


# ------------------------------------------- a degraded feature must say so

def test_trailing_text_after_the_json_no_longer_breaks_journal_ranking(monkeypatch):
    r"""`re.search(r"\{.*\}", DOTALL)` is greedy — it ran from the first brace to the
    last one anywhere in the response, so one trailing sentence produced "Extra data"
    and the manuscript silently dropped to semantic-only ranking. Observed live."""
    import editor as E

    payload = ('{"recommended_journals":[{"rank":1,"journal_id":7,'
               '"journal_name":"J Materials","overall_fit_score":81}]}\n'
               'Note: scores are approximate. {not json}')
    monkeypatch.setattr(E, "_generate_text", lambda *a, **k: payload)

    picks = E._llm_rank_journals("an abstract", [{"_cid": 7, "id": 7, "name": "J Materials"}], {})
    assert [p["journal_id"] for p in picks] == [7]


def test_a_failed_ranking_is_reported_to_the_editor(monkeypatch):
    """It used to be a `print` to the container log. The editor was shown
    semantic-only picks with nothing anywhere saying the ranking had degraded."""
    import editor as E

    monkeypatch.setattr(E, "_generate_text", lambda *a, **k: "no json here at all")
    warnings: list = []
    picks = E._llm_rank_journals("an abstract", [{"_cid": 1, "id": 1, "name": "J"}], {},
                                 warnings=warnings)
    assert picks == []
    assert warnings and "semantic matching only" in warnings[0]


def test_ranking_success_adds_no_warning(monkeypatch):
    import editor as E

    monkeypatch.setattr(E, "_generate_text", lambda *a, **k:
                        '{"recommended_journals":[{"rank":1,"journal_id":1}]}')
    warnings: list = []
    E._llm_rank_journals("an abstract", [{"_cid": 1, "id": 1, "name": "J"}], {}, warnings=warnings)
    assert warnings == []


# -------------------------------------------------- fractional twips, everywhere else

def _pPr(doc, index=0):
    """The `w:pPr` of one paragraph, created if Word did not write one."""
    p = doc.paragraphs[index]._p
    return p.get_or_add_pPr()


def test_a_fractional_page_margin_is_read_not_raised(plain, tmp_path):
    """8 of 1,606. A page converted from A4 lands as `w:left="1275.5905511811022"`,
    and `Section.left_margin` puts that through `int()`. The manuscript then lost its
    entire house-style check over the page geometry the check is largely about."""
    d = docx.Document(plain)
    pgMar = d.sections[0]._sectPr.find(qn("w:pgMar"))
    pgMar.set(qn("w:left"), "1275.5905511811022")
    out = tmp_path / "frac-margin.docx"
    d.save(str(out))

    st = M.read_structure(str(out))
    assert st.sections[0].left_margin_in == round(1275.5905511811022 / 1440, 3)


def test_a_fractional_first_line_indent_is_read_not_raised(plain, tmp_path):
    d = docx.Document(plain)
    ind = _pPr(d).get_or_add_ind()
    ind.set(qn("w:firstLine"), "216.00000000000003")
    out = tmp_path / "frac-indent.docx"
    d.save(str(out))

    st = M.read_structure(str(out))
    assert st.paragraphs[0].first_line_in == 0.15


def test_a_hanging_indent_keeps_its_sign(plain, tmp_path):
    """`w:hanging` is the same measurement negated — the reference list's 0.25" hang.
    Reading it unsigned would turn a hanging indent into a first-line one."""
    d = docx.Document(plain)
    ind = _pPr(d).get_or_add_ind()
    ind.set(qn("w:hanging"), "360.00000000000006")
    out = tmp_path / "hanging.docx"
    d.save(str(out))

    st = M.read_structure(str(out))
    assert st.paragraphs[0].first_line_in == -0.25


def test_fractional_paragraph_spacing_is_points_not_inches(plain, tmp_path):
    """Word writes spacing in twips and the house spec is in points: 20 twips to the
    point, 1440 to the inch. Crossing the two conversions makes the checker
    confidently wrong rather than silent."""
    d = docx.Document(plain)
    spacing = _pPr(d).get_or_add_spacing()
    spacing.set(qn("w:before"), "359.00000000000006")
    out = tmp_path / "frac-spacing.docx"
    d.save(str(out))

    st = M.read_structure(str(out))
    assert st.paragraphs[0].space_before_pt == 18.0


def test_the_newer_alignment_spelling_is_understood(plain, tmp_path):
    """Word writes the direction-neutral `w:jc val="start"`; python-docx's enum knows
    only `left` and raises. One paragraph like this cost a whole manuscript."""
    d = docx.Document(plain)
    jc = _pPr(d).get_or_add_jc()
    jc.set(qn("w:val"), "start")
    out = tmp_path / "jc-start.docx"
    d.save(str(out))

    st = M.read_structure(str(out))
    assert st.paragraphs[0].alignment == "left"


def test_an_unknown_alignment_is_unknown_not_a_crash(plain, tmp_path):
    d = docx.Document(plain)
    _pPr(d).get_or_add_jc().set(qn("w:val"), "distribute")
    out = tmp_path / "jc-odd.docx"
    d.save(str(out))

    assert M.read_structure(str(out)).paragraphs[0].alignment is None


def test_a_failed_job_never_reaches_the_user_with_an_empty_reason(monkeypatch):
    """`app.py` renders "Processing failed for **paper.docx**: {error_message}" and
    the worker passed it `str(exc)` — empty for `NotImplementedError`, which is what
    python-docx raises on a document that has never contained a list. The user got a
    failure with nothing after the colon, and the history row was blank too, so there
    was nowhere else to look. Same bug as the house-style warning, on the path that
    actually stops someone's work."""
    import pipeline

    import json as _json

    recorded = {}
    monkeypatch.setattr(pipeline.auth, "fail_job",
                        lambda jid, msg: recorded.setdefault("fail", msg))
    # `log_job`'s error message is its last positional argument.
    monkeypatch.setattr(pipeline.auth, "log_job",
                        lambda *a, **k: recorded.setdefault("history", a[-1]))
    monkeypatch.setattr(pipeline, "run_pipeline",
                        lambda *a, **k: (_ for _ in ()).throw(NotImplementedError))

    pipeline._process_job({
        "id": 1,
        "input_path": "/nonexistent.docx",
        "options_json": _json.dumps(dict(_OPTS, filename="paper.docx")),
    })

    assert recorded["fail"] == "NotImplementedError"
    assert recorded["history"] == "NotImplementedError"


# ----------------------------------------------- the copyedit deleting the manuscript

_GEAR = ("2.1 Gear Design and Geometry A 24-tooth spur gear was modeled in Autodesk "
         "Fusion 360 following ISO metric gear design principles. The selected "
         "parameters were:")


def _chunk_returning(monkeypatch, payload):
    """Run `ai_edit_chunk` against a fixed model response."""
    import editor as E
    monkeypatch.setattr(E, "_generate_text",
                        lambda *a, **k: json.dumps(payload), raising=True)
    return E.ai_edit_chunk([_GEAR], {}, "CMOS", "Vancouver", "US English", "", False)


def test_a_truncated_paragraph_is_refused_not_applied(monkeypatch):
    """Measured on a real manuscript: the model read the run-on section heading, kept
    it, and threw the two sentences after it away — with no query. `ai_edit_chunk`
    checked the array length and the type of `edited` and nothing else, so it reached
    the redline looking like a deletion the editor had asked for."""
    edited, queries, failure = _chunk_returning(
        monkeypatch, [{"edited": "Gear Design and Geometry"}])

    assert failure is None
    assert edited == [_GEAR], "the author's paragraph must survive"
    assert len(queries) == 1
    assert queries[0]["local_index"] == 0
    assert "left unchanged" in queries[0]["query"]


def test_an_emptied_paragraph_is_refused(monkeypatch):
    edited, queries, _ = _chunk_returning(monkeypatch, [{"edited": "   "}])
    assert edited == [_GEAR]
    assert len(queries) == 1


def test_an_ordinary_copyedit_is_applied_untouched(monkeypatch):
    """Narrowed, not blocking: a real edit must still go through, with its own query."""
    fixed = _GEAR.replace("modeled", "modelled")
    edited, queries, _ = _chunk_returning(
        monkeypatch, [{"edited": fixed, "query": "US or UK spelling?",
                       "suggestion": "modeled"}])
    assert edited == [fixed]
    assert len(queries) == 1 and queries[0]["query"] == "US or UK spelling?"


def test_a_short_heading_may_lose_its_number(monkeypatch):
    """The guard must not fire on headings and captions, where a large proportional
    cut is usually the correct edit."""
    import editor as E
    monkeypatch.setattr(E, "_generate_text",
                        lambda *a, **k: json.dumps([{"edited": "Material characteristics"}]),
                        raising=True)
    edited, queries, _ = E.ai_edit_chunk(
        ["2.2 Material characteristics"], {}, "CMOS", "Vancouver", "US English", "", False)
    assert edited == ["Material characteristics"]
    assert queries == []


def test_an_emptied_paragraph_is_caught_at_any_length():
    """The short-paragraph exemption is about *proportional* cuts on headings. It was
    never meant to allow deleting a paragraph outright, and on 2026-09-04 it did: an
    85-character sentence carrying a citation came back empty and reached the redline
    with nothing reported — no error, no skipped chunk, no warning."""
    import editor
    short_sentence = ("real-time, flagging unusual behavior and preventing "
                      "unauthorized access attempts [6].")
    assert len(short_sentence) < editor._GUARD_MIN_CHARS
    assert editor._lost_content(short_sentence, "") is True
    assert editor._lost_content(short_sentence, "   ") is True


def test_a_heading_may_still_lose_its_number():
    """The exemption must keep working for what it was for."""
    import editor
    assert editor._lost_content("2.2 Material characteristics",
                                "Material characteristics") is False
