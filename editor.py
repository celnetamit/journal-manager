"""Document editing and LLM-driven helpers.

Uses provider-specific adapters for Gemini, OpenRouter, and
OpenAI-compatible endpoints. All API key / path lookups go through
`config.py`.
"""
from __future__ import annotations

import datetime
import difflib
import json
import os
import re
import threading
import time
import concurrent.futures
import xml.etree.ElementTree as ET
from typing import Any, Dict, List, Optional, Tuple

import docx
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import numpy as np
import requests

import config as app_config


def _normalize_settings(settings: Optional[Dict[str, Any]] = None) -> Dict[str, str]:
    return app_config.normalize_llm_settings(settings or app_config.get_llm_settings())


def _llm_temperature() -> float:
    """Sampling temperature for copyediting/review generation. Defaults to a low
    value so the SAME input produces consistent corrections across runs and
    across parallel chunks (copyediting is a deterministic task, not creative
    writing). Override with LLM_TEMPERATURE if needed."""
    try:
        return max(0.0, float(os.getenv("LLM_TEMPERATURE", "0.1")))
    except (TypeError, ValueError):
        return 0.1


# --- Concurrency / rate-limit control ---
#
# A single process-wide semaphore caps how many LLM text/embedding calls are
# in flight at once, shared across the copyedit worker pool AND the parallel
# AI reviewer. This keeps peak load on the provider fixed no matter how many
# tasks run concurrently, so adding the reviewer can't push us past the
# provider's concurrency ceiling. Tunable per provider tier via env.
def _llm_max_concurrency() -> int:
    try:
        return max(1, int(os.getenv("LLM_MAX_CONCURRENCY", "3")))
    except (TypeError, ValueError):
        return 3


def _chunk_pool_size() -> int:
    """Per-job paragraph-chunk worker pool. Defaults to the global LLM
    concurrency cap (no point spawning more threads than the semaphore allows),
    overridable via LLM_CHUNK_WORKERS for finer control."""
    raw = os.getenv("LLM_CHUNK_WORKERS")
    if raw:
        try:
            return max(1, int(raw))
        except (TypeError, ValueError):
            pass
    return _llm_max_concurrency()


_LLM_SEMAPHORE = threading.Semaphore(_llm_max_concurrency())

# Max retries for a 429 (Too Many Requests) before giving up, and the initial
# backoff delay in seconds (doubled after each failed attempt).
_RATE_LIMIT_MAX_RETRIES = int(os.getenv("LLM_RATE_LIMIT_RETRIES", "4"))
_RATE_LIMIT_BASE_DELAY = float(os.getenv("LLM_RATE_LIMIT_BASE_DELAY", "2.0"))


# --- Provider helpers ---

def _client(api_key: Optional[str] = None):
    return app_config.get_gemini_client(api_key=api_key)


def _retry_after_seconds(resp: "requests.Response", fallback: float) -> float:
    """Honor a numeric Retry-After header (seconds) when the server provides one."""
    header = resp.headers.get("Retry-After")
    if header:
        try:
            return max(0.0, float(header))
        except (TypeError, ValueError):
            pass
    return fallback


def _post_json(url: str, payload: Dict[str, Any], headers: Optional[Dict[str, str]] = None,
               timeout: int = 120) -> Dict[str, Any]:
    delay = _RATE_LIMIT_BASE_DELAY
    for attempt in range(_RATE_LIMIT_MAX_RETRIES):
        resp = requests.post(url, json=payload, headers=headers or {}, timeout=timeout)
        # On 429 (Too Many Requests), wait and retry with exponential backoff,
        # preferring the server's Retry-After hint when present.
        if resp.status_code == 429:
            if attempt < _RATE_LIMIT_MAX_RETRIES - 1:
                wait = _retry_after_seconds(resp, delay)
                print(f"Rate limited (429); retrying in {wait:.1f}s "
                      f"(attempt {attempt + 1}/{_RATE_LIMIT_MAX_RETRIES})")
                time.sleep(wait)
                delay *= 2
                continue
            raise RuntimeError(
                f"Rate limited (429): still throttled after "
                f"{_RATE_LIMIT_MAX_RETRIES} attempts."
            )
        resp.raise_for_status()
        data = resp.json()
        if not isinstance(data, dict):
            raise RuntimeError("Unexpected non-object JSON response")
        return data
    # Unreachable: the loop always returns or raises above.
    raise RuntimeError("Rate limited (429): exceeded retry attempts.")


def _generate_text(prompt: str, settings: Optional[Dict[str, Any]] = None,
                   response_mime_type: Optional[str] = None) -> str:
    """Generate text through the selected provider."""
    cfg = _normalize_settings(settings)
    provider = cfg["provider"]

    # Cap concurrent provider calls (shared across the copyedit pool and the
    # parallel AI reviewer) so total in-flight load stays bounded.
    with _LLM_SEMAPHORE:
        if provider == "gemini":
            from google.genai import types

            # Always pin a low temperature so corrections are consistent run to
            # run and across parallel chunks; carry the JSON mime type when asked.
            config_kwargs: Dict[str, Any] = {"temperature": _llm_temperature()}
            if response_mime_type:
                config_kwargs["response_mime_type"] = response_mime_type
            kwargs = {"config": types.GenerateContentConfig(**config_kwargs)}
            # Hold the client in a local so it stays alive for the SDK's internal
            # retry loop; a temporary would be GC'd mid-call and close its httpx
            # transport ("Cannot send a request, as the client has been closed.").
            client = _client(cfg["api_key"])
            resp = client.models.generate_content(
                model=cfg["text_model"], contents=prompt, **kwargs,
            )
            time.sleep(1)
            return (resp.text or "").strip()

        if provider in {"openrouter", "openai-compatible", "custom"}:
            base = cfg["base_url"].rstrip("/")
            if not base:
                raise RuntimeError("Base URL is required for this provider.")
            payload: Dict[str, Any] = {
                "model": cfg["text_model"],
                "messages": [{"role": "user", "content": prompt}],
                "temperature": _llm_temperature(),
            }
            if response_mime_type == "application/json":
                payload["response_format"] = {"type": "json_object"}
            headers = {"Content-Type": "application/json"}
            if cfg["api_key"]:
                headers["Authorization"] = f"Bearer {cfg['api_key']}"
            data = _post_json(f"{base}/chat/completions", payload, headers=headers)
            try:
                text = data["choices"][0]["message"]["content"]
            except Exception as exc:
                raise RuntimeError(f"Unexpected chat completion response: {data}") from exc
            time.sleep(0.6)
            return (text or "").strip()

        raise RuntimeError(f"Unsupported LLM provider: {provider}")


def _embed(text: str, settings: Optional[Dict[str, Any]] = None) -> List[float]:
    cfg = _normalize_settings(settings)
    provider = cfg["provider"]

    # Share the same concurrency cap as text generation.
    with _LLM_SEMAPHORE:
        if provider == "gemini":
            client = _client(cfg["api_key"])
            resp = client.models.embed_content(
                model=cfg["embed_model"], contents=text,
            )
            time.sleep(0.3)
            if not resp.embeddings:
                raise RuntimeError("Empty embedding response")
            return list(resp.embeddings[0].values)

        if provider in {"openrouter", "openai-compatible", "custom"}:
            base = cfg["base_url"].rstrip("/")
            if not base:
                raise RuntimeError("Base URL is required for this provider.")
            payload = {"model": cfg["embed_model"], "input": text}
            headers = {"Content-Type": "application/json"}
            if cfg["api_key"]:
                headers["Authorization"] = f"Bearer {cfg['api_key']}"
            data = _post_json(f"{base}/embeddings", payload, headers=headers, timeout=180)
            try:
                vector = data["data"][0]["embedding"]
            except Exception as exc:
                raise RuntimeError(f"Unexpected embeddings response: {data}") from exc
            time.sleep(0.3)
            return list(vector)

        raise RuntimeError(f"Unsupported LLM provider: {provider}")


def test_text_generation(prompt: str, settings: Optional[Dict[str, Any]] = None) -> str:
    """Run a small text-generation probe with the selected provider."""
    prompt = prompt.strip()
    if not prompt:
        raise ValueError("Prompt cannot be empty.")
    return _generate_text(prompt, settings=settings)


def test_embedding(text: str, settings: Optional[Dict[str, Any]] = None) -> List[float]:
    """Run an embedding probe with the selected provider."""
    text = text.strip()
    if not text:
        raise ValueError("Text cannot be empty.")
    return _embed(text, settings=settings)


# --- High-level operations ---

def align_global_citations(
    paras: List[str], settings: Dict[str, Any], ref_style: str,
    enabled_rule_ids: Optional[List[str]] = None,
) -> List[str]:
    """Second pass: convert in-text citations to sequential [1], [2] and
    re-sort the bibliography to match."""
    indexed_paras = {str(i): p for i, p in enumerate(paras) if p.strip()}

    citation_format_block = ""
    if enabled_rule_ids is None or "citation" in enabled_rule_ids:
        citation_format_block = (
            "\nIN-TEXT CITATION FORMATTING (apply to every in-text bracketed citation you output):\n"
            '- Put a single space after each comma when listing multiple citations, e.g. "[1, 2, 3]" (NOT "[1,2,3]").\n'
            '- For a consecutive range, use an en dash (–) with no surrounding spaces, e.g. "[3–7]" (NOT "[3-7]" or "[3—7]").\n'
        )

    prompt = f"""You are a strict reference alignment engine.
The selected reference style is {ref_style}.

Task:
1. Scan the provided manuscript text for any in-text citations (e.g., "(Smith et al., 2020)").
2. Convert all in-text citations to sequential bracketed numbers (e.g., "[1]", "[1, 2, 3]") based on the EXACT order they first appear in the text.
3. Locate the bibliography/reference list at the end of the text.
4. Re-order and re-number the actual bibliography entries so that they match the new numerical sequence of the in-text citations.
{citation_format_block}
CRITICAL OUTPUT INSTRUCTIONS:
- Return ONLY a valid JSON dictionary.
- The keys MUST be the original paragraph index strings.
- The values MUST be the updated paragraph text.
- ONLY include paragraphs that you modified (i.e., paragraphs containing an in-text citation, or bibliography entries).
- Do NOT include paragraphs that were untouched.
- Do NOT wrap the JSON in markdown code blocks.

Input JSON dictionary (Key = Index, Value = Paragraph Text):
{json.dumps(indexed_paras)}
"""
    try:
        text = _generate_text(prompt, settings=settings, response_mime_type="application/json")
        match = re.search(r"\{.*\}", text, re.DOTALL)
        if not match:
            print("Warning: No JSON dict found for global citations.")
            return paras
        updates = json.loads(match.group(0))
        new_paras = list(paras)
        for idx_str, new_text in updates.items():
            idx = int(idx_str)
            if 0 <= idx < len(new_paras):
                new_paras[idx] = new_text
        return new_paras
    except Exception as e:
        print(f"Global citation alignment failed: {e}")
        return paras


def read_docx(file_path: str) -> List[str]:
    doc = docx.Document(file_path)
    return [p.text for p in doc.paragraphs]


# --- Serper.dev (Google Scholar) DOI fallback — OPTIONAL ---
# Used only when a Serper API key is configured. It is a fallback for Crossref:
# when Crossref cannot find a reference's DOI, we search Google Scholar via
# Serper. Returns the same "https://doi.org/<DOI>" format as fetch_crossref_doi.

_SERPER_SCHOLAR_URL = "https://google.serper.dev/scholar"
_DOI_RE = re.compile(r"10\.\d{4,9}/[-._;()/:A-Za-z0-9]+")


def verify_serper_key(api_key: str) -> Tuple[bool, str]:
    """Validate a Serper API key with one tiny query. Returns (ok, message).
    Called once per run so a bad/expired token is reported to the user instead
    of silently failing on every reference."""
    if not api_key:
        return False, "No Serper API key configured."
    try:
        r = requests.post(
            _SERPER_SCHOLAR_URL,
            headers={"X-API-KEY": api_key, "Content-Type": "application/json"},
            json={"q": "test"}, timeout=15,
        )
        if r.status_code in (401, 403):
            return False, (
                f"Serper API key was rejected (HTTP {r.status_code}). "
                "Check the SERPER_API_KEY token."
            )
        if r.status_code == 429:
            return False, "Serper API quota exceeded (HTTP 429)."
        if r.status_code != 200:
            return False, f"Serper API returned HTTP {r.status_code}."
        return True, "ok"
    except requests.exceptions.RequestException as e:
        return False, f"Serper API unreachable: {e}"


def fetch_serper_scholar_doi(citation_text: str, api_key: str) -> Optional[str]:
    """Search Google Scholar (via Serper) for a citation and return its DOI as
    'https://doi.org/<DOI>', or None. Conservative: only returns a DOI from a
    result whose title meaningfully overlaps the citation, to avoid attaching a
    wrong DOI."""
    if not api_key:
        return None
    try:
        r = requests.post(
            _SERPER_SCHOLAR_URL,
            headers={"X-API-KEY": api_key, "Content-Type": "application/json"},
            json={"q": citation_text[:300]}, timeout=20,
        )
        if r.status_code != 200:
            return None
        organic = (r.json() or {}).get("organic") or []
        clean_cit = re.sub(r"[^a-z0-9]", "", citation_text.lower())
        for item in organic[:3]:
            title = item.get("title") or ""
            clean_title = re.sub(r"[^a-z0-9]", "", title.lower())
            if not clean_title:
                continue
            # Require title overlap in either direction before trusting a DOI.
            if clean_title[:30] not in clean_cit and clean_cit[:30] not in clean_title:
                continue
            blob = " ".join(
                str(item.get(k, "")) for k in ("doi", "link", "snippet", "title")
            )
            m = _DOI_RE.search(blob)
            if m:
                return "https://doi.org/" + m.group(0).rstrip(").,;")
    except requests.exceptions.RequestException:
        return None
    return None


def fetch_crossref_doi(citation_text: str) -> Optional[str]:
    try:
        url = "https://api.crossref.org/works"
        params = {
            "query.bibliographic": citation_text,
            "select": "DOI,score,type,title",
            "rows": 3,
        }
        r = requests.get(url, params=params, timeout=4)
        if r.status_code != 200:
            return None
        items = (r.json().get("message") or {}).get("items") or []

        for item in items:
            item_type = (item.get("type") or "").lower()
            if "review" in item_type and "review" not in citation_text.lower():
                continue

            titles = item.get("title") or []
            if titles:
                clean_title = re.sub(r"[^a-z0-9]", "", titles[0].lower())
                clean_cit = re.sub(r"[^a-z0-9]", "", citation_text.lower())
                if clean_title not in clean_cit and item.get("score", 0) < 85:
                    continue

            if item.get("score", 0) > 55:
                return "https://doi.org/" + item["DOI"]
    except Exception:
        return None
    return None


# --- Preliminary originality (web-match) scan via Serper — OPTIONAL ---
# IMPORTANT: this is NOT a real plagiarism check. It only finds VERBATIM matches
# on the public web (via exact-phrase Google search). It cannot see paywalled
# journals or paraphrased text and produces no formal similarity score. It is a
# cheap, preliminary signal — not a substitute for iThenticate/Turnitin.

_SERPER_SEARCH_URL = "https://google.serper.dev/search"
_SENT_SPLIT_RE = re.compile(r"(?<=[.!?])\s+")
PLAGIARISM_DISCLAIMER = (
    "Preliminary web-match scan only — finds verbatim text on the public web. "
    "It does NOT cover paywalled/journal content or paraphrasing and is NOT a "
    "substitute for iThenticate/Turnitin."
)


def _looks_like_reference(sentence: str) -> bool:
    """A bibliography/citation line legitimately matches published sources, so we
    skip it — counting it as 'plagiarism' would be wrong."""
    has_year = re.search(r"\b(19|20)\d{2}\b", sentence)
    if has_year and re.search(r"\bet\s+al\b", sentence, re.I):
        return True
    if has_year and sentence.count(",") >= 4:
        return True
    if "doi.org" in sentence.lower() or "https://" in sentence.lower():
        return True
    return False


def _candidate_sentences(paragraphs: List[str], min_words: int = 10,
                         max_words: int = 45, max_checks: int = 25) -> List[str]:
    """Pick substantive body sentences worth checking. Skips headings, the
    keywords line, and reference/citation lines, then samples evenly across the
    document so coverage is spread out and the query budget is bounded."""
    sents: List[str] = []
    for p in paragraphs:
        t = (p or "").strip()
        if not t or t.lower().startswith(("keyword", "key word")):
            continue
        # Skip reference/bibliography paragraphs wholesale — sentence-splitting a
        # reference (e.g. on the period after "et al.") would otherwise leak a
        # citation fragment past the per-sentence guard.
        if _looks_like_reference(t):
            continue
        for s in _SENT_SPLIT_RE.split(t):
            s = s.strip()
            if min_words <= len(s.split()) <= max_words and not _looks_like_reference(s):
                sents.append(s)
    if len(sents) > max_checks:
        step = len(sents) / max_checks
        sents = [sents[int(i * step)] for i in range(max_checks)]
    return sents


def _serper_search_phrase(phrase: str, api_key: str) -> Optional[List[Dict[str, Any]]]:
    """Exact-phrase Google search via Serper. Returns the organic results list
    (possibly empty) or None on a network/API error."""
    try:
        r = requests.post(
            _SERPER_SEARCH_URL,
            headers={"X-API-KEY": api_key, "Content-Type": "application/json"},
            json={"q": f'"{phrase}"', "num": 5}, timeout=20,
        )
        if r.status_code != 200:
            return None
        return (r.json() or {}).get("organic") or []
    except requests.exceptions.RequestException:
        return None


def plagiarism_scan(paragraphs: List[str], api_key: str,
                    max_checks: int = 25, max_workers: int = 5) -> Dict[str, Any]:
    """Preliminary web-match originality scan. For a bounded sample of body
    sentences, exact-phrase search the web; a sentence with any verbatim hit is
    flagged with its top source links. Returns a dict with an originality score
    (% of checked sentences with NO web match) and the flagged sentences.

    Optional: with no api_key this no-ops and reports it. See PLAGIARISM_DISCLAIMER
    — this is a preliminary signal, not a formal plagiarism report."""
    result: Dict[str, Any] = {
        "available": False, "checked": 0, "flagged": [], "score": None,
        "disclaimer": PLAGIARISM_DISCLAIMER, "note": "",
    }
    if not api_key:
        result["note"] = "No Serper API key — originality scan skipped."
        return result

    sentences = _candidate_sentences(paragraphs, max_checks=max_checks)
    if not sentences:
        result["note"] = "No substantive body sentences found to scan."
        return result

    def check(s: str) -> Optional[Dict[str, Any]]:
        organic = _serper_search_phrase(s, api_key)
        if organic:  # exact-phrase query returned hits => verbatim match on web
            sources = [
                {"title": o.get("title", ""), "link": o.get("link", "")}
                for o in organic[:3] if o.get("link")
            ]
            return {"sentence": s, "sources": sources}
        return None

    flagged: List[Dict[str, Any]] = []
    with concurrent.futures.ThreadPoolExecutor(max_workers=max_workers) as ex:
        for res in ex.map(check, sentences):
            if res:
                flagged.append(res)

    checked = len(sentences)
    result.update(
        available=True, checked=checked, flagged=flagged,
        score=round(100 * (1 - len(flagged) / checked)) if checked else None,
    )
    return result


# --- In-house / publisher rules (displayed in UI; injected into the edit prompt) ---
#
# Each group can be toggled on/off per user and is surfaced in the frontend.
# `body` is the exact text injected into the copyediting prompt when the group
# is enabled. Add new built-in groups here and they default to ON for everyone.

HOUSE_RULE_GROUPS: List[Dict[str, str]] = [
    {
        "id": "reference",
        "title": "In-House Reference Rules",
        "summary": "Author limit (6 + et al.), strip honorific titles, ISO-4 journal abbreviation, year-only publication date, source-type structure.",
        "body": """IN-HOUSE REFERENCE RULES (these OVERRIDE the base reference style above for any reference/bibliography entry):
0. PUBLICATION DATE — YEAR ONLY: The publication date of a reference must show the YEAR ONLY (a 4-digit year), never a month, season, or day. Remove any month name or abbreviation, season, or day that is attached to the publication year, and keep only the year. Examples: "Vol. 175, July.2025" -> "Vol. 175, 2025"; "June 2025" -> "2025"; "Jan 2024" -> "2024"; "Apr.2021" -> "2021"; "Spring 2019" -> "2019"; "15 March 2020" -> "2020". Also drop any now-stray punctuation left behind (e.g. the "." in "July.2025"). TWO EXCEPTIONS where a month is REQUIRED and must be kept: (a) the event date of a CONFERENCE reference ("Year, Month Date.") and (b) the "[Accessed on Month Year]" date of a WEB PAGE reference — only these two retain the month. Do NOT touch months that appear inside the article/book/chapter TITLE itself (e.g. a paper titled "Trends in March 2020 Lockdowns").
1. AUTHOR LIMIT (6 + et al.): Count the authors in each reference. If a reference lists MORE THAN 6 authors, keep ONLY the first 6 authors in their original order, DELETE every author after the 6th, and append "et al." after the 6th author. The output must contain EXACTLY 6 names followed by "et al." — never 7 or more names before "et al.". Example with 8 authors: "Smith A, Jones B, Lee C, Patel D, Garcia E, Kim F, Brown G, Davis H" -> "Smith A, Jones B, Lee C, Patel D, Garcia E, Kim F, et al." (authors 7 and 8 are removed). If there are 6 or fewer authors, list them ALL and do NOT add "et al.". Never add "et al." to a list that already shows all authors.
1a. STRIP TITLES: Remove any honorific or professional title prefixes attached to an author name (e.g. "Mr", "Mrs", "Ms", "Dr", "Dr.", "Prof", "Prof.", "Professor", "Er", "Er.", "Sir", "Sri", "Smt", "Md" when used as a courtesy title, "PhD"/"MD"/"FRCS" when used as a leading prefix). Keep only the actual name. Example: "Dr. Smith JA, Prof Jones BR" -> "Smith JA, Jones BR". Do this in BOTH the bibliography entries and any in-text author-date citations.
2. JOURNAL ABBREVIATION: For JOURNAL references, abbreviate the journal name using standard ISO 4 / Index Medicus (NLM) abbreviations (e.g., "Journal of Science" -> "J Sci"; "New England Journal of Medicine" -> "N Engl J Med"; "Nature" stays "Nature"). Do NOT abbreviate book titles, publisher names, or chapter titles.
3. SOURCE-TYPE STRUCTURE: Identify each reference's type and format it with EXACTLY this punctuation and order (journal house style):
   - JOURNAL ARTICLE: Authors. Article title. Abbreviated Journal Title. Year; Volume (Issue): start–endp. doi: <doi>.
     (page range uses an en dash and a trailing "p"; include the doi if present)
     Example: Hsu SC, Wang SJ, Liu CY, Juang YY, Yang CH, Hung CI. The impact of anxiety and migraine on quality of sleep in patients with major depressive disorder. Compr Psychiatry. 2009; 50 (2): 151–157p. doi: 10.1016/j.comppsych.2008.07.002.
   - BOOK: Authors. Book Title. Publisher Location: Publisher Name; Year. (do NOT abbreviate the title)
     Example: Singh S. Textbook of Medical Surgical Nursing. New Delhi: Wolters Kluwer; 2005.
   - EDITED BOOK / CHAPTER: Authors. Chapter title, In: Editors, editors. Book Title. Edition. Publisher Location: Publisher Name; Year. pp. start–end.
     Example: Singh S, Sharma A. Surgery in nursing, In: Janice L. Hinkle, Kerry H. Cheever, editor. Brunner & Suddarth's Textbook of Medical-Surgical Nursing. 14th edition. New Delhi: Wolters Kluwer; 2005. pp. 2198–2201.
   - CONFERENCE: Authors. Article title, Title of Conference, Location of Conference. Year, Month Date.
     Example: Bowden FJ, Fairley CK. Endemic STDs in the Northern Territory: estimations of effective rates of partner exchange. The Scientific Meeting of the Royal Australian College of Physicians. Darwin, Australia. 1996, June 24–25.
   - WEB PAGE (ONLINE): Authors. (Year). Title of article. [Online] Title of Webpage. Available at URL [Accessed on Month Year].
     Example: World Health Organization. (2015). Drinking water. [Online] World Health Organization. Available at http://www.who.int/mediacentre/factsheets/fs391/en/ [Accessed on July 2015].
   IMPORTANT: rules 1 and 2 ALWAYS apply on top of these formats — cap the author list at 6 + "et al." and abbreviate journal titles to ISO-4, even though this journal's own stylesheet otherwise lists all authors and full journal titles.
   Only treat an item as a reference if it is clearly a citation/bibliography entry; do NOT restructure ordinary body text.""",
    },
    {
        "id": "heading",
        "title": "In-House Heading Rules",
        "summary": "Level-based case (H1 ALL CAPS, sub-levels Title Case), remove leading numbering, preserve hierarchy, no trailing period/colon.",
        "body": """IN-HOUSE HEADING RULES (apply to any paragraph that is a section/subsection heading):
1. CASE BY LEVEL:
   - Top-level / main section headings (Level 1, e.g. "INTRODUCTION", "MATERIALS AND METHODS", "RESULTS", "DISCUSSION", "CONCLUSION") -> ALL CAPS.
   - Sub-level headings (Levels 2, 3, 4) -> Title Case (capitalize the first and last word and all major words; keep short articles/conjunctions/prepositions such as "a", "an", "the", "of", "and", "in", "for", "to" lowercase unless first/last). Example: "fundamentals of functional genomics" -> "Fundamentals of Functional Genomics".
   (The journal styles these visually as H1 all-caps bold, H2 bold, H3 bold italic, H4 italic — apply the CASING here; the bold/italic is the journal's visual style applied at layout.)
2. NO NUMBERING: Remove any leading numbering or auto-number prefixes from headings (e.g. "1. Introduction", "1.2 Methods", "Chapter 3: Results", "IV. Discussion" -> "INTRODUCTION", "Methods", "Results", "Discussion"). Keep the heading text itself intact.
3. STRUCTURE: Preserve the heading's hierarchy/level and order; only fix its capitalization and remove numbering. Do NOT merge a heading into body text or turn body text into a heading. Do NOT add a trailing period or colon to a heading.
4. RUN-IN LABELS: A run-in heading or label that opens a paragraph (e.g. "OBJECTIVE:", "Background:", "Aim:" followed by body text on the same line) must be KEPT, not deleted. Reformat its case per the level rules and you MAY keep its colon (a run-in label is not a standalone heading), but never remove the label word itself.""",
    },
    {
        "id": "citation",
        "title": "In-House In-Text Citation Rules",
        "summary": "Single space after commas; en dash for consecutive ranges; drop the redundant parenthetical author-date when a numbered marker is present (narrative citations untouched).",
        "body": """IN-HOUSE IN-TEXT CITATION RULES (apply to citation markers that appear inside body sentences, e.g. "[1]", "[1,2]", "[3-7]"):
1. SPACE AFTER COMMA: When multiple citations are listed, put a single space after each comma, e.g. "[1,2,3]" -> "[1, 2, 3]".
2. EN-DASH FOR RANGES: For a consecutive citation range, use an en dash (–) with no surrounding spaces between the first and last number, e.g. "[3-7]", "[3 to 7]", "[3—7]" -> "[3–7]".
3. REDUNDANT PARENTHETICAL AUTHOR-DATE: This journal uses a NUMBERED citation system (the "[n]" marker carries the citation). When an author-date appears INSIDE PARENTHESES right next to a numbered marker — e.g. "...older approaches (Hasan et al., 2025) [1]." or "...as shown [1] (Hasan et al., 2025)." — the parenthetical "(Author, Year)" is redundant with the number, so DELETE the parenthetical author-date and KEEP the "[n]" marker: "...older approaches [1].". Remove the now-empty parentheses and fix spacing/punctuation so exactly one space sits before "[n]" and the sentence still ends cleanly.
   CRITICAL — NARRATIVE CITATIONS ARE EXEMPT: Only remove an author-date that is genuinely PARENTHETICAL (enclosed in "(...)") and is NOT the grammatical subject/object of the sentence. If the author name is woven into the running text as the subject — e.g. "Hasan et al. (2025) [1] offered a detailed examination." — KEEP the author name; that is a narrative citation and removing it would break the sentence (here you may still keep the "(2025)" and the "[1]"). When unsure whether removing the parenthetical leaves a grammatical, meaning-preserving sentence, do NOT remove it — leave both and, if needed, raise a query.
4. Do NOT change the citation numbers themselves or alter reference-list (bibliography) entries with these rules; they apply only to in-text citation markers.""",
    },
    {
        "id": "citation_tense",
        "title": "In-House Citation Tense Rules",
        "summary": "ONLY the single reporting verb directly attached to a citation opener (\"Author et al. (Year) [n] <verb>\") is put into simple past; the rest of the paragraph is never touched.",
        "body": """IN-HOUSE CITATION TENSE RULES (VERY NARROW SCOPE — read carefully):
This rule applies to EXACTLY ONE verb: the reporting verb that directly follows a citation subject at the START of a sentence, where the sentence opens with the literal pattern "Author et al. (Year) [n] <verb> ..." (an author name + year + bracketed number, then the verb). Nothing else.
1. REPORTING VERB TENSE: In such an opener, put that ONE reporting verb into the SIMPLE PAST tense, because it describes a completed action in a specific past study. Examples (note only the opener verb changes): "Khan et al. (2024) [8] point out key challenges" -> "Khan et al. (2024) [8] pointed out key challenges"; "Huo et al. (2023) [6] introduces an intelligent system" -> "Huo et al. (2023) [6] introduced an intelligent system"; "Kielhauser et al. (2020) [5] offering a hands-on example" -> "Kielhauser et al. (2020) [5] offered a hands-on example".
2. DO NOT CASCADE: This is the single most important limit. Do NOT change the tense of ANY other verb in the paragraph. Every later sentence — even though it describes the same cited study (e.g. "The study involves...", "These images get processed...", "The researchers stack...", "All this helps...", "The team notes...") — MUST keep the author's original tense exactly. A literature-review paragraph is about one study throughout, but that does NOT make the whole paragraph reportable: change only the one opener verb and leave the remaining sentences completely untouched.
3. "CONSISTENT" MEANS ACROSS OPENERS, NOT WITHIN A PARAGRAPH: Apply the simple-past opener fix the same way to each DIFFERENT citation opener (typically one per related-work paragraph). It does NOT mean making all verbs within a paragraph share one tense.
4. NEVER touch general, established, or still-true statements, methodology descriptions, or any sentence whose subject is not an explicit "Author et al. (Year) [n]" citation. When in doubt, leave the verb unchanged.
5. Do NOT alter, MOVE, or DELETE the author names, the year, or the citation number — the "Author et al. (Year)" phrase and any "[n]" marker MUST remain exactly in place. Do NOT rewrite the sentence or remove its subject.""",
    },
    {
        "id": "abbreviation",
        "title": "In-House Abbreviation Expansion Rules",
        "summary": "Expand element labels (Fig.→Figure, Tab.→Table, Eq.→Equation, etc.) to their full word when they refer to a numbered element.",
        "body": """IN-HOUSE ABBREVIATION EXPANSION RULES (apply to body text AND figure/table captions; spell out abbreviated labels that refer to a numbered manuscript element so they read unambiguously):
1. EXPAND TO THE FULL WORD when the abbreviation refers to a specific element (i.e. it is followed by a number, a letter, or a number range, with or without a period). Use these exact, capitalized full forms and KEEP the identifier that follows:
   - "Fig", "Fig.", "fig", "fig." -> "Figure"   (e.g. "Fig. 3", "fig 3", "Fig.3" -> "Figure 3")
   - "Figs", "Figs.", "figs." -> "Figures"      (e.g. "Figs. 2–4" -> "Figures 2–4")
   - "Tab", "Tab.", "tab.", "Tbl", "Tbl." -> "Table"   (e.g. "Tab. 2" -> "Table 2")
   - "Tabs", "Tabs.", "Tbls." -> "Tables"
   - "Img", "img", "img." -> "Image"            (e.g. "img 5" -> "Image 5")
   - "Eq", "Eq.", "Eqn", "Eqn.", "eq." -> "Equation"  (e.g. "Eq. (4)" -> "Equation (4)")
   - "Eqs", "Eqs.", "Eqns." -> "Equations"
   - "Sec", "Sec.", "Sect", "Sect." -> "Section"
   - "Ch", "Ch.", "Chap", "Chap." -> "Chapter"
   - "App.", "Appx", "Appx." -> "Appendix"
   - "Suppl.", "Supp." -> "Supplementary"
   - "Photo", "photo" referring to a numbered plate -> "Photo" (capitalized) and keep the number (e.g. "photo 2" -> "Photo 2")
2. CAPITALIZATION: When the expanded label is immediately followed by its number/identifier, capitalize it ("Figure 3", "Table 2", "Equation 4"), even mid-sentence — this is the element's name. At the start of a sentence always capitalize. A general mention with no number (e.g. "as shown in the figure above") stays lowercase and is left as the ordinary word.
3. SPACING: Ensure exactly one space between the full word and its identifier ("Fig.3" -> "Figure 3"); preserve ranges/lists and any surrounding parentheses ("(Fig. 1a)" -> "(Figure 1a)").
4. SCOPE / DO NOT TOUCH: Do NOT expand these inside reference/bibliography entries, inside URLs/DOIs, or where the letters are part of another word (e.g. "fight", "tablet", "imgur", "section" already spelled out, "figure" already spelled out). Do NOT change established discipline abbreviations that are not element labels (units, gene names, chemical symbols, "et al.", "vs.", "etc.", "i.e.", "e.g."). Never alter the numbers/identifiers themselves.""",
    },
    {
        "id": "title",
        "title": "In-House Title Rules",
        "summary": "Article title in title case (Upper and Lower); subtitle after a colon capitalized; no trailing period.",
        "body": """IN-HOUSE TITLE RULES (apply ONLY to the article title):
1. TITLE CASE (Upper and Lower): Capitalize the first and last word and all major words; keep short articles, coordinating conjunctions, and short prepositions ("a", "an", "the", "of", "and", "for", "to", "in", "with", "on") lowercase unless they are the first or last word. After a colon, capitalize the first word of the subtitle. Example: "assessment of knowledge: regarding migraine among patients with migraine" -> "Assessment of Knowledge: Regarding Migraine Among Patients with Migraine".
2. Do NOT add a trailing period, and do NOT change the wording or meaning of the title.""",
    },
    {
        "id": "keywords",
        "title": "In-House Keywords Rules",
        "summary": "Alphabetical order; only the first keyword capitalized, rest lower case except proper nouns; comma + space separated; five to ten keywords; no trailing period.",
        "body": """IN-HOUSE KEYWORDS RULES (apply to the "Keywords:" line):
1. ORDER: Sort the keywords/phrases in alphabetical order FIRST (case-insensitive, by the first word of each phrase). Apply the case rule below AFTER sorting.
2. CASE: Only the FIRST keyword in the (now alphabetized) list has its first letter capitalized; every other keyword is lower case. The single EXCEPTION is proper nouns, which always keep their own capitalization wherever they appear (e.g. "Parkinson disease", "WhatsApp"). Example: "Keywords: Internet of Things, Artificial Intelligence, quality control" -> "Keywords: Artificial intelligence, internet of things, quality control".
3. SEPARATOR: Separate the keywords/phrases with a comma followed by a single space; no trailing period after the last keyword.
4. COUNT: There should be five to ten keywords or phrases. If there are fewer than five or more than ten, raise a query on the keywords paragraph — do NOT invent new keywords or delete meaningful ones to hit the count.""",
    },
    {
        "id": "abstract",
        "title": "In-House Abstract Rules",
        "summary": "Abstract heading in title case; no abbreviations, citations or references inside the abstract; ~300 words.",
        "body": """IN-HOUSE ABSTRACT RULES (apply to the abstract section):
1. HEADING: The "Abstract" heading is in Title Case.
2. CONTENT: The abstract text must NOT contain any abbreviations, in-text citations / reference numbers (e.g. "[1]"), or reference markers. Spell out any abbreviation in full and remove citation markers within the abstract only.
3. LENGTH: Aim for about 300 words. Do NOT pad the abstract or aggressively cut meaningful content just to hit the count; raise a query on the abstract paragraph if it is far outside the range.
(The journal sets the abstract body text in italic — that is its visual style, applied at layout.)""",
    },
    {
        "id": "front_matter",
        "title": "In-House Front-Matter Rules",
        "summary": "Author names + affiliations with superscript markers, corresponding author (*) + email, short title (≤60 chars), running short-author form.",
        "body": """IN-HOUSE FRONT-MATTER RULES (apply to the title block / author block):
1. AUTHOR NAMES: Give each author's name as initial(s) + full surname, each tagged with a superscript number linking to its affiliation. Example: "Saniya Jose1,*, Susan Kumar2, Thomas Mathew3".
2. AFFILIATION: Order each affiliation as Designation/Role, Department/School/Faculty, Institution/University/College, City, State/Province, Country — prefixed with its superscript number. Example: "2Professor, Department of Medical Surgical Nursing, St John's College of Nursing, Bangalore, Karnataka, India".
3. CORRESPONDING AUTHOR: Mark the corresponding author with "*"; provide a block "*Author for Correspondence" followed by the author's name and "E-mail: <address>".
4. SHORT TITLE: If the full title exceeds 70 characters, provide a short title of at most 60 characters (including spaces).
5. SHORT AUTHOR (running author line, non-italic): more than two authors -> "Lastname et al."; exactly two authors -> "Lastname and Lastname"; a single author -> the author's name.""",
    },
    {
        "id": "tables_figures",
        "title": "In-House Table & Figure Rules",
        "summary": "Arabic numbering; in-text \"Figure 2(b)\"/\"Table 1\"; captions sentence case with bold \"Figure 1.\"/\"Table 1.\" label; abbreviation lists alphabetical.",
        "body": """IN-HOUSE TABLE & FIGURE RULES:
1. IN-TEXT CITATION: Refer to tables/figures by the full word + an Arabic numeral, e.g. "Figure 2(b)", "(Figure 3a)", "Table 1" (never "Fig."/"Tab."). Numbering is Arabic and in sequential order.
2. CAPTION / LEGEND: Sentence case, non-italic, beginning with the label and number followed by a period — "Table 1." / "Figure 1." (label in bold at layout) — then the caption text in sentence case. Example: "Figure 1. Distribution of migraine patients on overall knowledge with regard to migraine.".
3. ABBREVIATIONS WITHIN A TABLE/FIGURE: List the abbreviation definitions in alphabetical order as "ABBR, expansion" separated by semicolons. Example: "CT, computed tomography; NMR, nuclear magnetic resonance; WHO, World Health Organization".""",
    },
    {
        "id": "text_abbreviations",
        "title": "In-House Text Abbreviation Rules",
        "summary": "Spell out each abbreviation in full at first use with the short form in parentheses, then use the short form thereafter.",
        "body": """IN-HOUSE TEXT ABBREVIATION RULES (general abbreviations in body text — distinct from figure/table element labels):
1. FIRST-USE EXPANSION: Spell out an abbreviation in full at its FIRST occurrence in the body text with the abbreviation in parentheses, then use the abbreviation throughout the rest of the text. Example: "... should be managed by multidisciplinary team (MDT). ... review by the MDT at least every 24 h.".
2. Do NOT re-expand the abbreviation on later occurrences, and do NOT apply this inside the abstract (the abstract carries no abbreviations). Units (e.g. "h", "min", "mg", "kg") and standard Latin abbreviations ("e.g.", "i.e.", "etc.", "et al.") are exempt and need no expansion.""",
    },
]

HOUSE_RULE_IDS = [g["id"] for g in HOUSE_RULE_GROUPS]


def build_house_rules_section(
    enabled_rule_ids: Optional[List[str]] = None, custom_rules: str = "",
) -> str:
    """Assemble the house-rules portion of the edit prompt from the enabled
    built-in groups plus any user-supplied custom rules. enabled_rule_ids=None
    means all built-in groups are active."""
    blocks: List[str] = []
    for group in HOUSE_RULE_GROUPS:
        if enabled_rule_ids is None or group["id"] in enabled_rule_ids:
            blocks.append(group["body"])
    if custom_rules and custom_rules.strip():
        blocks.append(
            "ADDITIONAL PUBLISHER RULES (user-provided; follow these exactly, "
            "they OVERRIDE the base style on any conflict):\n" + custom_rules.strip()
        )
    return "\n\n".join(blocks)


# --- Deterministic enforcement of the 6-author limit (safety net for the LLM) ---

# A single Vancouver / Index-Medicus author token, e.g. "Smith JA", "Lee C",
# "van der Berg AB", "O'Brien M", "Smith-Jones AB". The token must END in
# 1-4 capital initials so ordinary title/body words are not mistaken for authors.
_REF_AUTHOR_RE = re.compile(
    r"^[A-Z][A-Za-z'’.\-]*"               # first surname word (capitalized)
    r"(?:[ '’\-][A-Za-z][A-Za-z'’.\-]*)*"  # optional extra surname words / particles
    r" (?:[A-Z]\.?[ ]?){1,4}$"            # 1-4 initials (optional periods)
)
# An optional leading reference number such as "1.", "1)", "[1]", "(1)".
_REF_NUM_PREFIX_RE = re.compile(r"^\s*(?:\[\d+\]|\(\d+\)|\d+[.)])\s*")


def _trim_reference_authors(ref: str, max_authors: int = 6) -> str:
    """If `ref` is a reference entry whose author list exceeds `max_authors`,
    keep the first `max_authors` authors and append 'et al.'. Conservative: the
    text is returned unchanged unless the leading segment (before the title's
    terminating '. ') is unambiguously a comma-separated author list longer than
    the limit, with every token shaped like 'Surname Initials'."""
    prefix_m = _REF_NUM_PREFIX_RE.match(ref)
    prefix = prefix_m.group(0) if prefix_m else ""
    body = ref[len(prefix):]

    dot = body.find(". ")  # period+space that ends the author list before the title
    if dot == -1:
        return ref
    author_part = body[:dot]
    rest = body[dot:]  # begins with ". " — supplies the period after "et al"

    if re.search(r"\bet\s+al\b", author_part, re.IGNORECASE):
        return ref  # already abbreviated

    authors = [a.strip() for a in author_part.split(",")]
    if len(authors) <= max_authors:
        return ref
    if not all(_REF_AUTHOR_RE.match(a) for a in authors):
        return ref  # not confidently an author list -> leave untouched

    return prefix + ", ".join(authors[:max_authors]) + ", et al" + rest


def enforce_author_limit(
    paras: List[str], enabled_rule_ids: Optional[List[str]] = None,
    max_authors: int = 6,
) -> List[str]:
    """Deterministic safety net for the 6-author reference rule. Applied only
    when the reference rule group is active (enabled_rule_ids=None means all
    groups active). Runs after the LLM edit, so it operates on already
    normalized 'Surname Initials' author formatting."""
    if enabled_rule_ids is not None and "reference" not in enabled_rule_ids:
        return paras
    return [
        _trim_reference_authors(p, max_authors) if p and p.strip() else p
        for p in paras
    ]


# --- Deterministic enforcement of year-only reference dates (LLM safety net) ---

_MONTHS_RE = (
    r"(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|"
    r"Jul(?:y)?|Aug(?:ust)?|Sep(?:t)?(?:ember)?|Oct(?:ober)?|Nov(?:ember)?|"
    r"Dec(?:ember)?|Spring|Summer|Autumn|Fall|Winter)"
)
# A month (optionally preceded by a day) sitting immediately before a 4-digit
# year — i.e. a publication date like "July.2025", "June 2025", "15 Mar 2020".
# The year is captured so the whole run collapses to just the year.
_REF_MONTH_BEFORE_YEAR_RE = re.compile(
    r"(?:\b\d{1,2}(?:st|nd|rd|th)?\s+)?"
    r"\b" + _MONTHS_RE + r"\b\.?\s*((?:19|20)\d{2})\b",
    re.IGNORECASE,
)
# Quoted spans (straight or curly quotes) — protect a title's exact wording.
_QUOTED_SPAN_RE = re.compile("[\"“”][^\"“”]*[\"“”]")
# Entries whose month is intentional (event date / access date) and kept as-is.
_REF_KEEP_MONTH_RE = re.compile(
    r"accessed|conference|symposium|proceedings|congress|workshop|\bmeeting\b",
    re.IGNORECASE,
)


def _is_reference_entry(text: str) -> bool:
    """Heuristic: does this paragraph look like a bibliography entry (rather than
    body prose that merely mentions a year)? Conservative — we only strip months
    from text carrying a strong citation signal, so prose like 'in July 2025 we
    ran the trial' is never touched."""
    if not re.search(r"\b(19|20)\d{2}\b", text):
        return False
    if _REF_NUM_PREFIX_RE.match(text):
        return True
    t = text.lower()
    if '"' in text or "“" in text or "”" in text:   # quoted title
        return True
    if re.search(r"\bvol\.?\s*\d", t) or "doi" in t or re.search(r"\bet\s+al\b", t):
        return True
    if re.search(r"\bpp?\.\s*\d", t):   # page range, e.g. "pp. 2198–2201"
        return True
    return False


def _strip_reference_months(ref: str) -> str:
    """Collapse a publication 'Month Year' (or 'Day Month Year') to the year
    alone, but only OUTSIDE quoted titles so the title's wording is untouched."""
    def _drop(seg: str) -> str:
        return _REF_MONTH_BEFORE_YEAR_RE.sub(lambda mm: mm.group(1), seg)

    out: List[str] = []
    last = 0
    for m in _QUOTED_SPAN_RE.finditer(ref):
        out.append(_drop(ref[last:m.start()]))
        out.append(m.group(0))  # quoted title kept verbatim
        last = m.end()
    out.append(_drop(ref[last:]))
    return "".join(out)


def enforce_reference_year_only(
    paras: List[str], enabled_rule_ids: Optional[List[str]] = None,
) -> List[str]:
    """Deterministic safety net for the year-only reference-date rule. Applied
    only when the reference rule group is active (enabled_rule_ids=None means all
    groups active). Runs after the LLM edit. Skips conference and web
    'accessed on' entries — whose month is required — and never edits inside a
    quoted title."""
    if enabled_rule_ids is not None and "reference" not in enabled_rule_ids:
        return paras
    result: List[str] = []
    for p in paras:
        if (p and p.strip() and _is_reference_entry(p)
                and not _REF_KEEP_MONTH_RE.search(p)):
            result.append(_strip_reference_months(p))
        else:
            result.append(p)
    return result


# --- Deterministic removal of redundant parenthetical author-date citations ---
# Option 1: when the journal uses numbered markers "[n]", an "(Author, Year)"
# sitting in parentheses right next to the number duplicates it. A true
# parenthetical is by definition grammatically removable, and we only match
# parentheses that BEGIN with a capitalised author name and contain a 4-digit
# year — so NARRATIVE citations (author name in the running text, with only
# "(Year)" or nothing in the parentheses) never match and are preserved.

# An author-date enclosed in parentheses: opens with a capital-letter surname and
# ends at a 4-digit year (optional a/b disambiguation suffix). Bounded length and
# no nested parentheses keep it from over-reaching.
_PAREN_AUTHOR_DATE = r"\([A-Z][^()]{0,120}?(?:19|20)\d{2}[a-z]?\s*\)"
# A numeric in-text citation marker like [1], [1, 2], [3–7].
_NUM_CITE_MARKER = r"\[\s*\d+(?:\s*[–,\-]\s*\d+)*\s*\]"
# Parenthetical author-date immediately BEFORE the number — keep the number.
# Requires a preceding word char so we never strip a clause-opening subject.
_REDUNDANT_PAREN_BEFORE_RE = re.compile(
    r"(\w)\s+" + _PAREN_AUTHOR_DATE + r"\s*(" + _NUM_CITE_MARKER + r")"
)
# Parenthetical author-date immediately AFTER the number — keep the number.
_REDUNDANT_PAREN_AFTER_RE = re.compile(
    r"(" + _NUM_CITE_MARKER + r")\s*" + _PAREN_AUTHOR_DATE
)


def _drop_redundant_paren_citations(text: str) -> str:
    text = _REDUNDANT_PAREN_BEFORE_RE.sub(r"\1 \2", text)
    text = _REDUNDANT_PAREN_AFTER_RE.sub(r"\1", text)
    return text


def enforce_drop_redundant_paren_citation(
    paras: List[str], enabled_rule_ids: Optional[List[str]] = None,
) -> List[str]:
    """Deterministic safety net for Option 1: when a numbered marker "[n]" is
    present, drop an adjacent PARENTHETICAL "(Author, Year)" that duplicates it,
    keeping the number. Narrative citations (author name outside the parentheses)
    never match, so they are preserved. Applied only when the in-text citation
    rule group is active (enabled_rule_ids=None means all groups active)."""
    if enabled_rule_ids is not None and "citation" not in enabled_rule_ids:
        return paras
    return [
        _drop_redundant_paren_citations(p) if p and p.strip() else p
        for p in paras
    ]


# --- Deterministic enforcement of the keywords line (safety net for the LLM) ---

# Matches a leading "Keywords:" / "Key words -" label and captures (label, rest).
_KEYWORDS_LINE_RE = re.compile(r"^(\s*key\s*words?\s*[:\-–]\s*)(.+?)\s*$", re.I)


def _format_keywords_line(line: str) -> str:
    """Alphabetize the keywords and capitalize only the first one. Conservative:
    ordering, the first-keyword capital, comma+space separators and the no-
    trailing-period rule are enforced deterministically; the proper-noun-aware
    lower-casing of the remaining keywords is left to the LLM (code cannot tell
    'Internet' the common noun from 'Parkinson' the proper noun)."""
    m = _KEYWORDS_LINE_RE.match(line)
    if not m:
        return line
    label, rest = m.group(1), m.group(2)
    # Guard against a body sentence that merely starts with "Keyword: ..." — we
    # must not reorder/recapitalize prose. Real keyword items are short noun
    # phrases with no sentence structure. If anything looks sentence-like, bail
    # and defer to the LLM (under-applying is safe; corrupting a sentence is not).
    if re.search(r"[.!?]\s+\S", rest):  # mid-text sentence break
        return line
    items = [k.strip(" .") for k in re.split(r"[;,]", rest) if k.strip(" .")]
    if len(items) < 2:
        return line  # nothing to order; leave a single/empty keyword untouched
    _CONJ = {"and", "or", "but", "nor", "so", "yet", "because", "which", "that"}
    for it in items:
        words = it.split()
        if len(words) > 6:                 # keyword phrases are short
            return line
        if words and words[0].lower() in _CONJ:  # "..., and is robust" → prose
            return line
    items.sort(key=str.lower)
    first = items[0]
    head = first.split(" ", 1)[0]  # first token of the first keyword
    # Capitalize only a plainly-lowercase first word. Leave intentionally-cased
    # terms (mRNA, pH, iOS, eHealth, DNA, WHO) exactly as the LLM produced them —
    # the deterministic net must never override a correct edit.
    if head[:1].isalpha() and not any(c.isupper() for c in head):
        items[0] = first[0].upper() + first[1:]
    return label + ", ".join(items)


def enforce_keywords_format(
    paras: List[str], enabled_rule_ids: Optional[List[str]] = None,
) -> List[str]:
    """Deterministic safety net for the keywords rule. Applied only when the
    keywords rule group is active (enabled_rule_ids=None means all groups
    active). Runs after the LLM edit."""
    if enabled_rule_ids is not None and "keywords" not in enabled_rule_ids:
        return paras
    return [
        _format_keywords_line(p) if p and p.strip() else p
        for p in paras
    ]


def ai_edit_chunk(
    chunk_texts: List[str], settings: Dict[str, Any], edit_style: str, ref_style: str,
    lang: str, custom_dict: str, use_crossref: bool,
    enabled_rule_ids: Optional[List[str]] = None, custom_rules: str = "",
    use_serper: bool = False, serper_key: str = "",
) -> Tuple[List[str], List[Dict[str, Any]]]:
    """Returns (edited_texts, queries). `edited_texts` is the same length as
    `chunk_texts`. `queries` is a list of
    {"local_index", "snippet", "query", "suggestion"} for paragraphs the model
    flagged rather than silently changing ("suggestion" may be None)."""
    house_rules = build_house_rules_section(enabled_rule_ids, custom_rules)
    prompt = f"""You are a professional academic copyeditor.
Rules:
- Editing Style: {edit_style}
- Reference/Bibliography Style: {ref_style}
- Language Type: {lang}

I am providing a JSON array of text paragraphs from a manuscript.
Proofread them strictly following the specified guidelines. Fix spelling, grammar, and typography.
Properly format any references or bibliography items encountered.

EDIT CONSERVATIVELY — this is the most important rule:
- Make the MINIMUM changes necessary. Only alter text that is grammatically incorrect, misspelled, has a typographic error, or violates one of the rules below.
- PRESERVE the author's original wording, phrasing, sentence structure, and voice wherever the text is already correct. Do NOT rewrite, rephrase, reorder, or "improve" sentences for style, flow, or conciseness.
- If a sentence is awkward but grammatically correct and unambiguous, leave it unchanged.
- When a fix is needed, change only the specific words that are wrong; keep the rest of the sentence intact.
- NEVER delete a citation, an author name, or a sentence's grammatical subject. If a sentence begins with an author-date citation followed by a numeric marker (e.g. "Hasan et al. (2025) [1] ..."), KEEP BOTH the author-date AND the [n] marker exactly — they are not duplicates. Never remove the "Author et al. (Year)" part and never leave a sentence starting with a bare "[1] ..." that has no subject. To make a fragment like "Smith et al. (2020) [3] Offering an example." grammatical, change only the verb ("Offering" -> "offered"), keeping the author-date subject: "Smith et al. (2020) [3] offered an example."

{house_rules}
"""
    if custom_dict:
        prompt += f"\nCRITICAL CUSTOM DICTIONARY (DO NOT ALTER OR REFORMAT THESE TERMS):\n{custom_dict}\n"

    prompt += f"""
RAISING A QUERY INSTEAD OF GUESSING:
If something is genuinely ambiguous or you are NOT confident it is safe to change — e.g. a possible factual or numeric inconsistency, an unclear meaning, a missing/duplicated reference, or a house-rule issue you cannot resolve without the author's intent — do NOT silently rewrite it and do NOT guess. Leave the text as-is and raise a short QUERY for that paragraph (see the "query" field below). NEVER write the query inside the paragraph text itself.

CRITICAL OUTPUT FORMAT:
1. You MUST return ONLY a valid JSON array of the EXACT SAME LENGTH as the input array — no markdown, no extra text.
2. Each element corresponds to the input element at the SAME index and MUST be a JSON object with:
   - "edited": (string) the edited version of that paragraph. If nothing needs changing, return the original text unchanged.
   - "query": (string, OPTIONAL) a short note to the author/editor, used ONLY when you chose to flag rather than change. Omit the field or set it to null when there is no query.
   - "suggestion": (string, OPTIONAL) a concrete proposed correction the author can accept to resolve the query — wherever you can reasonably propose one. This is the SUGGESTED FIX, not a restatement of the problem: give the exact replacement wording, value, or reference the author would use (e.g. the corrected sentence, the right year, the missing citation). Only include it alongside a "query". Omit it (or set null) when the issue genuinely needs the author's own input and no fix can be proposed (e.g. "please confirm which value is intended"). Never apply the suggestion to the "edited" text yourself — it is a proposal for the author to accept.
3. Do not add markdown formatting or extra text. Return just the JSON array of objects.

Example output: [{{"edited": "The result was significant.", "query": null}}, {{"edited": "Smith (2020) reported the effect.", "query": "Reference [4] year reads 2020 here but 2002 in the bibliography — please confirm.", "suggestion": "Change the in-text year to 2002 to match the bibliography: \\"Smith (2002) reported the effect.\\""}}]

Input JSON:
{json.dumps(chunk_texts)}
"""

    try:
        text = _generate_text(prompt, settings=settings, response_mime_type="application/json")
        match = re.search(r"\[.*\]", text, re.DOTALL)
        if not match:
            print("Warning: No JSON array found in response.")
            return chunk_texts, []
        parsed = json.loads(match.group(0))

        if len(parsed) != len(chunk_texts):
            print("Warning: Array length mismatch.")
            return chunk_texts, []

        # Each element is normally an object {"edited", "query"}; tolerate a bare
        # string (older format) by treating it as the edited text with no query.
        result: List[str] = []
        queries: List[Dict[str, Any]] = []
        for i, elem in enumerate(parsed):
            if isinstance(elem, str):
                result.append(elem)
            elif isinstance(elem, dict):
                edited = elem.get("edited")
                result.append(edited if isinstance(edited, str) else chunk_texts[i])
                q = elem.get("query")
                if isinstance(q, str) and q.strip():
                    s = elem.get("suggestion")
                    queries.append({
                        "local_index": i,
                        "snippet": chunk_texts[i],
                        "query": q.strip(),
                        "suggestion": s.strip() if isinstance(s, str) and s.strip() else None,
                    })
            else:
                result.append(chunk_texts[i])

        if use_crossref or use_serper:
            for i in range(len(result)):
                t = result[i]
                if len(t) > 30 and re.search(r"\b(19|20)\d{2}\b", t) and "doi.org" not in t:
                    doi = fetch_crossref_doi(t) if use_crossref else None
                    if not doi and use_serper:
                        # Crossref came up empty (or is off) — try Google Scholar.
                        doi = fetch_serper_scholar_doi(t, serper_key)
                    if doi:
                        result[i] = t + f" {doi}"
        return result, queries
    except json.JSONDecodeError as e:
        print(f"JSON Parse Error: {e}")
        time.sleep(1)
        return chunk_texts, []
    except Exception as e:
        print(f"Gemini API Error: {e}")
        time.sleep(1)
        return chunk_texts, []


# --- Redline (.docx with native Word track changes) ---

def add_track_change_run(paragraph, text: str, change_type: str, tc_id: int,
                         author: str = "AI Editor", date: Optional[str] = None) -> None:
    if date is None:
        date = datetime.datetime.now().strftime("%Y-%m-%dT%H:%M:%SZ")

    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    r.append(rPr)

    if change_type == "insert":
        color = OxmlElement("w:color")
        color.set(qn("w:val"), "00B050")
        rPr.append(color)
        ins = OxmlElement("w:ins")
        ins.set(qn("w:id"), str(tc_id))
        ins.set(qn("w:author"), author)
        ins.set(qn("w:date"), date)
        t = OxmlElement("w:t")
        t.text = text
        t.set(qn("xml:space"), "preserve")
        r.append(t)
        ins.append(r)
        paragraph._p.append(ins)
    elif change_type == "delete":
        color = OxmlElement("w:color")
        color.set(qn("w:val"), "FF0000")
        rPr.append(color)
        strike = OxmlElement("w:strike")
        strike.set(qn("w:val"), "true")
        rPr.append(strike)
        delete = OxmlElement("w:del")
        delete.set(qn("w:id"), str(tc_id))
        delete.set(qn("w:author"), author)
        delete.set(qn("w:date"), date)
        delText = OxmlElement("w:delText")
        delText.text = text
        delText.set(qn("xml:space"), "preserve")
        r.append(delText)
        delete.append(r)
        paragraph._p.append(delete)
    else:
        t = OxmlElement("w:t")
        t.text = text
        t.set(qn("xml:space"), "preserve")
        r.append(t)
        paragraph._p.append(r)


#: The address of one paragraph inside a table: (table, row, cell, paragraph).
#: Four parts because all four are needed to be unambiguous — a cell can hold several
#: paragraphs, and a merged cell appears in more than one row.
TableAddress = Tuple[int, int, int, int]


#: Below this there is nothing to edit: "N", "pH", "Mean", a column letter.
#: Deliberately generous — the cost of skipping a two-word header is nil, and the cost
#: of sending 1,260 of them is 250 LLM calls and a real chance of one coming back
#: changed.
_MIN_CELL_CHARS = 15

#: Words of this length or more, made only of letters. Units and symbols are shorter
#: than this almost without exception, so counting them is how a measurement is told
#: from a sentence without listing every unit in existence.
_WORD = re.compile(r"[^\W\d_]{3,}", re.UNICODE)


def is_editable_cell(text: str) -> bool:
    """Whether a table cell holds prose worth sending to a copyeditor.

    Measured on a real manuscript before this existed: of 1,556 cell paragraphs, 71%
    were pure numbers and 81% were under fifteen characters. Sending them all took 311
    LLM calls and left the job on "Copyediting tables..." for a quarter of an hour, to
    edit text that has no grammar in it. Worse than the cost: a model asked to
    copyedit `0.15` may return `0.150`, and a silently altered number in a results
    table is the most damaging thing this tool could do to a manuscript.

    Decided by counting real words rather than by listing what a number looks like.
    A pattern of permitted characters has to be widened for every unit and every
    dash — the first version missed U+2212, which is the minus sign every results
    table actually uses, and called "−0.469 to 5.133" prose.
    """
    t = (text or "").strip()
    if len(t) < _MIN_CELL_CHARS:
        return False
    # Three real words. "4.5 N/mm2" has none, "Standard deviation" has two (a label),
    # "Table 1 Number of Ranks Given by Sample Respondents" has seven.
    return len(_WORD.findall(t)) >= 3


def collect_table_texts(structure) -> List[Tuple[TableAddress, str]]:
    """Cell paragraphs worth copyediting, with the address to write each one back to.

    This is the 11.3% of a real manuscript the copyeditor has never seen: `read_docx`
    returns `doc.paragraphs`, which skips table cells entirely. The addresses are what
    `generate_redline_docx` writes the edits back through, so the two are defined in
    the same file and tested together — invented separately is how the edits end up
    written nowhere.

    **A merged cell is returned once.** python-docx hands back the same cell for every
    grid position it spans, so a caption merged across seven columns arrives seven
    times; without this it is edited seven times, billed seven times, and — because
    each pass diffs against the same original — written back seven times over.
    """
    out: List[Tuple[TableAddress, str]] = []
    for t in structure.tables:
        seen: set = set()
        for row in t.grid:
            for cell in row:
                for pi, para in enumerate(cell.paragraphs):
                    if not is_editable_cell(para.text):
                        continue
                    # Keyed on the text and its position in the cell: a merged cell
                    # repeats both, a genuinely repeated value in two separate cells
                    # differs in neither — so this also collapses "Yes" in twenty
                    # rows, which is correct, since editing it once is editing it.
                    key = (pi, para.text)
                    if key in seen:
                        continue
                    seen.add(key)
                    out.append(((t.index, cell.row, cell.col, pi), para.text))
    return out


def _mark_up_paragraph(p, orig: str, edited: str, tc_id: int) -> int:
    """Rewrite one paragraph as track changes. Returns the next change id.

    Extracted so body paragraphs and table cells go through exactly the same code.
    A second copy of this diff loop for tables would be a second place for the
    author's text to be mangled, and only one of them would ever be looked at.
    """
    if orig.strip() == edited.strip():
        return tc_id

    p.clear()

    token_pattern = r"(\s+|\b|[.,!?;:])"
    orig_tokens = [t for t in re.split(token_pattern, orig) if t]
    edited_tokens = [t for t in re.split(token_pattern, edited) if t]

    matcher = difflib.SequenceMatcher(None, orig_tokens, edited_tokens)
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == "equal":
            add_track_change_run(p, "".join(orig_tokens[i1:i2]), "equal", tc_id)
        elif tag == "delete":
            add_track_change_run(p, "".join(orig_tokens[i1:i2]), "delete", tc_id)
            tc_id += 1
        elif tag == "insert":
            add_track_change_run(p, "".join(edited_tokens[j1:j2]), "insert", tc_id)
            tc_id += 1
        elif tag == "replace":
            add_track_change_run(p, "".join(orig_tokens[i1:i2]), "delete", tc_id)
            tc_id += 1
            add_track_change_run(p, "".join(edited_tokens[j1:j2]), "insert", tc_id)
            tc_id += 1
    return tc_id


def generate_redline_docx(
    original_path: str, edited_paragraphs: List[str], output_path: str,
    queries: Optional[List[Dict[str, Any]]] = None,
    table_edits: Optional[Dict[TableAddress, str]] = None,
) -> None:
    doc = docx.Document(original_path)
    tc_id = 1

    settings = doc.settings.element
    track_changes = OxmlElement("w:trackRevisions")
    settings.append(track_changes)

    for p, edited in zip(doc.paragraphs, edited_paragraphs):
        tc_id = _mark_up_paragraph(p, p.text, edited, tc_id)

    # Table cells, addressed rather than zipped. They are not in `doc.paragraphs`, so
    # the body list above keeps exactly the meaning it always had and an edit here
    # cannot shift it.
    for (ti, ri, ci, pi), edited in (table_edits or {}).items():
        try:
            cell = doc.tables[ti].rows[ri].cells[ci]
            para = cell.paragraphs[pi]
        except (IndexError, KeyError):
            # A stale address — the document changed under the job. Skipped rather
            # than raised: losing the whole redline over one cell is the worse
            # outcome, and writing into a neighbouring cell is worse still.
            continue
        tc_id = _mark_up_paragraph(para, para.text, edited, tc_id)

    # Anchor editor queries as native Word comments on their paragraphs. Done
    # after the edit loop so it covers both changed and unchanged paragraphs.
    if queries:
        query_map: Dict[int, List[str]] = {}
        for q in queries:
            idx = q.get("index")
            note = (q.get("query") or "").strip()
            if isinstance(idx, int) and note:
                suggestion = (q.get("suggestion") or "").strip()
                if suggestion:
                    note = f"{note}\nSuggested fix: {suggestion}"
                query_map.setdefault(idx, []).append(note)
        paragraphs = doc.paragraphs
        for idx, notes in query_map.items():
            if 0 <= idx < len(paragraphs):
                p = paragraphs[idx]
                anchor = p.add_run("")  # invisible anchor at paragraph end
                doc.add_comment(
                    anchor, text="\n".join(notes),
                    author="AI Editor", initials="AE",
                )

    doc.save(output_path)


# --- Journal recommendation ---

def cosine_similarity(v1, v2) -> float:
    n1 = np.linalg.norm(v1)
    n2 = np.linalg.norm(v2)
    if n1 == 0 or n2 == 0:
        return 0.0
    return float(np.dot(v1, v2) / (n1 * n2))


def _build_journal_embeddings(settings: Dict[str, Any]) -> str:
    in_path = app_config.journals_path()
    out_path = app_config.journals_embedded_path_for_settings(settings)

    with in_path.open("r") as f:
        journals = json.load(f)

    texts = [j["name"] + " " + " ".join(j.get("topics", [])) for j in journals]
    embeddings = []
    for i in range(0, len(texts), 100):
        batch = texts[i:i + 100]
        for t in batch:
            embeddings.append(_embed(t, settings=settings))
        time.sleep(0.5)
        print(f"  embedded {min(i + 100, len(texts))}/{len(texts)}")

    for j, emb in zip(journals, embeddings):
        j["embedding"] = emb

    out_path.parent.mkdir(parents=True, exist_ok=True)
    with out_path.open("w") as f:
        json.dump(journals, f)
    return str(out_path)


def _topic_keyword_hits(abstract: str, topics: List[str]) -> Dict[str, List[str]]:
    """Map each journal focus topic to the manuscript keywords that actually
    appear in the text — concrete, explainable evidence behind a match."""
    text = (abstract or "").lower()
    hits: Dict[str, List[str]] = {}
    for t in topics:
        phrase = (t or "").lower().strip()
        if not phrase:
            continue
        found: List[str] = []
        if phrase in text:
            found.append(phrase)
        for w in re.findall(r"[a-z]{4,}", phrase):
            if w in text and w not in found:
                found.append(w)
        if found:
            hits[t] = found
    return hits


def _matched_topics(abstract: str, topics: List[str]) -> List[str]:
    """Journal focus topics whose keywords actually appear in the manuscript
    text — concrete, explainable evidence behind a recommendation."""
    return list(_topic_keyword_hits(abstract, topics).keys())


def _score_tier(score: float) -> str:
    if score >= 0.75:
        return "a strong topical match"
    if score >= 0.5:
        return "a good topical match"
    if score >= 0.35:
        return "a moderate topical match"
    if score > 0:
        return "a loose topical match"
    return "a fallback candidate"


def _fit_label(score: float) -> str:
    """Short submission-guidance verdict derived from the similarity score."""
    if score >= 0.75:
        return "Primary target"
    if score >= 0.55:
        return "Strong fit"
    if score >= 0.4:
        return "Reasonable fit"
    if score > 0:
        return "Worth considering"
    return "Fallback candidate"


def _impact_context(impact_factor: Any) -> Optional[str]:
    """Interpret a journal's impact factor relative to this catalog's range
    (roughly 0.5–4.2), so the number carries meaning, not just a value."""
    if not impact_factor:
        return None
    try:
        v = float(impact_factor)
    except (TypeError, ValueError):
        return None
    if v >= 3.5:
        tier = "relatively high for this catalog, signalling strong visibility and selectivity"
    elif v >= 2.0:
        tier = "solid mid-range, a good balance of reach and acceptance odds"
    elif v >= 1.0:
        tier = "modest, which often means a faster, more accessible route to publication"
    else:
        tier = "on the lower end, typically the most accessible to submit to"
    return (
        f"Impact factor {impact_factor} is {tier}. It is shown for context only and "
        "does not influence the ranking, which is driven purely by topical fit."
    )


def _gap_phrase(score: float, next_score: Optional[float]) -> Optional[str]:
    """Describe how decisively this journal beats the next-ranked candidate."""
    if next_score is None or score <= 0 or next_score <= 0:
        return None
    gap = score - next_score
    if gap >= 0.1:
        return "It is clearly ahead of the next candidate, making it a confident first choice."
    if gap >= 0.03:
        return "It edges out the next candidate, but the alternatives below are close runners-up."
    return (
        "It only narrowly beats the next candidate — the journals below are nearly as good a "
        "fit, so weigh them on scope, speed, and impact factor too."
    )


def _explain_journal_match(abstract: str, journal: Dict[str, Any], rank: int = 1,
                           total: int = 0, next_score: Optional[float] = None) -> str:
    """Build a grounded, human-readable reason for why this journal was ranked,
    reflecting the actual logic (semantic similarity + topic overlap). Also
    records `matched_topics`, `matched_keywords`, `fit_label`, and `rank` on the
    journal dict so the UI and report can surface them directly."""
    score = journal.get("score", 0) or 0
    topics = journal.get("topics", [])
    hits = _topic_keyword_hits(abstract, topics)
    matched = list(hits.keys())
    keywords = sorted({w for ws in hits.values() for w in ws})
    journal["matched_topics"] = matched
    journal["matched_keywords"] = keywords
    # When the LLM ranking ran, the blended fit_score (0-100) is the headline
    # signal, so derive the label from it; otherwise fall back to the cosine score.
    fit_score = journal.get("fit_score")
    journal["fit_label"] = _fit_label(fit_score / 100 if fit_score is not None else score)
    journal["rank"] = rank

    parts: List[str] = []

    # 0. The AI advisor's holistic verdict, when the LLM ranking stage ran.
    llm_reason = journal.get("llm_reason")
    if llm_reason:
        verdict = journal.get("fit_verdict") or "a fit"
        confidence = journal.get("confidence")
        conf = f" ({confidence.lower()} confidence)" if confidence else ""
        parts.append(f"AI advisor verdict — {verdict}{conf}: {llm_reason}")

    # 1. Rank + verdict, with how many journals were considered.
    screened = f" of {total} journals screened" if total else ""
    if score > 0:
        parts.append(
            f"Ranked #{rank}{screened} and classified as a “{_fit_label(score)}” for your "
            "manuscript."
        )
    else:
        parts.append(
            f"Listed as a fallback candidate (#{rank}{screened}): semantic scoring was "
            "unavailable, so a provisional ranking was used."
        )

    # 2. The core signal: semantic similarity and what it means.
    if score > 0:
        parts.append(
            f"Vector embeddings put the semantic similarity between your manuscript and this "
            f"journal's scope (its title and focus topics) at {int(score * 100)}% — "
            f"{_score_tier(score)}."
        )

    # 3. Concrete evidence: which of the journal's topics/keywords actually appear.
    if matched:
        if keywords:
            kw = ", ".join(f"“{w}”" for w in keywords[:6])
            parts.append(
                "This is grounded in direct overlap with the journal's focus areas ("
                + ", ".join(t.title() for t in matched)
                + f"): your text explicitly uses the term(s) {kw}."
            )
        else:
            parts.append(
                "It overlaps directly with the journal's focus areas: "
                + ", ".join(t.title() for t in matched) + "."
            )
    elif topics and score > 0:
        parts.append(
            "There is no exact keyword overlap, so the match is semantic rather than literal: "
            "the journal's focus areas (" + ", ".join(t.title() for t in topics)
            + ") sit close to your manuscript's meaning in embedding space."
        )

    # 4. How decisive the ranking is versus the next option.
    gap = _gap_phrase(score, next_score)
    if gap:
        parts.append(gap)

    # 5. Impact factor, clearly labelled as context not ranking input.
    impact = _impact_context(journal.get("impact_factor"))
    if impact:
        parts.append(impact)

    # 6. Any cautions raised by the AI advisor or the scope gate.
    risks = journal.get("risk_factors") or []
    if risks:
        parts.append("Caution: " + " ".join(risks))

    return " ".join(parts)


# --- Hybrid recommendation: heuristic pre-filter + LLM ranking + validation ---
#
# The deterministic semantic score (cosine similarity) decides *who is eligible*;
# the LLM decides the *final ranking and reasoning* but only from that candidate
# list; a validation + scope-gate layer guarantees every pick is a real journal
# from our own catalogue. See JOURNAL_RECOMMENDATION_BLUEPRINT.md. The whole LLM
# stage degrades gracefully: if it is unavailable or fails, we fall back to the
# pure-semantic ranking, so nothing downstream breaks.

CANDIDATE_POOL_SIZE = 40
_NAME_STOP_WORDS = {
    "the", "of", "and", "for", "in", "on", "a", "an", "to", "journal",
    "journals", "international", "research",
}


def _normalize_journal_name(name: str) -> str:
    """Lowercase, strip punctuation/stop-words, collapse whitespace (blueprint §7)."""
    s = re.sub(r"[^a-z0-9\s]", " ", (name or "").lower())
    return " ".join(t for t in s.split() if t and t not in _NAME_STOP_WORDS)


def _names_match(a: str, b: str) -> bool:
    """Conservative fuzzy match used for anti-hallucination validation.

    Accept on: normalised equality, high-confidence containment (>0.92), or a
    Dice token-overlap > 0.9 with >= 2 shared tokens. Anything weaker is treated
    as a hallucinated / wrong name and discarded by the caller."""
    na, nb = _normalize_journal_name(a), _normalize_journal_name(b)
    if not na or not nb:
        return False
    if na == nb:
        return True
    if na in nb or nb in na:
        shorter, longer = sorted((len(na), len(nb)))
        if longer and shorter / longer > 0.92:
            return True
    ta, tb = set(na.split()), set(nb.split())
    if len(ta) >= 2 and len(tb) >= 2:
        inter = len(ta & tb)
        dice = 2 * inter / (len(ta) + len(tb))
        if dice > 0.9 and inter >= 2:
            return True
    return False


def _prescreen_score(journal: Dict[str, Any], abstract: str) -> float:
    """Cheap heuristic prior (0-100): semantic cosine (dominant) plus a bonus for
    literal overlap between the manuscript and the journal's focus topics."""
    cosine = max(0.0, journal.get("score", 0) or 0)
    topics = journal.get("topics", [])
    overlap = len(_topic_keyword_hits(abstract, topics)) / len(topics) if topics else 0
    return round(min(100.0, cosine * 80 + overlap * 20), 1)


def _llm_rank_journals(abstract: str, candidates: List[Dict[str, Any]],
                       settings: Dict[str, Any], k: int = 3) -> List[dict]:
    """Ask the LLM to pick the best k journals, constrained to the candidate
    list. Returns the parsed ``recommended_journals`` array, or [] on any failure
    (the caller then falls back to the heuristic ranking)."""
    cand_lines = "\n".join(f'{c["_cid"]} | {c["name"]}' for c in candidates)
    cand_details = json.dumps([
        {
            "id": c["_cid"],
            "name": c["name"],
            "topics": c.get("topics", []),
            "publisher": c.get("publisher"),
            "impact_factor": c.get("impact_factor"),
            "preScreenScore": c.get("prescreen_score", 0),
        }
        for c in candidates
    ])
    prompt = f"""You are an expert scholarly publishing advisor and journal-fit recommendation engine.
Recommend the {k} journals that best fit the manuscript, chosen ONLY from the supplied candidate list.

=== HARD CONSTRAINTS ===
1. ONLY recommend journals whose id and name appear EXACTLY in the Candidate Journals list. Copy the id (integer) and name verbatim. Any recommendation whose name is not found in that list will be discarded.
2. Do NOT invent, abbreviate, merge, or modify journal names or ids.
3. If fewer than {k} candidates are a reasonable fit, return fewer — never pad with poor matches.
4. Judge fit HOLISTICALLY, not by keyword overlap alone. Terminology can differ (e.g. "COVID-19" vs "pandemic infectious disease") — reward true topical/domain alignment even when exact words differ. Treat closely related fields as compatible.
=== END CONSTRAINTS ===

overall_fit_score (0-100) rubric: scope alignment 35%, domain/field match 25%, audience 15%, methodology 15%, contribution/article type 10%. Use preScreenScore as a prior, not a final answer.

MANUSCRIPT (title + abstract excerpt):
{abstract[:30000]}

CANDIDATE JOURNALS (id | name) — recommend only from these:
{cand_lines}

CANDIDATE DETAILS (evidence to judge fit; preScreenScore = heuristic prior 0-100):
{cand_details}

Return ONLY valid JSON (no markdown, no commentary) of this exact shape:
{{"recommended_journals":[{{"rank":1,"journal_id":0,"journal_name":"","overall_fit_score":0,"confidence":"High|Medium|Low","fit_verdict":"Excellent Fit|Good Fit|Possible Fit","why_this_journal":"concrete, evidence-based reasons grounded in the journal's own topics — not generic praise","fit_risks_or_cautions":[]}}]}}"""
    try:
        text = _generate_text(prompt, settings=settings, response_mime_type="application/json")
        match = re.search(r"\{.*\}", text, re.DOTALL)
        if not match:
            print("Journal LLM ranking: no JSON found in response.")
            return []
        data = json.loads(match.group(0))
        picks = data.get("recommended_journals", [])
        return picks if isinstance(picks, list) else []
    except Exception as e:
        print(f"Journal LLM ranking failed, falling back to semantic ranking: {e}")
        return []


def _validate_and_gate(picks: List[dict], candidates: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """Anti-hallucination validation (blueprint §7) + scope gate (§8).

    Each LLM pick must map to a real candidate (exact id+name, else conservative
    fuzzy name match), otherwise it is discarded. Survivors get a blended final
    ``fit_score``; picks the heuristic disagrees with are penalised (not deleted)
    and flagged."""
    by_id = {c["_cid"]: c for c in candidates}
    matched: List[Dict[str, Any]] = []
    seen: set = set()
    for p in picks:
        try:
            jid = int(p.get("journal_id"))
        except (TypeError, ValueError):
            jid = None
        name = p.get("journal_name", "")

        journal = None
        if jid in by_id and _names_match(name, by_id[jid]["name"]):
            journal = by_id[jid]
        else:
            for c in candidates:  # id was wrong/missing — trust a strong name match
                if _names_match(name, c["name"]):
                    journal = c
                    break
        if journal is None:
            print(f"Discarded hallucinated journal pick: {name!r} (id={p.get('journal_id')})")
            continue
        if journal["_cid"] in seen:
            continue
        seen.add(journal["_cid"])

        heuristic = journal.get("prescreen_score", 0)
        cosine = journal.get("score", 0) or 0
        llm_fit = float(p.get("overall_fit_score", 0) or 0)
        risks = list(p.get("fit_risks_or_cautions", []) or [])

        # Scope gate: a soft cross-check against our deterministic heuristic.
        if cosine < 0.2 and heuristic < 40:
            final = llm_fit * 0.8
            risks.append(
                "Heuristic scope check was weak — the AI advisor selected this on holistic "
                "fit rather than strong term/scope overlap, so confirm the scope yourself."
            )
        else:
            final = 0.45 * llm_fit + 0.55 * heuristic

        journal["fit_score"] = round(final, 1)
        journal["llm_fit_score"] = llm_fit
        journal["fit_verdict"] = p.get("fit_verdict", "")
        journal["confidence"] = p.get("confidence", "")
        journal["llm_reason"] = (p.get("why_this_journal", "") or "").strip()
        journal["risk_factors"] = risks
        matched.append(journal)
    return matched


def _finalize(top: List[Dict[str, Any]], abstract: str, total: int) -> List[Dict[str, Any]]:
    """Attach the deterministic, human-readable explanation to each pick."""
    for idx, j in enumerate(top):
        next_score = top[idx + 1].get("score") if idx + 1 < len(top) else None
        j["reason"] = _explain_journal_match(
            abstract, j, rank=idx + 1, total=total, next_score=next_score
        )
    return top


def recommend_journals(abstract: str, settings: Dict[str, Any], k: int = 3) -> List[dict]:
    try:
        abstract_emb = np.array(_embed(abstract, settings=settings))
    except Exception as e:
        print(f"Embedding error: {e}")
        abstract_emb = None

    journals_path = app_config.journals_embedded_path_for_settings(settings)
    if not journals_path.exists():
        try:
            print(f"Building journal embeddings at {journals_path} ...")
            _build_journal_embeddings(settings)
        except Exception as e:
            print(f"Warning: {journals_path} not found and rebuild failed: {e}")
            return []

    with journals_path.open("r") as f:
        journals = json.load(f)

    # Stage 1-2: score every journal and build the heuristic candidate pool.
    for idx, j in enumerate(journals):
        j["_cid"] = idx
        if abstract_emb is not None and "embedding" in j:
            j["score"] = cosine_similarity(abstract_emb, np.array(j["embedding"]))
        else:
            # No real embedding available — use 0.0 as the "semantic scoring
            # unavailable" sentinel (handled by _explain_journal_match as a
            # fallback candidate). NEVER use random.random() here: that value is
            # later rendered as a real similarity percentage and sorted on,
            # which would surface confidently-wrong recommendations. Ranking
            # then falls back to the deterministic prescreen_score (keyword
            # overlap), which is a legitimate signal.
            j["score"] = 0.0
        j["prescreen_score"] = _prescreen_score(j, abstract)

    total = len(journals)
    ranked = sorted(
        journals, key=lambda x: (x.get("prescreen_score", 0), x.get("score", 0)), reverse=True
    )
    candidates = ranked[:CANDIDATE_POOL_SIZE]

    # Stage 3-5: LLM ranking constrained to candidates, then validate + scope-gate.
    picks = _llm_rank_journals(abstract, candidates, settings, k=k) if candidates else []
    gated = _validate_and_gate(picks, candidates)

    # Stage 6: LLM picks first (sorted by blended fit), backfill from the
    # heuristic pool, cap to k. If the LLM stage produced nothing, this is just
    # the pure-semantic top-k — identical to the previous behaviour.
    gated.sort(key=lambda j: j.get("fit_score", 0), reverse=True)
    chosen = {j["_cid"] for j in gated}
    for c in ranked:
        if len(gated) >= k:
            break
        if c["_cid"] not in chosen:
            gated.append(c)
            chosen.add(c["_cid"])
    top = gated[:k]

    return _finalize(top, abstract, total)


# --- Report / cover letter / polish ---

def generate_report(edit_style: str, ref_style: str, lang: str,
                    used_crossref: bool, custom_dict: str,
                    enabled_rule_ids: Optional[List[str]] = None,
                    custom_rules: str = "",
                    queries: Optional[List[Dict[str, Any]]] = None,
                    warnings: Optional[List[str]] = None,
                    plagiarism: Optional[Dict[str, Any]] = None) -> str:
    report = "### 📑 Editorial Report\n\n"
    if warnings:
        report += "**⚠️ Notices:**\n"
        for w in warnings:
            report += f"- {w}\n"
        report += "\n"
    report += "**Configurations Applied:**\n"
    report += f"- **Copyediting:** {edit_style}\n"
    report += f"- **References:** {ref_style}\n"
    report += f"- **Language:** {lang}\n"
    if custom_dict:
        report += "- **Custom Dictionary:** Active constraints applied.\n"
    report += "\n**Summary of Interventions:**\n"
    report += "- Corrected typography and narrative consistency.\n"
    if used_crossref:
        report += "- **Live Crossref Validation:** Officially verified DOIs embedded into bibliography.\n"
    else:
        report += "- **References Formatted:** AI applied selected style structure.\n"

    report += "\n**Publisher / House Rules Applied:**\n"
    for group in HOUSE_RULE_GROUPS:
        applied = enabled_rule_ids is None or group["id"] in enabled_rule_ids
        mark = "✅ Applied" if applied else "⬜ Skipped"
        report += f"- **{group['title']}:** {mark} — {group['summary']}\n"
    if custom_rules and custom_rules.strip():
        report += "\n**Additional Publisher Rules (custom):**\n"
        for line in custom_rules.strip().splitlines():
            line = line.strip()
            if line:
                report += f"- {line}\n"

    if queries:
        report += "\n**Editor Queries (please review):**\n"
        report += (
            f"_The editor flagged {len(queries)} item(s) rather than changing them "
            "silently. Where a fix could be proposed, a **Suggested fix** is given "
            "for you to accept. Each query also appears as a comment in the redline "
            "document._\n"
        )
        for q in queries:
            snippet = (q.get("snippet") or "").strip()
            snippet = re.sub(r"\s+", " ", snippet)
            if len(snippet) > 80:
                snippet = snippet[:77] + "…"
            loc = f"¶{q['index'] + 1}" if isinstance(q.get("index"), int) else "—"
            prefix = f"**{snippet}** ({loc})" if snippet else f"({loc})"
            report += f"- {prefix}: {q['query']}\n"
            suggestion = (q.get("suggestion") or "").strip()
            if suggestion:
                report += f"    - 💡 **Suggested fix:** {suggestion}\n"

    if plagiarism and plagiarism.get("available"):
        flagged = plagiarism.get("flagged") or []
        checked = plagiarism.get("checked", 0)
        score = plagiarism.get("score")
        report += "\n**🔎 Preliminary Originality Scan:**\n"
        report += f"_{plagiarism.get('disclaimer', '')}_\n"
        if score is not None:
            report += (
                f"- Originality (preliminary): **{score}%** — {len(flagged)} of "
                f"{checked} checked sentence(s) had a verbatim web match.\n"
            )
        for item in flagged:
            s = re.sub(r"\s+", " ", item.get("sentence", "")).strip()
            if len(s) > 100:
                s = s[:97] + "…"
            report += f"- Possible match: “{s}”\n"
            for src in item.get("sources", [])[:3]:
                title = (src.get("title") or src.get("link") or "source").strip()
                report += f"    - [{title}]({src.get('link', '')})\n"
        if not flagged:
            report += "- No verbatim web matches found in the checked sample.\n"
    elif plagiarism and plagiarism.get("note"):
        report += f"\n**🔎 Preliminary Originality Scan:** _{plagiarism['note']}_\n"

    return report


def build_plagiarism_report(plagiarism: Optional[Dict[str, Any]]) -> str:
    """Render the preliminary originality scan as a standalone, downloadable
    Markdown report (full detail: every flagged sentence and all its sources).
    Returns '' when there is nothing to report so the caller can skip writing a
    file."""
    if not plagiarism or not plagiarism.get("available"):
        return ""
    flagged = plagiarism.get("flagged") or []
    checked = plagiarism.get("checked", 0)
    score = plagiarism.get("score")

    lines = ["# 🔎 Preliminary Originality Report", ""]
    lines.append(f"> **{plagiarism.get('disclaimer', '')}**")
    lines.append("")
    lines.append("## Summary")
    if score is not None:
        lines.append(f"- **Preliminary originality score:** {score}%")
    lines.append(f"- **Sentences checked:** {checked}")
    lines.append(f"- **Sentences with a verbatim web match:** {len(flagged)}")
    lines.append("")
    lines.append(
        "_The score is the share of checked sentences with NO verbatim match on "
        "the public web. A high score is not a guarantee of originality, and a "
        "flagged sentence is not proof of plagiarism — common phrasing, correct "
        "quotation, and an author's own prior posts all match too. Review each "
        "item in context._"
    )
    lines.append("")
    lines.append("## Flagged sentences")
    if not flagged:
        lines.append("")
        lines.append("_No verbatim web matches were found in the checked sample._")
        return "\n".join(lines) + "\n"

    for i, item in enumerate(flagged, 1):
        sentence = re.sub(r"\s+", " ", item.get("sentence", "")).strip()
        lines.append("")
        lines.append(f"### {i}. “{sentence}”")
        sources = item.get("sources") or []
        if sources:
            lines.append("Matching web sources:")
            for src in sources:
                title = (src.get("title") or src.get("link") or "source").strip()
                link = src.get("link", "")
                lines.append(f"- [{title}]({link})")
        else:
            lines.append("- _(match detected; no source link available)_")
    return "\n".join(lines) + "\n"


def build_journal_report(recommended: List[dict]) -> str:
    """Render the top journal recommendations as a downloadable Markdown report."""
    lines = ["# 📚 Top Journal Recommendations", ""]
    lines.append(
        "_How these are chosen: a deterministic semantic score (vector-embedding "
        "similarity between your manuscript and each journal's scope) selects the "
        "candidate journals; an AI publishing advisor then ranks the best fits from "
        "that shortlist, and every pick is validated against our real journal "
        "catalogue. A higher fit score means a closer match. Impact factor is shown "
        "for context only and does not affect the ranking._"
    )
    lines.append("")
    if not recommended:
        lines.append("_No journal recommendations were generated for this manuscript._")
        return "\n".join(lines) + "\n"
    for i, j in enumerate(recommended, 1):
        fit_score = j.get("fit_score")
        score = j.get("score", 0)
        if fit_score is not None:
            match = f"{int(round(fit_score))}/100 fit"
        elif score > 0:
            match = f"{int(score * 100)}% match"
        else:
            match = "Recommended"
        fit = j.get("fit_label")
        verdict = j.get("fit_verdict")
        heading = f"## {i}. {j.get('name', 'Unknown')}"
        if fit:
            heading += f" — {fit}"
        lines.append(heading)
        lines.append(f"- **Relevance:** {match}")
        if verdict:
            conf = j.get("confidence")
            lines.append(f"- **AI Advisor Verdict:** {verdict}" + (f" ({conf} confidence)" if conf else ""))
        if j.get("impact_factor"):
            lines.append(f"- **Impact Factor:** {j.get('impact_factor')}")
        lines.append(f"- **Publisher:** {j.get('publisher', 'Unknown')}")
        topics = j.get("topics", [])
        if topics:
            lines.append(f"- **Focus Topics:** {', '.join(topics).title()}")
        matched = j.get("matched_topics", [])
        if matched:
            lines.append(f"- **Matched Topics:** {', '.join(t.title() for t in matched)}")
        keywords = j.get("matched_keywords", [])
        if keywords:
            lines.append(f"- **Matched Terms:** {', '.join(keywords)}")
        reason = j.get("reason", "")
        if reason:
            lines.append(f"- **Why recommended:** {reason}")
        lines.append("")
    return "\n".join(lines) + "\n"


def _add_markdown_runs(paragraph, text: str) -> None:
    """Add runs to a paragraph, rendering **bold** spans in the limited markdown."""
    for i, segment in enumerate(re.split(r"\*\*(.+?)\*\*", text)):
        if not segment:
            continue
        run = paragraph.add_run(segment)
        run.bold = i % 2 == 1  # odd segments were captured inside ** **


def markdown_to_docx(md_text: str, out_path: str) -> str:
    """Render the limited markdown the platform produces (# / ## headings,
    - bullets, **bold**) into a Word .docx file. Returns out_path."""
    document = docx.Document()
    for raw_line in md_text.splitlines():
        line = raw_line.rstrip()
        if not line.strip():
            continue
        if line.startswith("# "):
            document.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            document.add_heading(line[3:].strip(), level=2)
        elif line.startswith("### "):
            document.add_heading(line[4:].strip(), level=3)
        elif line.lstrip().startswith("- "):
            para = document.add_paragraph(style="List Bullet")
            _add_markdown_runs(para, line.lstrip()[2:].strip())
        else:
            para = document.add_paragraph()
            # strip surrounding underscores used for italic placeholders
            _add_markdown_runs(para, line.strip().strip("_"))
    document.save(out_path)
    return out_path


_DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


# --- JATS / XML export (publisher production format) ---

# Section names recognized as structural headings when building JATS.
_SECTION_KEYWORDS = {
    "abstract", "keywords", "introduction", "background", "related work",
    "literature review", "materials and methods", "methods", "methodology",
    "experimental", "results", "results and discussion", "discussion",
    "conclusion", "conclusions", "summary", "acknowledgements",
    "acknowledgments", "references", "bibliography", "reference list",
}
_REFERENCE_HEADINGS = {"references", "bibliography", "reference list"}
JATS_DOCTYPE = (
    '<!DOCTYPE article PUBLIC '
    '"-//NLM//DTD JATS (Z39.96) Journal Publishing DTD v1.3 20210610//EN" '
    '"JATS-journalpublishing1-3.dtd">'
)
JATS_MIME = "application/xml"
_XLINK_NS = "http://www.w3.org/1999/xlink"

# Figure / table caption patterns, e.g. "Figure 1. ...", "Fig. 2: ...", "Table 3 - ...".
_FIG_RE = re.compile(r"^(Fig(?:ure)?\.?\s*([0-9]+|[IVXLC]+))\s*[.:)\-]\s*(.+)$", re.I)
_TAB_RE = re.compile(r"^(Tab(?:le)?\.?\s*([0-9]+|[IVXLC]+))\s*[.:)\-]\s*(.+)$", re.I)


def _normalize_heading(text: str) -> str:
    return re.sub(r"\s+", " ", (text or "").strip().lower()).rstrip(":.")


def _is_jats_heading(text: str) -> bool:
    """Heuristically decide whether a paragraph is a section heading. Conservative:
    known section names, numbered headings (e.g. '2.1 Methods'), or a short
    capitalized line with no terminal sentence punctuation."""
    t = (text or "").strip()
    if not t:
        return False
    if _FIG_RE.match(t) or _TAB_RE.match(t):
        return False  # figure/table captions are blocks, not headings
    if _normalize_heading(t) in _SECTION_KEYWORDS:
        return True
    if re.match(r"^\d+(\.\d+)*\.?\s+\S", t) and len(t) <= 80 and not t.endswith("."):
        return True
    if len(t) <= 70 and len(t.split()) <= 8 and not re.search(r"[.!?]$", t):
        if t.isupper() or t[:1].isupper():
            return True
    return False


_DOI_RE = re.compile(r"(10\.\d{4,9}/[^\s,;]+)", re.I)
_YEAR_VOL_RE = re.compile(
    r"(\d{4})\s*;\s*(\d+)\s*(?:\(\s*([^)]+?)\s*\))?\s*:\s*"
    r"([A-Za-z]?\d+(?:\s*[-–]\s*[A-Za-z]?\d+)?)"
)
# In-text numeric citation group, e.g. "[1]", "[1, 2, 3]", "[3-7]", "[3–7]".
_CITE_GROUP_RE = re.compile(r"\[(\d+(?:\s*[-–,]\s*\d+)*)\]")
_NUM_PREFIX_RE = re.compile(r"^(\d+(?:\.\d+)*)\.?\s+(.*)$")
_AFFIL_HINTS = (
    "university", "department", "institute", "college", "laborator",
    "hospital", "school of", "faculty", "centre", "center", "dept", "academy",
)


def _split_authors(authors_str: str) -> List[str]:
    return [a.strip() for a in (authors_str or "").split(",") if a.strip()]


def _ref_name_elem(parent, token: str) -> None:
    """Vancouver-style 'Surname Initials' -> <name><surname>/<given-names>."""
    token = token.strip()
    if re.fullmatch(r"et\s+al\.?", token, re.I):
        ET.SubElement(parent, "etal")
        return
    m = re.match(r"^(.+?)\s+([A-Z][A-Za-z.\-]{0,5})$", token)
    name = ET.SubElement(parent, "name")
    if m:
        ET.SubElement(name, "surname").text = m.group(1)
        ET.SubElement(name, "given-names").text = m.group(2)
    else:
        ET.SubElement(name, "surname").text = token


def _contrib_name_elem(parent, token: str) -> None:
    """Byline-style 'First M. Last' -> surname = last word, given-names = rest."""
    words = token.split()
    name = ET.SubElement(parent, "name")
    if len(words) >= 2:
        ET.SubElement(name, "surname").text = words[-1]
        ET.SubElement(name, "given-names").text = " ".join(words[:-1])
    else:
        ET.SubElement(name, "surname").text = token


def _author_tokens(line: str) -> List[str]:
    cleaned = re.sub(r"\S+@\S+", "", line)                      # emails
    cleaned = re.sub(r"[0-9\*†‡§¶#]+", "", cleaned)  # affil markers
    cleaned = re.sub(r"\s+and\s+", ",", cleaned, flags=re.I)
    return [t.strip(" .") for t in cleaned.split(",") if t.strip(" .")]


def _looks_like_author_line(line: str) -> bool:
    """Conservative byline detection: a comma/and-separated list of at least two
    capitalized name-like tokens (so a 2-word heading is never mistaken for an
    author line). Single-author bylines are intentionally not auto-detected."""
    t = (line or "").strip()
    if not t or _normalize_heading(t) in _SECTION_KEYWORDS:
        return False
    if len(t) > 200 or re.search(r"[.!?]$", t):
        return False
    tokens = _author_tokens(t)
    if not (2 <= len(tokens) <= 20):
        return False
    good = 0
    for tok in tokens:
        words = tok.split()
        if 1 <= len(words) <= 5 and all(w[:1].isupper() for w in words if w):
            good += 1
    return good >= 2 and good >= len(tokens) - 1


def _looks_like_affiliation(line: str) -> bool:
    low = (line or "").lower()
    return any(h in low for h in _AFFIL_HINTS)


def _looks_like_single_author(line: str) -> bool:
    """A single byline name (no comma). Only trusted when the following line is
    an affiliation, to avoid mistaking a 2-3 word heading for an author."""
    t = (line or "").strip()
    if not t or "," in t or _normalize_heading(t) in _SECTION_KEYWORDS:
        return False
    if re.search(r"[.!?]$", t):
        return False
    words = t.split()
    return 2 <= len(words) <= 5 and all(w[:1].isupper() for w in words)


def _extract_keywords(items: List[str]):
    """Pull an inline 'Keywords: a, b, c' line out of the body items."""
    kept: List[str] = []
    keywords: Optional[List[str]] = None
    for it in items:
        m = re.match(r"^\s*key\s*words?\s*[:\-–]\s*(.+)$", it, re.I)
        if m and keywords is None:
            keywords = [k.strip(" .") for k in re.split(r"[;,]", m.group(1)) if k.strip(" .")]
        else:
            kept.append(it)
    return keywords, kept


def _split_dot(text: str) -> List[str]:
    return [p.strip().rstrip(".").strip() for p in re.split(r"\.\s+", text) if p.strip()]


def _parse_journal_reference(raw: str, doi: Optional[str]):
    yvp = _YEAR_VOL_RE.search(raw)
    if not yvp:
        return None
    parts = _split_dot(raw[:yvp.start()].strip())
    authors = title = source = None
    if len(parts) >= 3:
        authors, source = parts[0], parts[-1]
        title = " ".join(parts[1:-1])
    elif len(parts) == 2:
        authors, title = parts
    elif len(parts) == 1:
        authors = parts[0]
    return {
        "kind": "journal", "raw": raw, "authors": authors, "title": title,
        "source": source, "year": yvp.group(1), "volume": yvp.group(2),
        "issue": yvp.group(3), "pages": yvp.group(4), "doi": doi,
    }


def _parse_book_chapter(raw: str, doi: Optional[str]):
    if not re.search(r"\bIn:\s", raw, re.I):
        return None
    before, after = re.split(r"\bIn:\s*", raw, maxsplit=1, flags=re.I)
    bparts = _split_dot(before)
    authors = bparts[0] if bparts else None
    chapter_title = " ".join(bparts[1:]) if len(bparts) > 1 else None

    editors = None
    rest_after = after
    em = re.match(r"^(.*?),?\s*editors?\.\s*(.*)$", after, re.I)
    if em:
        editors, rest_after = em.group(1).strip(), em.group(2).strip()
    aparts = _split_dot(rest_after)
    source = aparts[0] if aparts else None
    place = publisher = year = pages = None
    for seg in aparts[1:]:
        pm = re.search(r"(?:([^:;]+):\s*)?([^;]+);\s*(\d{4})", seg)
        if pm:
            place = pm.group(1).strip() if pm.group(1) else None
            publisher = pm.group(2).strip()
            year = pm.group(3)
    # Extract pages from the whole tail (the "p. 33-58" token can be split by
    # the sentence splitter, so search the raw string instead).
    pg = re.search(r"\bp\.?\s*([0-9]+(?:\s*[-–]\s*[0-9]+)?)", rest_after)
    if pg:
        pages = pg.group(1)
    if not year:
        ym = re.search(r"(\d{4})", after)
        year = ym.group(1) if ym else None
    return {
        "kind": "book-chapter", "raw": raw, "authors": authors,
        "chapter_title": chapter_title, "editors": editors, "source": source,
        "place": place, "publisher": publisher, "year": year, "pages": pages, "doi": doi,
    }


def _parse_book(raw: str, doi: Optional[str]):
    ym = re.search(r";\s*(\d{4})", raw) or re.search(r"(\d{4})\.?\s*$", raw)
    if not ym:
        return None
    year = ym.group(1)
    parts = _split_dot(raw)
    if len(parts) < 2:
        return None
    authors, title = parts[0], parts[1]
    edition = publisher = place = None
    for seg in parts[2:]:
        if re.match(r"^\d+(?:st|nd|rd|th)\s+ed\.?$", seg, re.I):
            edition = seg
            continue
        pm = re.search(r"(?:([^:;]+):\s*)?([^;]+);\s*\d{4}", seg)
        if pm:
            place = pm.group(1).strip() if pm.group(1) else None
            publisher = pm.group(2).strip()
    return {
        "kind": "book", "raw": raw, "authors": authors, "source": title,
        "edition": edition, "publisher": publisher, "place": place,
        "year": year, "doi": doi,
    }


def _parse_reference(text: str) -> Dict[str, Any]:
    """Parse a normalized Vancouver-style reference into structured fields,
    detecting journal articles, book chapters, and books. Falls back to
    {'kind': 'other'} when unrecognized, so the caller emits a <mixed-citation>."""
    raw = text.strip()
    doi_m = _DOI_RE.search(raw)
    doi = doi_m.group(1).rstrip(".") if doi_m else None
    return (
        _parse_journal_reference(raw, doi)
        or _parse_book_chapter(raw, doi)
        or _parse_book(raw, doi)
        or {"kind": "other", "raw": raw, "doi": doi}
    )


def _add_person_group(parent, group_type: str, authors_str: str) -> None:
    pg = ET.SubElement(parent, "person-group", {"person-group-type": group_type})
    for tok in _split_authors(authors_str):
        _ref_name_elem(pg, tok)


def _add_pages(ec, pages: str) -> None:
    parts = re.split(r"\s*[-–]\s*", pages)
    ET.SubElement(ec, "fpage").text = parts[0]
    if len(parts) > 1:
        ET.SubElement(ec, "lpage").text = parts[1]


def _add_reference(ref_list_el, n: int, text: str) -> None:
    info = _parse_reference(text)
    kind = info["kind"]
    ref = ET.SubElement(ref_list_el, "ref", {"id": f"ref{n}"})

    if kind == "other":
        mc = ET.SubElement(ref, "mixed-citation")
        mc.text = info["raw"]
        if info.get("doi"):
            ET.SubElement(mc, "pub-id", {"pub-id-type": "doi"}).text = info["doi"]
        return

    ec = ET.SubElement(ref, "element-citation", {"publication-type": kind})
    if info.get("authors"):
        _add_person_group(ec, "author", info["authors"])

    if kind == "journal":
        if info.get("title"):
            ET.SubElement(ec, "article-title").text = info["title"]
        if info.get("source"):
            ET.SubElement(ec, "source").text = info["source"]
    elif kind == "book-chapter":
        if info.get("chapter_title"):
            ET.SubElement(ec, "chapter-title").text = info["chapter_title"]
        if info.get("editors"):
            _add_person_group(ec, "editor", info["editors"])
        if info.get("source"):
            ET.SubElement(ec, "source").text = info["source"]
    elif kind == "book":
        if info.get("source"):
            ET.SubElement(ec, "source").text = info["source"]
        if info.get("edition"):
            ET.SubElement(ec, "edition").text = info["edition"]

    if info.get("place"):
        ET.SubElement(ec, "publisher-loc").text = info["place"]
    if info.get("publisher"):
        ET.SubElement(ec, "publisher-name").text = info["publisher"]
    if info.get("year"):
        ET.SubElement(ec, "year").text = info["year"]
    if info.get("volume"):
        ET.SubElement(ec, "volume").text = info["volume"]
    if info.get("issue"):
        ET.SubElement(ec, "issue").text = info["issue"]
    if info.get("pages"):
        _add_pages(ec, info["pages"])
    if info.get("doi"):
        ET.SubElement(ec, "pub-id", {"pub-id-type": "doi"}).text = info["doi"]


class _MixedContent:
    """Helper to append interleaved text and child elements to a parent,
    using ElementTree's text/tail model."""

    def __init__(self, parent):
        self.parent = parent
        self.last = None

    def text(self, s: str) -> None:
        if not s:
            return
        if self.last is None:
            self.parent.text = (self.parent.text or "") + s
        else:
            self.last.tail = (self.last.tail or "") + s

    def xref(self, n: int, label: str) -> None:
        x = ET.SubElement(self.parent, "xref", {"ref-type": "bibr", "rid": f"ref{n}"})
        x.text = label
        self.last = x


def _add_body_paragraph(parent, text: str, valid_refs) -> None:
    """Add a <p>, linking numeric in-text citations [n] to <xref rid='refn'>."""
    p = ET.SubElement(parent, "p")
    if not valid_refs:
        p.text = text
        return
    mc = _MixedContent(p)
    pos = 0
    for m in _CITE_GROUP_RE.finditer(text):
        mc.text(text[pos:m.start()])
        mc.text("[")
        for tok in re.findall(r"\d+|[^\d]+", m.group(1)):
            if tok.isdigit() and int(tok) in valid_refs:
                mc.xref(int(tok), tok)
            else:
                mc.text(tok)
        mc.text("]")
        pos = m.end()
    mc.text(text[pos:])


def _add_block(parent, text: str, valid_refs) -> None:
    """Emit a body block: a figure caption -> <fig>, a table caption ->
    <table-wrap>, otherwise a <p> (with in-text citation linking)."""
    fm = _FIG_RE.match(text)
    if fm:
        label = re.sub(r"\s+", " ", fm.group(1)).strip()
        fig = ET.SubElement(parent, "fig", {"id": f"fig{fm.group(2)}"})
        ET.SubElement(fig, "label").text = label
        caption = ET.SubElement(fig, "caption")
        _add_body_paragraph(caption, fm.group(3).strip(), valid_refs)
        return
    tm = _TAB_RE.match(text)
    if tm:
        label = re.sub(r"\s+", " ", tm.group(1)).strip()
        tw = ET.SubElement(parent, "table-wrap", {"id": f"tab{tm.group(2)}"})
        ET.SubElement(tw, "label").text = label
        caption = ET.SubElement(tw, "caption")
        _add_body_paragraph(caption, tm.group(3).strip(), valid_refs)
        return
    _add_body_paragraph(parent, text, valid_refs)


def _emit_sections(body_el, sections, valid_refs) -> None:
    """Emit sections, nesting numbered subsections (e.g. 2.1 under 2) and
    splitting a leading number into <label>. Untitled leading paragraphs go
    directly under <body>."""
    stack = []  # list of (depth, element)
    for s in sections:
        title = s["title"]
        if not title:
            for p in s["paras"]:
                _add_block(body_el, p, valid_refs)
            continue
        label = None
        title_text = title
        depth = 1
        m = _NUM_PREFIX_RE.match(title)
        if m:
            label, title_text, depth = m.group(1), m.group(2), m.group(1).count(".") + 1
        while stack and stack[-1][0] >= depth:
            stack.pop()
        parent = stack[-1][1] if stack else body_el
        sec = ET.SubElement(parent, "sec")
        if label:
            ET.SubElement(sec, "label").text = label
        ET.SubElement(sec, "title").text = title_text
        for p in s["paras"]:
            _add_block(sec, p, valid_refs)
        stack.append((depth, sec))


def _add_structured_contrib(contrib_group, author: Dict[str, Any]) -> None:
    """Build a <contrib> from a structured author dict supporting
    surname/given_names (or full name), ORCID, corresponding flag, and email."""
    attrs = {"contrib-type": "author"}
    if author.get("corresponding"):
        attrs["corresp"] = "yes"
    contrib = ET.SubElement(contrib_group, "contrib", attrs)
    orcid = author.get("orcid")
    if orcid:
        if not orcid.startswith("http"):
            orcid = "https://orcid.org/" + orcid
        ET.SubElement(contrib, "contrib-id", {"contrib-id-type": "orcid"}).text = orcid
    name = ET.SubElement(contrib, "name")
    if author.get("surname") or author.get("given_names"):
        ET.SubElement(name, "surname").text = author.get("surname", "")
        if author.get("given_names"):
            ET.SubElement(name, "given-names").text = author["given_names"]
    else:
        words = (author.get("name", "") or "").split()
        if len(words) >= 2:
            ET.SubElement(name, "surname").text = words[-1]
            ET.SubElement(name, "given-names").text = " ".join(words[:-1])
        else:
            ET.SubElement(name, "surname").text = " ".join(words)
    if author.get("email"):
        ET.SubElement(contrib, "email").text = author["email"]


def _add_pub_date(article_meta, metadata: Dict[str, Any]) -> None:
    date = metadata.get("pub_date")
    year = month = day = None
    if date:
        parts = str(date).split("-")
        year = parts[0] if len(parts) >= 1 else None
        month = parts[1] if len(parts) >= 2 else None
        day = parts[2] if len(parts) >= 3 else None
    elif metadata.get("pub_year"):
        year = str(metadata["pub_year"])
    if not year:
        return
    pd = ET.SubElement(article_meta, "pub-date",
                       {"publication-format": "electronic", "date-type": "pub"})
    if day:
        ET.SubElement(pd, "day").text = day
    if month:
        ET.SubElement(pd, "month").text = month
    ET.SubElement(pd, "year").text = year


def _add_permissions(article_meta, metadata: Dict[str, Any], article) -> None:
    holder = metadata.get("copyright_holder")
    cyear = metadata.get("copyright_year")
    statement = metadata.get("copyright_statement")
    lurl = metadata.get("license_url")
    ltext = metadata.get("license_text")
    if not any([holder, cyear, statement, lurl, ltext]):
        return
    perm = ET.SubElement(article_meta, "permissions")
    if statement or holder or cyear:
        if not statement:
            statement = ("© " + (f"{cyear} " if cyear else "") + (holder or "")).strip()
        ET.SubElement(perm, "copyright-statement").text = statement
    if cyear:
        ET.SubElement(perm, "copyright-year").text = str(cyear)
    if holder:
        ET.SubElement(perm, "copyright-holder").text = holder
    if lurl or ltext:
        lic_attrs = {}
        if lurl:
            lic_attrs["xlink:href"] = lurl
            article.set("xmlns:xlink", _XLINK_NS)
        lic = ET.SubElement(perm, "license", lic_attrs)
        ET.SubElement(lic, "license-p").text = ltext or lurl


def _add_funding(article_meta, metadata: Dict[str, Any]) -> None:
    funding = metadata.get("funding")
    if not funding:
        return
    fg = ET.SubElement(article_meta, "funding-group")
    for item in funding:
        ag = ET.SubElement(fg, "award-group")
        if isinstance(item, dict):
            if item.get("funder"):
                ET.SubElement(ag, "funding-source").text = item["funder"]
            if item.get("award_id"):
                ET.SubElement(ag, "award-id").text = item["award_id"]
        elif item:
            ET.SubElement(ag, "funding-source").text = str(item)


def build_jats_xml(
    paragraphs: List[str], title: Optional[str] = None,
    journal_title: Optional[str] = None,
    authors: Optional[List[Dict[str, Any]]] = None,
    metadata: Optional[Dict[str, Any]] = None,
) -> str:
    """Convert an edited manuscript (list of paragraph strings) into JATS
    Journal Publishing XML. Structure is inferred heuristically:
    - first paragraph -> article title (unless provided);
    - a following byline -> structured <contrib-group> authors (+ affiliation);
    - an inline 'Keywords:' line -> <kwd-group>;
    - section headings split the body (numbered subsections nest);
    - an 'Abstract' section is lifted into front matter;
    - paragraphs after a 'References' heading become a structured reference
      list (<element-citation> for journal articles, else <mixed-citation>);
    - in-text [n] citations are linked to their <ref> via <xref>.
    Optional `authors` (list of structured dicts: surname/given_names or name,
    orcid, corresponding, email) override byline parsing. Optional `metadata`
    adds article DOI, pub-date, copyright/license permissions, and funding.
    ElementTree escapes all special characters automatically."""
    metadata = metadata or {}
    nonblank = [p.strip() for p in (paragraphs or []) if p and p.strip()]

    article_title = title
    rest = nonblank
    if not article_title and nonblank:
        article_title = nonblank[0]
        rest = nonblank[1:]

    # Optional byline + affiliation immediately after the title. Multi-author
    # bylines are detected directly; a single-author byline is trusted only when
    # the next line is an affiliation.
    authors_line = affiliation_line = None
    if rest and _looks_like_author_line(rest[0]):
        authors_line = rest[0]
        rest = rest[1:]
        if rest and _looks_like_affiliation(rest[0]):
            affiliation_line = rest[0]
            rest = rest[1:]
    elif len(rest) >= 2 and _looks_like_affiliation(rest[1]) and _looks_like_single_author(rest[0]):
        authors_line, affiliation_line = rest[0], rest[1]
        rest = rest[2:]

    # Split off the reference list at the first References-style heading.
    ref_start = None
    for idx, text in enumerate(rest):
        if _normalize_heading(text) in _REFERENCE_HEADINGS:
            ref_start = idx
            break
    if ref_start is not None:
        body_items = rest[:ref_start]
        ref_items = rest[ref_start + 1:]
    else:
        body_items = rest
        ref_items = []

    keywords, body_items = _extract_keywords(body_items)
    valid_refs = set(range(1, len(ref_items) + 1))

    # Group body paragraphs into sections by heading.
    sections: List[Dict[str, Any]] = []
    current: Dict[str, Any] = {"title": None, "paras": []}
    for text in body_items:
        if _is_jats_heading(text):
            if current["title"] is not None or current["paras"]:
                sections.append(current)
            current = {"title": text, "paras": []}
        else:
            current["paras"].append(text)
    if current["title"] is not None or current["paras"]:
        sections.append(current)

    # Build the JATS tree.
    article = ET.Element("article", {
        "article-type": "research-article",
        "dtd-version": "1.3",
    })
    front = ET.SubElement(article, "front")
    if journal_title:
        journal_meta = ET.SubElement(front, "journal-meta")
        jt_group = ET.SubElement(journal_meta, "journal-title-group")
        ET.SubElement(jt_group, "journal-title").text = journal_title
    article_meta = ET.SubElement(front, "article-meta")
    # JATS article-meta child order: article-id, title-group, contrib-group, aff,
    # pub-date, permissions, abstract, kwd-group, funding-group.
    if metadata.get("article_doi"):
        ET.SubElement(article_meta, "article-id",
                      {"pub-id-type": "doi"}).text = metadata["article_doi"]

    title_group = ET.SubElement(article_meta, "title-group")
    ET.SubElement(title_group, "article-title").text = article_title or "Untitled Manuscript"

    if authors:
        contrib_group = ET.SubElement(article_meta, "contrib-group")
        for a in authors:
            _add_structured_contrib(contrib_group, a)
        if affiliation_line:
            ET.SubElement(article_meta, "aff", {"id": "aff1"}).text = affiliation_line
    elif authors_line:
        contrib_group = ET.SubElement(article_meta, "contrib-group")
        for tok in _author_tokens(authors_line):
            contrib = ET.SubElement(contrib_group, "contrib", {"contrib-type": "author"})
            _contrib_name_elem(contrib, tok)
        if affiliation_line:
            ET.SubElement(article_meta, "aff", {"id": "aff1"}).text = affiliation_line

    _add_pub_date(article_meta, metadata)
    _add_permissions(article_meta, metadata, article)

    # Lift an Abstract section into the front matter.
    abstract_section = next(
        (s for s in sections if s["title"] and _normalize_heading(s["title"]) == "abstract"),
        None,
    )
    if abstract_section:
        sections.remove(abstract_section)
        abstract_el = ET.SubElement(article_meta, "abstract")
        for p in abstract_section["paras"]:
            ET.SubElement(abstract_el, "p").text = p

    if keywords:
        kwd_group = ET.SubElement(article_meta, "kwd-group")
        for kw in keywords:
            ET.SubElement(kwd_group, "kwd").text = kw

    _add_funding(article_meta, metadata)

    body = ET.SubElement(article, "body")
    _emit_sections(body, sections, valid_refs)

    back = ET.SubElement(article, "back")
    if ref_items:
        ref_list = ET.SubElement(back, "ref-list")
        ET.SubElement(ref_list, "title").text = "References"
        for n, text in enumerate(ref_items, 1):
            _add_reference(ref_list, n, text)

    ET.indent(article)  # pretty-print (Python 3.9+)
    xml_body = ET.tostring(article, encoding="unicode")
    return f'<?xml version="1.0" encoding="UTF-8"?>\n{JATS_DOCTYPE}\n{xml_body}\n'


def write_jats_xml(paragraphs: List[str], out_path: str, **kwargs) -> str:
    """Build JATS XML and write it to out_path. Returns out_path."""
    with open(out_path, "w", encoding="utf-8") as f:
        f.write(build_jats_xml(paragraphs, **kwargs))
    return out_path


def validate_jats(xml_text: str):
    """Structural validation of generated JATS XML. Checks well-formedness and
    the presence of required elements; verifies every in-text <xref> points to a
    real <ref>. Returns (ok, issues). This is a pragmatic structural check, not
    full DTD validation (which would require the JATS DTD files)."""
    issues: List[str] = []
    body = "\n".join(
        ln for ln in (xml_text or "").splitlines() if not ln.strip().startswith("<!DOCTYPE")
    )
    try:
        root = ET.fromstring(body)
    except ET.ParseError as e:
        return False, [f"Not well-formed XML: {e}"]

    if root.tag != "article":
        issues.append("Root element is not <article>.")
    title = root.find("./front/article-meta/title-group/article-title")
    if title is None or not (title.text and title.text.strip()):
        issues.append("Missing or empty <article-title>.")
    if root.find("./body") is None:
        issues.append("Missing <body>.")

    ref_ids = {r.get("id") for r in root.findall(".//ref")}
    for x in root.findall('.//xref[@ref-type="bibr"]'):
        if x.get("rid") not in ref_ids:
            issues.append(f"<xref> points to a missing reference: {x.get('rid')}")

    return (len(issues) == 0), issues


def generate_cover_letter(abstract: str, journal_name: str, settings: Dict[str, Any]) -> str:
    prompt = (
        f"Write a professional, compelling manuscript submission cover letter "
        f"tailored to the journal '{journal_name}' based on the following "
        f"abstract/intro. Return only the letter content without any markdown boxes.\n\n"
        f"Abstract:\n{abstract}"
    )
    try:
        return _generate_text(prompt, settings=settings)
    except Exception as e:
        print(f"Cover letter error: {e}")
        return "Error generating cover letter."


def generate_title_abstract_polish(abstract: str, settings: Dict[str, Any]) -> str:
    prompt = (
        "Based on this draft abstract, generate 3 highly optimized, impactful "
        "title options and 1 fully polished, compelling abstract that maximizes "
        "chances of journal acceptance. Format it beautifully in Markdown.\n\n"
        f"Text:\n{abstract}"
    )
    try:
        return _generate_text(prompt, settings=settings)
    except Exception as e:
        print(f"Title/abstract polish error: {e}")
        return "Error generating polish."


# --- AI peer reviewer ---

# Cap the manuscript text sent to the reviewer to keep within model context.
_AI_REVIEW_MAX_CHARS = 16000


def generate_ai_review(manuscript_text: str, settings: Dict[str, Any]) -> str:
    """Act as an expert academic peer reviewer and produce a structured review
    report in Markdown (summary, significance, strengths, weaknesses,
    methodology, minor issues, and a clear recommendation)."""
    excerpt = (manuscript_text or "").strip()[:_AI_REVIEW_MAX_CHARS]
    if not excerpt:
        return "_No manuscript text was available to review._"
    prompt = (
        "You are an expert academic peer reviewer for a scholarly journal. "
        "Review the manuscript below as a referee would and write a rigorous, "
        "constructive review report. Be specific and cite concrete issues from "
        "the text. Use this exact Markdown structure with these headings:\n\n"
        "# AI Peer Review Report\n"
        "## Summary\n(2-4 sentences summarizing the manuscript's aims and findings.)\n"
        "## Significance & Originality\n(Is the contribution novel and important to its field?)\n"
        "## Strengths\n(Bulleted list.)\n"
        "## Major Concerns\n(Bulleted list of substantive issues: methodology, evidence, "
        "logic, unsupported claims, missing controls, statistics.)\n"
        "## Minor Issues\n(Bulleted list: clarity, figures, typos, references.)\n"
        "## Recommendation\n(State exactly ONE of: Accept, Minor Revision, Major Revision, "
        "Reject — then one sentence justifying it.)\n"
        "## Confidence\n(High / Medium / Low — your confidence in this assessment.)\n\n"
        "If the text is only a partial manuscript, review what is provided and say so. "
        "Do not invent content that is not present.\n\n"
        f"MANUSCRIPT:\n{excerpt}"
    )
    try:
        return _generate_text(prompt, settings=settings)
    except Exception as e:
        print(f"AI review error: {e}")
        return "Error generating AI peer review."


# --- Orchestration ---

def process_document_async(
    paras: List[str], settings: Dict[str, Any], edit_style: str, ref_style: str,
    lang: str, custom_dict: str, use_crossref: bool, progress_callback,
    enabled_rule_ids: Optional[List[str]] = None, custom_rules: str = "",
    use_serper: bool = False, serper_key: str = "",
) -> Tuple[List[str], List[Dict[str, Any]]]:
    """Returns (edited_paras, queries). Each query is
    {"index": <paragraph index in paras>, "snippet", "query", "suggestion"}
    ("suggestion" may be None)."""
    edited_paras = [""] * len(paras)
    all_queries: List[Dict[str, Any]] = []
    task_indices = [i for i, p in enumerate(paras) if p.strip()]

    for i in range(len(paras)):
        if not paras[i].strip():
            edited_paras[i] = paras[i]

    CHUNK_SIZE = 5
    chunks = []
    for i in range(0, len(task_indices), CHUNK_SIZE):
        chunk_inds = task_indices[i:i + CHUNK_SIZE]
        chunk_texts = [paras[idx] for idx in chunk_inds]
        chunks.append((chunk_inds, chunk_texts))

    total_chunks = len(chunks)
    if total_chunks == 0:
        return edited_paras, all_queries

    completed = 0
    with concurrent.futures.ThreadPoolExecutor(max_workers=_chunk_pool_size()) as ex:
        future_to_chunk = {
            ex.submit(ai_edit_chunk, chunk_texts, settings, edit_style, ref_style,
                      lang, custom_dict, use_crossref,
                      enabled_rule_ids, custom_rules,
                      use_serper, serper_key): chunk_inds
            for chunk_inds, chunk_texts in chunks
        }
        for future in concurrent.futures.as_completed(future_to_chunk):
            chunk_inds = future_to_chunk[future]
            try:
                result, chunk_queries = future.result()
                for idx, edited_text in zip(chunk_inds, result):
                    edited_paras[idx] = edited_text
                for q in chunk_queries:
                    all_queries.append({
                        "index": chunk_inds[q["local_index"]],
                        "snippet": q["snippet"],
                        "query": q["query"],
                        "suggestion": q.get("suggestion"),
                    })
            except Exception as exc:
                print(f"Chunk generated an exception: {exc}")
                for idx in chunk_inds:
                    edited_paras[idx] = paras[idx]
            completed += 1
            if progress_callback:
                progress_callback(completed / total_chunks)
    all_queries.sort(key=lambda q: q["index"])
    return edited_paras, all_queries
