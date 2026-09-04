"""Make DOIs, URLs and e-mail addresses clickable in the redline.

Asked for by the editorial team. It is the riskiest thing in the plan because it
touches the finished redline, and the redline's one absolute requirement is that the
paragraph list stays exactly as long and exactly in order as the original — every
tracked change is positional, so a paragraph gained or lost moves all of them onto the
wrong text and the file still looks perfect.

**So this never adds, removes or reorders a paragraph, and never rewrites text.** It
wraps runs that already exist in a `w:hyperlink` element, in place. A run's text,
formatting and tracked-change state are untouched; the only change is that Word now
knows where it points.

Three rules that are not obvious:

* A link inside `w:del` must not be created. Deleted text is text the editor is being
  asked to remove — making it clickable invites them to follow a reference that the
  redline is proposing to delete.
* A run has to be *split* when the address is only part of it. `Contact
  jaya@gbu.ac.in for data` is usually one run, and wrapping the whole run would make
  the sentence a link.
* The relationship is created on the part, not the document, and reused per URL —
  a manuscript citing the same DOI eight times should carry one relationship, not
  eight.
"""

from __future__ import annotations

import re
from typing import Dict, List, Optional

from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

#: What to linkify, in the order they are tried. A DOI is matched before a bare URL so
#: `https://doi.org/10.1000/xyz` is recognised as the DOI it is.
_PATTERNS = (
    ("doi", re.compile(r"https?://(?:dx\.)?doi\.org/10\.\d{4,9}/[^\s,;)\]}>\"']+")),
    ("doi", re.compile(r"(?<![\w/.])doi\s*:\s*(10\.\d{4,9}/[^\s,;)\]}>\"']+)", re.I)),
    ("doi", re.compile(r"(?<![\w/.:])10\.\d{4,9}/[^\s,;)\]}>\"']+")),
    ("url", re.compile(r"https?://[^\s,;)\]}>\"']+")),
    ("url", re.compile(r"(?<![\w@.])www\.[^\s,;)\]}>\"']+")),
    ("email", re.compile(r"(?<![\w.+-])[\w.+-]+@[\w-]+\.[\w.-]+[\w]")),
)

#: Trailing punctuation that belongs to the sentence, not to the address. A DOI at the
#: end of a reference is followed by a full stop and the stop is not part of it.
_TRAILING = ".,;:)]}>\"'"


def find_targets(text: str) -> List[tuple]:
    """[(start, end, href)] for every address in `text`, non-overlapping.

    Earlier matches win, so a DOI inside a `doi.org` URL is not also matched as a bare
    DOI and wrapped twice.
    """
    found: List[tuple] = []
    taken: List[range] = []
    for kind, pattern in _PATTERNS:
        for m in pattern.finditer(text):
            start, end = m.span()
            raw = text[start:end]
            while raw and raw[-1] in _TRAILING:
                raw, end = raw[:-1], end - 1
            if not raw or any(start < r.stop and end > r.start for r in taken):
                continue
            taken.append(range(start, end))
            if kind == "email":
                href = f"mailto:{raw}"
            elif kind == "doi":
                bare = re.sub(r"^(?:https?://(?:dx\.)?doi\.org/|doi\s*:\s*)", "", raw,
                              flags=re.I)
                href = f"https://doi.org/{bare}"
            else:
                href = raw if raw.lower().startswith("http") else f"https://{raw}"
            found.append((start, end, href))
    return sorted(found)


def _split_run(run_el, start: int, end: int):
    """Split one `w:r` so the characters between `start` and `end` are their own run.

    Returns the middle run. The run's properties are copied onto each piece, so a
    hyperlink inside a bold sentence stays bold.
    """
    import copy

    t = run_el.find(qn("w:t"))
    if t is None or t.text is None:
        return None
    text = t.text
    before, middle, after = text[:start], text[start:end], text[end:]

    t.text = middle
    t.set(qn("xml:space"), "preserve")
    parent = run_el.getparent()
    index = list(parent).index(run_el)

    def piece(content):
        el = copy.deepcopy(run_el)
        el.find(qn("w:t")).text = content
        el.find(qn("w:t")).set(qn("xml:space"), "preserve")
        return el

    if after:
        parent.insert(index + 1, piece(after))
    if before:
        parent.insert(index, piece(before))
    return run_el


def _groups(container) -> List[tuple]:
    """The paragraph's runs, split into stretches that may share one hyperlink.

    A group is a run of consecutive `w:r` siblings under the same parent: either the
    paragraph's own runs, or the runs inside one `w:ins`. A `w:del` ends the current
    group, so an address is never assembled across text the redline is deleting, and a
    link never spans two parents — which would mean a `w:hyperlink` half inside a
    tracked insertion and half outside it.
    """
    groups: List[tuple] = []
    current: List = []
    for child in container.iterchildren():
        if child.tag == qn("w:r"):
            current.append(child)
            continue
        if current:
            groups.append((container, current))
            current = []
        if child.tag == qn("w:ins"):
            runs = [r for r in child.iterchildren(qn("w:r"))]
            if runs:
                groups.append((child, runs))
    if current:
        groups.append((container, current))
    return groups


def _run_text(run_el) -> str:
    t = run_el.find(qn("w:t"))
    return (t.text or "") if t is not None else ""


def _wrap_in_hyperlink(span: List, rid: str) -> None:
    """Put the consecutive runs in `span` inside one `w:hyperlink`, nested as Word does.

    The obvious version — insert the hyperlink where the runs are and move them into
    it — produces `<w:ins><w:hyperlink>` when they are inserted runs, and that is **not
    valid OOXML**: `w:ins` is a CT_RunTrackChange, whose content group has no
    `hyperlink` in it. Word writes the other order, `<w:hyperlink><w:ins><w:r>`, and a
    file with the nesting inverted is one Word may quietly "repair" — dropping exactly
    the links this function exists to add, without an error anywhere. Every DOI in a
    real redline is inside `w:ins`, because the copyedit reformats the references, so
    that bug hit 26 links out of 30 and nothing in the output looked wrong.

    So inserted runs are lifted out: the `w:ins` is split around them and they get
    their own `w:ins`, carrying the same author, date and id, inside the hyperlink.
    """
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), rid)
    parent = span[0].getparent()

    if parent.tag != qn("w:ins"):
        parent.insert(list(parent).index(span[0]), link)
        for run_el in span:
            link.append(run_el)
        return

    ins = parent
    siblings = list(ins)
    at = siblings.index(span[0])
    grand = ins.getparent()
    ins_at = list(grand).index(ins)

    def ins_shell():
        el = OxmlElement("w:ins")
        for name, value in ins.attrib.items():
            el.set(name, value)
        return el

    inner = ins_shell()
    for run_el in span:
        inner.append(run_el)
    link.append(inner)

    trailing = siblings[at + len(span):]
    if trailing:
        tail = ins_shell()
        for el in trailing:
            tail.append(el)
        grand.insert(ins_at + 1, tail)
    grand.insert(ins_at + 1, link)
    if len(ins) == 0:
        grand.remove(ins)


def linkify_paragraph(paragraph, cache: Optional[Dict[str, str]] = None) -> int:
    """Wrap every address in this paragraph in a hyperlink. Returns how many.

    Runs inside `w:del` are skipped, and so are runs already inside a `w:hyperlink` —
    the author's own links are left exactly as they were.
    """
    cache = cache if cache is not None else {}
    part = paragraph.part
    made = 0

    for parent, runs in _groups(paragraph._p):
        text = "".join(_run_text(r) for r in runs)
        targets = find_targets(text)
        if not targets:
            continue
        # Last address first. Wrapping one splits runs and renumbers everything after
        # it; working backwards means the offsets computed here stay true.
        for start, end, href in reversed(targets):
            span = _runs_covering(runs, start, end)
            if span is None:
                continue
            rid = cache.get(href)
            if rid is None:
                rid = part.relate_to(href, RT.HYPERLINK, is_external=True)
                cache[href] = rid
            _wrap_in_hyperlink(span, rid)
            made += 1
    return made


def _runs_covering(runs: List, start: int, end: int) -> Optional[List]:
    """Split `runs` so the characters `[start, end)` are exactly one run sequence.

    Addresses do not respect run boundaries. Word splits runs wherever it likes — a
    spell-check boundary, an rsid, a stray formatting change — and in a real manuscript
    `doi:10.1016/j.ajhg.2016.09.015` arrived as two runs, `doi:10.1016/j.ajhg` and
    `.2016.09.015.`. Detecting per run linked the first half to a DOI that does not
    exist: a link that looks right in Word and resolves to nothing, which is worse than
    leaving the text plain.

    So detection happens on the group's whole text and the span is cut out of however
    many runs it crosses.
    """
    pos = 0
    first = last = None
    for i, run in enumerate(runs):
        length = len(_run_text(run))
        if first is None and pos + length > start:
            first, first_off = i, start - pos
        if first is not None and pos + length >= end:
            last, last_off = i, end - pos
            break
        pos += length
    if first is None or last is None:
        return None

    if first == last:
        middle = _split_run(runs[first], first_off, last_off)
        return [middle] if middle is not None else None

    head = _split_run(runs[first], first_off, len(_run_text(runs[first])))
    tail = _split_run(runs[last], 0, last_off)
    if head is None or tail is None:
        return None
    return [head] + runs[first + 1:last] + [tail]


def linkify_document(document) -> int:
    """Linkify every body paragraph and table cell. Returns how many links were made.

    The paragraph list is only ever read. Nothing here appends, removes or reorders a
    paragraph, which is the one thing that would corrupt the redline.
    """
    cache: Dict[str, str] = {}
    made = 0
    for p in document.paragraphs:
        made += linkify_paragraph(p, cache)
    for table in document.tables:
        for row in table.rows:
            try:
                cells = list(row.cells)
            except ValueError:               # irregular vertical merges
                from docx.table import _Cell
                cells = [_Cell(tc, table) for tc in row._tr.tc_lst]
            for cell in cells:
                for p in cell.paragraphs:
                    made += linkify_paragraph(p, cache)
    return made
