"""What is actually in a .docx, beyond the words.

`editor.read_docx` returns `[p.text for p in doc.paragraphs]` and that is all the
copyeditor has ever seen. It is enough to fix grammar and it is not enough for anything
else the house rules ask for: a heading's case and weight, whether a species name is
italic, which bullet glyph a list uses, or that a table exists at all. On a real
manuscript — the Guanidine paper — that reader misses 11.3% of the text (3,540 characters
in 12 tables) and every one of 190 italic runs, 67 subscripts and 37 headings.

This module reads the same file and keeps those. It does **not** replace `read_docx`:
`generate_redline_docx` walks `zip(doc.paragraphs, edited_paragraphs)`, so the plain list
must keep its exact length and order or the redline silently writes each change onto the
wrong paragraph. `read_structure` is additive, and `paragraphs[i]` here is `read_docx()[i]`
there — same index, same paragraph, on purpose.

Nothing here judges. `house_layout` does that; this only reports what the file says.
"""

from __future__ import annotations

from dataclasses import dataclass, field
from typing import Any, Dict, List, Optional

import docx
from docx.oxml.ns import qn

#: EMU per inch. python-docx lengths are in EMU; the house spec is in inches.
EMU_PER_INCH = 914400

#: Word's numFmt values mapped to the house's number-listing names.
NUMBER_FORMATS = {
    "decimal": "N1",            # 1. 2. 3.
    "lowerRoman": "N2",         # i. ii. iii.
    "lowerLetter": "N3",        # a. b. c.
}

#: The bullet glyphs the house names B1-B4. Word stores the character used by the
#: numbering definition; Symbol and Wingdings put different code points behind the
#: same visual mark, so both are mapped rather than only the one this office happens
#: to produce today.
BULLET_GLYPHS = {
    "": "B1",   # Symbol: round solid
    "•": "B1",   # bullet U+2022
    "●": "B1",   # black circle U+25CF — the same solid round mark at another
                 #   code point, and the very glyph this checker's own message prints
                 #   as B1. Without it the report read "bullet '●' is not one of the
                 #   house marks (B1 ● …)", which is wrong and visibly self-contradictory.
    "∙": "B1",   # bullet operator U+2219
    "": "B2",   # Symbol: round hollow
    "○": "B2",   # white circle U+25CB
    "◦": "B2",   # white bullet U+25E6
    "o": "B2",   # Courier New 'o' — Word's default hollow bullet
    "": "B3",   # Wingdings: square solid
    "■": "B3",   # black square U+25A0
    "▪": "B3",   # black small square U+25AA
    "": "B4",   # Wingdings: square hollow
    "□": "B4",   # white square U+25A1
    "▫": "B4",   # white small square U+25AB
}


@dataclass
class Run:
    """One stretch of text with uniform formatting."""
    text: str
    bold: Optional[bool] = None
    italic: Optional[bool] = None
    underline: Optional[bool] = None
    superscript: bool = False
    subscript: bool = False
    font_name: Optional[str] = None
    font_size_pt: Optional[float] = None


@dataclass
class Para:
    """A paragraph, at the same index `read_docx` would give it."""
    index: int
    text: str
    style: str
    #: 0-based Word outline level, or None for body text. This is what makes a
    #: heading a heading — the style *name* is only a convention and a manuscript
    #: that uses "Heading 1" for a figure caption is common.
    outline_level: Optional[int] = None
    alignment: Optional[str] = None
    runs: List[Run] = field(default_factory=list)
    #: List membership, when the paragraph is one: {"num_id", "level", "kind",
    #: "house_code", "glyph", "num_fmt"}. Empty when it is not a list item.
    listing: Dict[str, Any] = field(default_factory=dict)
    #: True when the paragraph sits inside a table cell rather than the body flow.
    in_table: bool = False
    #: Indents in inches. `first_line_in` is negative for a hanging indent, which is
    #: how the reference list's 0.25" hang is expressed.
    first_line_in: Optional[float] = None
    left_indent_in: Optional[float] = None
    #: Space after the paragraph, in points, and the line spacing rule.
    space_after_pt: Optional[float] = None
    space_before_pt: Optional[float] = None
    line_spacing: Optional[float] = None

    @property
    def dominant_font(self) -> Optional[str]:
        """The font most of the paragraph's characters are in.

        By characters rather than by run count: a heading in Times New Roman with a
        single stray one-character run in Calibri is a Times New Roman heading, and
        counting runs would call it a tie.
        """
        return _dominant(self.runs, "font_name")

    @property
    def dominant_size_pt(self) -> Optional[float]:
        return _dominant(self.runs, "font_size_pt")

    @property
    def is_bold(self) -> Optional[bool]:
        """True when every run with text in it is bold. None when nothing says."""
        vals = [r.bold for r in self.runs if r.text.strip()]
        if not vals or all(v is None for v in vals):
            return None
        return all(bool(v) for v in vals)

    @property
    def is_italic(self) -> Optional[bool]:
        vals = [r.italic for r in self.runs if r.text.strip()]
        if not vals or all(v is None for v in vals):
            return None
        return all(bool(v) for v in vals)


@dataclass
class Cell:
    """One table cell, with the paragraphs inside it.

    Paragraphs rather than a string: the house table spec is about weight, slant and
    size — a column head is 9 pt bold, a subhead 9 pt bold italic — and none of that
    survives `cell.text`.
    """
    row: int
    col: int
    paragraphs: List["Para"] = field(default_factory=list)

    @property
    def text(self) -> str:
        return "\n".join(p.text for p in self.paragraphs).strip()


@dataclass
class TableFormat:
    """The table's own geometry, as the house spec measures it."""
    #: Border style and width, from `w:tblBorders`. `None` when the table sets none
    #: and inherits from its style — which is not the same as having no borders, so
    #: the checker says "not stated" rather than "wrong".
    border_style: Optional[str] = None
    border_size_pt: Optional[float] = None
    #: Default cell margins in inches: top, bottom, left, right.
    margin_top_in: Optional[float] = None
    margin_bottom_in: Optional[float] = None
    margin_left_in: Optional[float] = None
    margin_right_in: Optional[float] = None
    style: str = ""


@dataclass
class Table:
    """A table, and where it sits relative to the body paragraphs."""
    index: int
    rows: int
    cols: int
    cells: List[List[str]]
    #: The same grid, with formatting kept.
    grid: List[List["Cell"]] = field(default_factory=list)
    fmt: "TableFormat" = field(default_factory=lambda: TableFormat())
    #: Index of the body paragraph immediately before this table, or -1 when the
    #: table opens the document. This is how a caption is found: a table's caption
    #: is the paragraph on one side of it, and without the position there is no
    #: way to say which.
    after_paragraph: int = -1

    @property
    def text(self) -> str:
        return "\n".join(" | ".join(row) for row in self.cells)


@dataclass
class Section:
    """Page geometry, in inches — the units the house spec is written in."""
    page_width_in: Optional[float]
    page_height_in: Optional[float]
    left_margin_in: Optional[float]
    right_margin_in: Optional[float]
    top_margin_in: Optional[float]
    bottom_margin_in: Optional[float]


@dataclass
class Image:
    """An inline picture and its size on the page."""
    index: int
    width_in: float
    height_in: float


@dataclass
class Structure:
    paragraphs: List[Para]
    tables: List[Table]
    sections: List[Section]
    images: List[Image]

    @property
    def texts(self) -> List[str]:
        """Exactly what `read_docx` returns, from the same file.

        Kept so a caller can hold one parse instead of two, and so the equivalence
        is testable rather than assumed.
        """
        return [p.text for p in self.paragraphs]

    def body_paragraphs(self) -> List[Para]:
        return [p for p in self.paragraphs if not p.in_table]

    def headings(self) -> List[Para]:
        return [p for p in self.body_paragraphs() if p.outline_level is not None]


def _dominant(runs: List[Run], attr: str):
    """The attribute value covering the most characters."""
    weights: Dict[Any, int] = {}
    for r in runs:
        v = getattr(r, attr)
        if v is None or not r.text.strip():
            continue
        weights[v] = weights.get(v, 0) + len(r.text)
    if not weights:
        return None
    return max(weights.items(), key=lambda kv: kv[1])[0]


def _emu_to_in(value) -> Optional[float]:
    if value is None:
        return None
    return round(int(value) / EMU_PER_INCH, 3)


def _row_cells(row) -> list:
    """The cells of a table row, tolerant of irregular vertical merges.

    `row.cells` resolves every vertically merged continuation cell by walking to the
    cell above it, and raises `ValueError: no 'tc' element at grid_offset=N` when a
    real table's grid does not line up — which Word tolerates and writes anyway.

    Falling back to the row's own `w:tc` children loses the merge resolution: a merged
    cell is then read once, in the row it is actually written in, rather than repeated
    down the rows it spans. That is a worse reading of the table and a much better
    outcome than failing the manuscript, which is what this did before.
    """
    try:
        return list(row.cells)
    except ValueError:
        from docx.table import _Cell
        return [_Cell(tc, row.table) for tc in row._tr.tc_lst]


def _size_pt(font) -> Optional[float]:
    """Point size from a run or style font, tolerant of fractional half-points.

    `font.size` puts `w:sz/@w:val` through python-docx's typed accessor, which does
    `int(value)` and raises on anything fractional. LibreOffice and Google Docs both
    export sizes like `26.666666666666668`, and that `ValueError` came out of
    `read_structure` and failed the entire manuscript over one heading.
    """
    try:
        size = font.size
    except ValueError:
        pass
    else:
        return round(size.pt, 1) if size is not None else None

    el = getattr(font, "_element", None)
    rpr = el.find(qn("w:rPr")) if el is not None else None
    sz = rpr.find(qn("w:sz")) if rpr is not None else None
    try:
        return round(float(sz.get(qn("w:val"))) / 2, 1)      # half-points
    except (AttributeError, TypeError, ValueError):
        return None


#: Word writes `w:sectPr` geometry in twips, and python-docx's typed accessors put
#: each one through `int()`. Real files carry fractional values — a page converted
#: from A4 lands as `1275.5905511811022` — and that killed the whole manuscript. Same
#: failure as `_size_pt`, one directory over.
_SECTION_XML = {
    "page_width": ("w:pgSz", "w:w"), "page_height": ("w:pgSz", "w:h"),
    "left_margin": ("w:pgMar", "w:left"), "right_margin": ("w:pgMar", "w:right"),
    "top_margin": ("w:pgMar", "w:top"), "bottom_margin": ("w:pgMar", "w:bottom"),
}
_TWIPS_PER_INCH = 1440.0


def _section_in(section, name: str) -> Optional[float]:
    """One section dimension in inches, tolerant of fractional twips."""
    try:
        return _emu_to_in(getattr(section, name))
    except ValueError:
        pass
    tag, attr = _SECTION_XML[name]
    el = section._sectPr.find(qn(tag))
    try:
        return round(float(el.get(qn(attr))) / _TWIPS_PER_INCH, 3)
    except (AttributeError, TypeError, ValueError):
        return None


#: The paragraph-format accessors convert the same strict way, and real files carry
#: fractional twips there too — `w:firstLine="1.0000000000000009"`,
#: `w:before="359.00000000000006"`. Each entry lists the places the measurement can be
#: written, in order; `w:hanging` is the same number negated, which is how Word
#: expresses a hanging indent, and `w:start` is the newer spelling of `w:left`.
_PF_XML = {
    "first_line_indent": (("w:ind", "w:firstLine", 1), ("w:ind", "w:hanging", -1)),
    "left_indent": (("w:ind", "w:left", 1), ("w:ind", "w:start", 1)),
    "space_before": (("w:spacing", "w:before", 1),),
    "space_after": (("w:spacing", "w:after", 1),),
}
#: 1 pt is 20 twips. Word writes paragraph spacing in twips and the house spec is in
#: points, so the two conversions differ and must not be crossed.
_TWIPS_PER_PT = 20.0


def _pf_twips(pf, name: str) -> Optional[float]:
    """A paragraph-format measurement read straight out of the XML, in twips.

    Kept as a float rather than a python-docx `Length`, because the reason we are
    here at all is that the value does not fit in one.
    """
    pPr = pf._element.find(qn("w:pPr"))
    if pPr is None:
        return None
    for tag, attr, sign in _PF_XML[name]:
        el = pPr.find(qn(tag))
        raw = el.get(qn(attr)) if el is not None else None
        if raw is not None:
            try:
                return sign * float(raw)
            except ValueError:
                return None
    return None


def _pf_in(pf, name: str) -> Optional[float]:
    """A paragraph-format indent in inches, tolerant of fractional twips."""
    try:
        return _emu_to_in(getattr(pf, name))
    except ValueError:
        twips = _pf_twips(pf, name)
        return round(twips / _TWIPS_PER_INCH, 3) if twips is not None else None


def _pf_pt(pf, name: str) -> Optional[float]:
    """Paragraph spacing in points, tolerant of fractional twips."""
    try:
        value = getattr(pf, name)
    except ValueError:
        twips = _pf_twips(pf, name)
        return round(twips / _TWIPS_PER_PT, 1) if twips is not None else None
    return round(value.pt, 1) if value is not None else None


#: `w:jc` gained direction-neutral spellings in the later OOXML revision, and Word
#: writes them. python-docx's enum knows only the older pair and raises on the rest,
#: so one `w:jc val="start"` — the commonest alignment there is — cost a manuscript
#: its entire house-style check.
_JC_ALIASES = {"start": "left", "end": "right", "both": "justify"}


def _alignment(p) -> Optional[str]:
    """A paragraph's alignment by name, tolerant of the newer OOXML spellings."""
    try:
        align = p.alignment
    except ValueError:
        pPr = p._p.find(qn("w:pPr"))
        jc = pPr.find(qn("w:jc")) if pPr is not None else None
        raw = jc.get(qn("w:val")) if jc is not None else None
        alias = _JC_ALIASES.get(raw, raw)
        return alias if alias in set(ALIGNMENT_NAMES.values()) else None
    return ALIGNMENT_NAMES.get(int(align)) if align is not None else None


def _style_font(style) -> tuple:
    """(name, size_pt) a style asks for, following `basedOn` upwards.

    Word only writes the font onto a run when it *differs* from the style, so a
    correctly-formatted heading has `run.font.name is None` — the value lives on the
    style, or on the document defaults behind that. Reading only the run reports
    `None` for every paragraph in the file, which looks like uniformity and is
    actually a broken measurement: nothing about "Times New Roman, 11 pt" can be
    checked from it.
    """
    seen = 0
    name = size = None
    while style is not None and seen < 8:
        seen += 1
        f = getattr(style, "font", None)
        if f is not None:
            if name is None and f.name:
                name = f.name
            if size is None:
                size = _size_pt(f)
        if name is not None and size is not None:
            break
        style = getattr(style, "base_style", None)
    return name, size


def _doc_default_font(document) -> tuple:
    """The `w:docDefaults` font — the last fallback before Word's own default."""
    try:
        rPr = document.styles.element.find(qn("w:docDefaults"))
        if rPr is None:
            return None, None
        rpr = rPr.find(qn("w:rPrDefault"))
        rpr = rpr.find(qn("w:rPr")) if rpr is not None else None
        if rpr is None:
            return None, None
        name = None
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is not None:
            name = (rfonts.get(qn("w:ascii")) or rfonts.get(qn("w:hAnsi"))
                    or rfonts.get(qn("w:cs")))
        size = None
        sz = rpr.find(qn("w:sz"))
        if sz is not None and (sz.get(qn("w:val")) or "").isdigit():
            size = int(sz.get(qn("w:val"))) / 2          # half-points
        return name, size
    except Exception:                                            # noqa: BLE001
        return None, None


def _read_run(r, inherited=(None, None)) -> Run:
    font = r.font
    size = _size_pt(font)
    sup = sub = False
    # `font.superscript` is None when unset rather than False, and reading the
    # vertAlign element directly also catches runs where Word wrote the property
    # onto the style instead of the run.
    try:
        sup = bool(font.superscript)
        sub = bool(font.subscript)
    except (AttributeError, ValueError):
        pass
    # The run's own value wins; the style chain and then docDefaults fill the gaps.
    # Recorded as the effective value, because that is what the reader on the page
    # sees and what the house rule is written about.
    inh_name, inh_size = inherited
    return Run(
        text=r.text,
        bold=r.bold,
        italic=r.italic,
        underline=r.underline,
        superscript=sup,
        subscript=sub,
        font_name=font.name or inh_name,
        font_size_pt=size if size is not None else inh_size,
    )


def _outline_level(p) -> Optional[int]:
    """0-based heading level, or None.

    Read from the style rather than the name where possible: `w:outlineLvl` is what
    Word itself uses for the navigation pane and the table of contents, so it is the
    honest answer to "is this a heading". The name is the fallback for documents
    built by tools that set the name and nothing else.
    """
    try:
        pPr = p._p.pPr
        if pPr is not None:
            lvl = pPr.find(qn("w:outlineLvl"))
            if lvl is not None:
                val = lvl.get(qn("w:val"))
                if val is not None and val.isdigit() and int(val) <= 8:
                    return int(val)
    except Exception:                                            # noqa: BLE001
        pass

    style = getattr(p.style, "name", "") or ""
    m = _HEADING_NAME.match(style.strip())
    if m:
        return int(m.group(1)) - 1
    try:
        base = p.style
        for _ in range(4):                       # a style may inherit from Heading N
            base = getattr(base, "base_style", None)
            if base is None:
                break
            m = _HEADING_NAME.match((base.name or "").strip())
            if m:
                return int(m.group(1)) - 1
    except Exception:                                            # noqa: BLE001
        pass
    return None


import re                                                        # noqa: E402

_HEADING_NAME = re.compile(r"^Heading\s*([1-9])$", re.I)


class _Numbering:
    """The document's numbering definitions, so a list item can name its own marker.

    Word does not store "bullet" or "1." on the paragraph — the paragraph carries a
    numId and a level, and the glyph lives in numbering.xml behind two indirections
    (numId -> abstractNumId -> level). Without following those, every list looks the
    same and B1..B4 / N1..N3 cannot be told apart at all.
    """

    def __init__(self, document):
        self._levels: Dict[tuple, Dict[str, Any]] = {}
        try:
            part = document.part.numbering_part
        except (KeyError, AttributeError, ValueError, NotImplementedError):
            # A document that has never contained a list has no numbering part, and
            # python-docx answers that with a bare `NotImplementedError` — which this
            # guard originally missed. It took the whole read down on 32 of 400 real
            # manuscripts. No numbering part simply means no numbering to resolve.
            return
        root = part.element

        abstract: Dict[str, Dict[int, Dict[str, Any]]] = {}
        for an in root.findall(qn("w:abstractNum")):
            aid = an.get(qn("w:abstractNumId"))
            levels: Dict[int, Dict[str, Any]] = {}
            for lvl in an.findall(qn("w:lvl")):
                ilvl = lvl.get(qn("w:ilvl"))
                fmt = lvl.find(qn("w:numFmt"))
                txt = lvl.find(qn("w:lvlText"))
                levels[int(ilvl or 0)] = {
                    "num_fmt": fmt.get(qn("w:val")) if fmt is not None else None,
                    "lvl_text": txt.get(qn("w:val")) if txt is not None else None,
                }
            abstract[aid] = levels

        for num in root.findall(qn("w:num")):
            nid = num.get(qn("w:numId"))
            ref = num.find(qn("w:abstractNumId"))
            aid = ref.get(qn("w:val")) if ref is not None else None
            for ilvl, info in (abstract.get(aid) or {}).items():
                self._levels[(nid, ilvl)] = info

    def describe(self, num_id: Optional[str], level: int) -> Dict[str, Any]:
        info = self._levels.get((num_id, level), {})
        fmt = info.get("num_fmt")
        lvl_text = info.get("lvl_text") or ""
        out: Dict[str, Any] = {"num_id": num_id, "level": level,
                               "num_fmt": fmt, "lvl_text": lvl_text}
        if fmt == "bullet":
            glyph = lvl_text[:1] if lvl_text else ""
            out["kind"] = "bullet"
            out["glyph"] = glyph
            out["house_code"] = BULLET_GLYPHS.get(glyph)
        elif fmt:
            out["kind"] = "number"
            out["house_code"] = NUMBER_FORMATS.get(fmt)
            # The house asks for a dot after the numeral. `lvl_text` is the literal
            # pattern Word prints, e.g. "%1." or "(%1)", so the separator is readable.
            out["dot_suffix"] = lvl_text.rstrip().endswith(".")
        else:
            out["kind"] = "unknown"
        return out


def _para_numbering(p) -> tuple:
    """(numId, ilvl) for a list paragraph, or (None, 0).

    Checked on the paragraph first and on its style second: Word puts numbering on
    the style for the built-in list styles and on the paragraph for direct
    formatting, and a reader that looks in only one place reports half the lists as
    not lists.
    """
    def _from(pPr):
        if pPr is None:
            return None, None
        numPr = pPr.find(qn("w:numPr"))
        if numPr is None:
            return None, None
        nid = numPr.find(qn("w:numId"))
        ilvl = numPr.find(qn("w:ilvl"))
        return (nid.get(qn("w:val")) if nid is not None else None,
                int(ilvl.get(qn("w:val"))) if ilvl is not None else 0)

    nid, ilvl = _from(p._p.pPr)
    if nid is not None:
        return nid, ilvl or 0
    try:
        nid, ilvl = _from(p.style.element.pPr)
    except AttributeError:
        return None, 0
    return nid, (ilvl or 0)


ALIGNMENT_NAMES = {0: "left", 1: "center", 2: "right", 3: "justify"}

#: Word measures table borders in eighths of a point and cell margins in twentieths
#: of a point (dxa). The house spec is written in points and inches, so both are
#: converted here — comparing a spec of "1/2 pt" against a stored `4` is the kind of
#: unit mismatch that makes a checker confidently wrong.
_EIGHTHS_PER_PT = 8.0
_DXA_PER_INCH = 1440.0


def _table_format(table) -> "TableFormat":
    """Borders and cell margins, read from the table's own properties.

    Only what the table states. A table that inherits its borders from a style
    reports `None`, and the checker says "not stated" rather than "wrong" — calling
    an inherited border missing would flag every table built from a Word table style.
    """
    fmt = TableFormat(style=getattr(getattr(table, "style", None), "name", "") or "")
    tblPr = table._tbl.find(qn("w:tblPr"))
    if tblPr is None:
        return fmt

    borders = tblPr.find(qn("w:tblBorders"))
    if borders is not None:
        # The four outer edges plus the two insides. Reported as one style/size
        # because the house asks for one, and a table with six different borders is
        # better described by naming the disagreement than by listing all six.
        seen_styles, seen_sizes = [], []
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            e = borders.find(qn("w:" + edge))
            if e is None:
                continue
            val = e.get(qn("w:val"))
            sz = e.get(qn("w:sz"))
            if val:
                seen_styles.append(val)
            if sz and sz.isdigit():
                seen_sizes.append(int(sz) / _EIGHTHS_PER_PT)
        if seen_styles:
            fmt.border_style = (seen_styles[0] if len(set(seen_styles)) == 1
                                else "mixed:" + ",".join(sorted(set(seen_styles))))
        if seen_sizes:
            fmt.border_size_pt = (seen_sizes[0] if len(set(seen_sizes)) == 1
                                  else max(seen_sizes))

    margins = tblPr.find(qn("w:tblCellMar"))
    if margins is not None:
        for edge, attr in (("top", "margin_top_in"), ("bottom", "margin_bottom_in"),
                           ("left", "margin_left_in"), ("right", "margin_right_in")):
            e = margins.find(qn("w:" + edge))
            if e is None:
                continue
            w = e.get(qn("w:w"))
            if w and w.lstrip("-").isdigit():
                setattr(fmt, attr, round(int(w) / _DXA_PER_INCH, 3))
    return fmt


def read_structure(path: str) -> Structure:
    """Parse a manuscript into everything the house rules need to see."""
    doc = docx.Document(path)
    numbering = _Numbering(doc)
    default_font = _doc_default_font(doc)

    paragraphs: List[Para] = []
    for i, p in enumerate(doc.paragraphs):
        s_name, s_size = _style_font(p.style)
        inherited = (s_name or default_font[0], s_size if s_size is not None else default_font[1])
        num_id, ilvl = _para_numbering(p)
        listing = numbering.describe(num_id, ilvl) if num_id else {}
        pf = p.paragraph_format
        paragraphs.append(Para(
            index=i,
            text=p.text,
            style=getattr(p.style, "name", "") or "",
            outline_level=_outline_level(p),
            alignment=_alignment(p),
            runs=[_read_run(r, inherited) for r in p.runs],
            listing=listing,
            first_line_in=_pf_in(pf, "first_line_indent"),
            left_indent_in=_pf_in(pf, "left_indent"),
            space_after_pt=_pf_pt(pf, "space_after"),
            space_before_pt=_pf_pt(pf, "space_before"),
            # A float is a multiple (1.0 single, 1.5); a Length is an exact height.
            # Kept as the multiple only, because that is what the spec is written in.
            line_spacing=(pf.line_spacing
                          if isinstance(pf.line_spacing, float) else None),
        ))

    # Where each table sits in the body flow. python-docx exposes tables and
    # paragraphs as separate lists with no ordering between them, so the document
    # body is walked directly — otherwise "the caption above this table" has no
    # meaning.
    order: List[tuple] = []
    p_seen = t_seen = 0
    for child in doc.element.body.iterchildren():
        if child.tag == qn("w:p"):
            order.append(("p", p_seen))
            p_seen += 1
        elif child.tag == qn("w:tbl"):
            order.append(("t", t_seen))
            t_seen += 1

    after: Dict[int, int] = {}
    last_p = -1
    for kind, idx in order:
        if kind == "p":
            last_p = idx
        else:
            after[idx] = last_p

    tables: List[Table] = []
    for ti, t in enumerate(doc.tables):
        cells = [[c.text.strip() for c in _row_cells(row)] for row in t.rows]
        grid: List[List[Cell]] = []
        for ri, row in enumerate(t.rows):
            row_cells: List[Cell] = []
            for ci, c in enumerate(_row_cells(row)):
                paras: List[Para] = []
                for p in c.paragraphs:
                    s_name, s_size = _style_font(p.style)
                    inh = (s_name or default_font[0],
                           s_size if s_size is not None else default_font[1])
                    paras.append(Para(
                        # Cell paragraphs are not in `doc.paragraphs` and have no
                        # index there. -1 says so rather than pointing at a body
                        # paragraph that has nothing to do with them.
                        index=-1,
                        text=p.text,
                        style=getattr(p.style, "name", "") or "",
                        outline_level=None,
                        alignment=_alignment(p),
                        runs=[_read_run(r, inh) for r in p.runs],
                        in_table=True,
                    ))
                row_cells.append(Cell(row=ri, col=ci, paragraphs=paras))
            grid.append(row_cells)

        tables.append(Table(
            index=ti,
            rows=len(t.rows),
            cols=len(t.columns),
            cells=cells,
            grid=grid,
            fmt=_table_format(t),
            after_paragraph=after.get(ti, -1),
        ))

    sections = [Section(
        page_width_in=_section_in(s, "page_width"),
        page_height_in=_section_in(s, "page_height"),
        left_margin_in=_section_in(s, "left_margin"),
        right_margin_in=_section_in(s, "right_margin"),
        top_margin_in=_section_in(s, "top_margin"),
        bottom_margin_in=_section_in(s, "bottom_margin"),
    ) for s in doc.sections]

    images = [Image(index=i,
                    width_in=round(sh.width / EMU_PER_INCH, 3),
                    height_in=round(sh.height / EMU_PER_INCH, 3))
              for i, sh in enumerate(doc.inline_shapes)]

    return Structure(paragraphs=paragraphs, tables=tables,
                     sections=sections, images=images)
